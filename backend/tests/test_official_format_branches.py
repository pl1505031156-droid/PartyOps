"""公文本机排版的异常路径、转换器和回环界面对抗测试。"""

from __future__ import annotations

import http.client
import io
import json
import os
import re
import socket
import sys
import threading
import time
import types
import uuid
import zipfile
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from lxml import etree

from app import official_format as formatter


def _docx_bytes(tmp_path: Path, *paragraphs: str) -> bytes:
    path = tmp_path / f"{uuid.uuid4().hex}.docx"
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)
    return path.read_bytes()


def _error_code(callable_: object) -> str:
    with pytest.raises(formatter.OfficialFormatError) as caught:
        callable_()  # type: ignore[operator]
    return caught.value.code


@pytest.mark.parametrize(
    ("text", "first_body", "role"),
    [
        ("公文标题", True, "title"),
        ("000001", True, "copy_number"),
        ("机密★1年", True, "security"),
        ("特急", True, "urgency"),
        ("中共某某委员会文件", True, "issuing_authority"),
        ("某办〔2026〕12号", True, "document_number"),
        ("签发人：张三", True, "signatory"),
        ("一、一级标题", False, "heading1"),
        ("（一）二级标题", False, "heading2"),
        ("1.三级标题", False, "heading3"),
        ("（1）四级标题", False, "heading4"),
        ("附件：材料", False, "attachment"),
        ("附件1", False, "attachment_heading"),
        ("某某单位：", False, "addressee"),
        ("抄送：有关单位。", False, "copy_recipient"),
        ("2026年8月24日印发", False, "imprint"),
        ("出席：张三、李四", False, "attendance"),
        ("（此件公开发布）", False, "note"),
        ("2026年8月24日", False, "date"),
        ("这是正文。", False, "body"),
        ("短句。", True, "body"),
    ],
)
def test_paragraph_role_matrix(text: str, first_body: bool, role: str) -> None:
    assert formatter._paragraph_role(text, first_body=first_body) == role


def test_contextual_classification_does_not_turn_document_header_into_title() -> None:
    root = etree.fromstring(
        f"""<w:document xmlns:w="{formatter.W}"><w:body>
        <w:p><w:r><w:t>000001</w:t></w:r></w:p>
        <w:p><w:r><w:t>机密★1年</w:t></w:r></w:p>
        <w:p><w:r><w:t>某某委员会文件</w:t></w:r></w:p>
        <w:p><w:r><w:t>某办〔2026〕12号</w:t></w:r></w:p>
        <w:p><w:r><w:t>关于开展工作的通知</w:t></w:r></w:p>
        <w:p><w:r><w:t>某某单位：</w:t></w:r></w:p>
        <w:p><w:r><w:t>请认真落实。</w:t></w:r></w:p>
        </w:body></w:document>""",
        parser=formatter.XML_PARSER,
    )
    roles = [
        role
        for _paragraph, role in formatter._classify_document_paragraphs(
            root.xpath(".//w:body/w:p", namespaces=formatter.NS)
        )
    ]
    assert roles == [
        "copy_number", "security", "issuing_authority", "document_number",
        "title", "addressee", "body",
    ]


@pytest.mark.parametrize(
    ("marker", "expected_role"),
    [("某某委员会令", "issuing_authority"), ("某某会议纪要", "issuing_authority")],
)
def test_order_and_minutes_headers_do_not_promote_following_body_to_title(
    marker: str, expected_role: str
) -> None:
    root = etree.fromstring(
        f"""<w:document xmlns:w="{formatter.W}"><w:body>
        <w:p><w:r><w:t>{marker}</w:t></w:r></w:p>
        <w:p><w:r><w:t>第12号</w:t></w:r></w:p>
        <w:p><w:r><w:t>请认真贯彻执行</w:t></w:r></w:p>
        <w:p/>
        </w:body></w:document>""",
        parser=formatter.XML_PARSER,
    )
    roles = [
        role
        for _paragraph, role in formatter._classify_document_paragraphs(
            root.xpath(".//w:body/w:p", namespaces=formatter.NS)
        )
    ]
    assert roles == [expected_role, "document_number", "body"]


def test_normal_style_creation_and_document_grid_diagnosis(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    assert formatter._configure_normal_style(None) == 0
    styles = etree.Element(formatter._qn("styles"), nsmap={"w": formatter.W})
    assert formatter._configure_normal_style(styles) == 1
    normal = styles.find(formatter._qn("style"))
    assert normal is not None and normal.get(formatter._qn("styleId")) == "Normal"

    source = tmp_path / "grid-source.docx"
    formatted = tmp_path / "grid-formatted.docx"
    document = Document()
    document.add_paragraph("标题")
    document.add_paragraph("正文内容。")
    document.save(source)
    monkeypatch.setattr(formatter, "_font_inventory", lambda: "方正小标宋 仿宋 楷体 黑体")
    formatter.format_docx(source, formatted)
    with zipfile.ZipFile(formatted, "a") as package:
        xml = etree.fromstring(package.read("word/document.xml"), parser=formatter.XML_PARSER)
        grid = xml.find(".//w:docGrid", namespaces=formatter.NS)
        assert grid is not None
        grid.set(formatter._qn("charSpace"), "0")
        package.writestr("word/document.xml", etree.tostring(xml, encoding="UTF-8"))
    warning = formatter.diagnose_docx(formatted, changed_count=0)
    error = formatter.diagnose_docx(formatted, changed_count=1)
    warning_issue = next(item for item in warning.issues if item.code == "DOCUMENT_GRID_INVALID")
    error_issue = next(item for item in error.issues if item.code == "DOCUMENT_GRID_INVALID")
    assert warning_issue.severity == "warning"
    assert error_issue.severity == "error"


def test_ooxml_safety_error_matrix(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    assert _error_code(lambda: formatter._safe_xml(b"<broken", "word/document.xml")) == "OOXML_XML_INVALID"
    assert _error_code(lambda: formatter._safe_xml(b"x" * (32 * 1024 * 1024 + 1), "word/document.xml")) == "OOXML_PART_TOO_LARGE"
    monkeypatch.setattr(formatter, "MAX_PACKAGE_BYTES", 8)
    expanded = tmp_path / "expanded.docx"
    with zipfile.ZipFile(expanded, "w") as package:
        package.writestr("[Content_Types].xml", b"12345")
        package.writestr("word/document.xml", b"12345")
    with zipfile.ZipFile(expanded) as package:
        assert _error_code(lambda: formatter._validated_members(package)) == "OOXML_EXPANSION_LIMIT"

    monkeypatch.setattr(formatter, "MAX_PACKAGE_BYTES", formatter.MAX_FILE_BYTES * 20)
    monkeypatch.setattr(formatter, "MAX_ZIP_RATIO", 2)
    compressed = tmp_path / "compressed.docx"
    with zipfile.ZipFile(compressed, "w", compression=zipfile.ZIP_DEFLATED) as package:
        package.writestr("[Content_Types].xml", b"<Types/>")
        package.writestr("word/document.xml", b"x" * 4096)
    with zipfile.ZipFile(compressed) as package:
        assert _error_code(lambda: formatter._validated_members(package)) == "OOXML_COMPRESSION_RATIO_UNSAFE"

    missing = tmp_path / "missing.docx"
    with zipfile.ZipFile(missing, "w") as package:
        package.writestr("[Content_Types].xml", b"<Types/>")
    with zipfile.ZipFile(missing) as package:
        assert _error_code(lambda: formatter._validated_members(package)) == "OOXML_REQUIRED_PART_MISSING"
    bad = tmp_path / "bad.docx"
    bad.write_bytes(b"not-a-zip")
    assert _error_code(lambda: formatter._read_core_parts(bad)) == "DOCX_PACKAGE_INVALID"


def test_diagnosis_reports_structural_review_items(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    path = tmp_path / "review.docx"
    document = Document()
    document.add_paragraph("标题")
    document.add_paragraph("正文没有预设样式")
    document.add_paragraph("1、需要人工确认")
    document.save(path)
    with zipfile.ZipFile(path, "a") as package:
        xml = etree.fromstring(package.read("word/document.xml"), parser=formatter.XML_PARSER)
        paragraph = xml.xpath(".//w:body/w:p", namespaces=formatter.NS)[1]
        text_box = etree.SubElement(paragraph, formatter._qn("txbxContent"))
        etree.SubElement(text_box, formatter._qn("p"))
        package.writestr(
            "word/document.xml",
            etree.tostring(xml, xml_declaration=True, encoding="UTF-8", standalone=True),
        )
    monkeypatch.setattr(formatter, "_font_inventory", lambda: "方正小标宋 仿宋 楷体")
    report = formatter.diagnose_docx(path, changed_count=1)
    codes = {issue.code for issue in report.issues}
    assert {"PAGE_GEOMETRY_INVALID", "TITLE_STYLE_INVALID", "BODY_STYLE_INVALID", "PAGE_NUMBER_MISSING", "TEXTBOX_REVIEW_REQUIRED", "NUMBERING_REVIEW_REQUIRED", "REQUIRED_FONT_MISSING"} <= codes
    assert report.compliant is False
    assert formatter.FormatReport(True, 1, 0, 0, ()).as_dict()["issues"] == []


def test_empty_document_and_missing_section_are_reported(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    path = tmp_path / "empty.docx"
    Document().save(path)
    with zipfile.ZipFile(path, "a") as package:
        xml = etree.fromstring(package.read("word/document.xml"), parser=formatter.XML_PARSER)
        for section in xml.xpath(".//w:sectPr", namespaces=formatter.NS):
            section.getparent().remove(section)
        package.writestr("word/document.xml", etree.tostring(xml, xml_declaration=True, encoding="UTF-8"))
    monkeypatch.setattr(formatter, "_font_inventory", lambda: "方正小标宋 仿宋 楷体 黑体")
    codes = {issue.code for issue in formatter.diagnose_docx(path).issues}
    assert "SECTION_MISSING" in codes
    assert "DOCUMENT_EMPTY" in codes


def test_font_inventory_and_issue_branches(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    class FontKey:
        def __enter__(self) -> "FontKey":
            return self

        def __exit__(self, *_args: object) -> None:
            return None

    def enum_value(_key: object, index: int) -> tuple[str, str, int]:
        if index:
            raise OSError
        return "FangSong SimHei Kaiti FZXiaoBiaoSong", "fonts.ttc", 1

    fake_winreg = SimpleNamespace(
        HKEY_LOCAL_MACHINE=object(),
        OpenKey=lambda *_args: FontKey(),
        EnumValue=enum_value,
    )
    monkeypatch.setitem(sys.modules, "winreg", fake_winreg)
    assert "fangsong" in formatter._font_inventory()
    assert formatter._font_issues() == []
    fake_winreg.OpenKey = lambda *_args: (_ for _ in ()).throw(OSError())
    assert formatter._font_inventory() == ""


def test_local_office_conversion_branches(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    source = tmp_path / "source.wps"
    source.write_bytes(b"wps")
    workspace = tmp_path / "work"
    workspace.mkdir()
    executable = tmp_path / "soffice"
    executable.write_bytes(b"binary")
    monkeypatch.setattr(formatter, "_office_candidates", lambda: [executable])

    observed: dict[str, object] = {}

    def successful_run(*args: object, **kwargs: object) -> SimpleNamespace:
        observed["command"] = args[0]
        observed["environment"] = kwargs["env"]
        (workspace / "source.docx").write_bytes(_docx_bytes(tmp_path, "标题"))
        return SimpleNamespace(returncode=0)

    monkeypatch.setattr(formatter.subprocess, "run", successful_run)
    assert formatter._convert_with_libreoffice(source, workspace) == workspace / "source.docx"
    assert os.environ.get("HTTP_PROXY") != ""
    command = observed["command"]
    environment = observed["environment"]
    assert isinstance(command, list) and "--safe-mode" in command
    assert isinstance(environment, dict)
    assert environment["HTTP_PROXY"] == "http://127.0.0.1:9"
    assert environment["NO_PROXY"] == "127.0.0.1,localhost"

    (workspace / "source.docx").unlink()
    monkeypatch.setattr(formatter.subprocess, "run", lambda *_args, **_kwargs: SimpleNamespace(returncode=1))
    assert _error_code(lambda: formatter._convert_with_libreoffice(source, workspace)) == "OFFICE_CONVERSION_FAILED"
    monkeypatch.setattr(formatter.subprocess, "run", lambda *_args, **_kwargs: (_ for _ in ()).throw(TimeoutError()))
    assert _error_code(lambda: formatter._convert_with_libreoffice(source, workspace)) == "OFFICE_CONVERSION_FAILED"
    monkeypatch.setattr(formatter, "_office_candidates", lambda: [])
    assert formatter._convert_with_libreoffice(source, workspace) is None


def test_office_candidate_and_windows_com_conversion_matrix(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    executable = tmp_path / "soffice.exe"
    executable.write_bytes(b"binary")
    program_files = tmp_path / "Program Files"
    installed = program_files / "LibreOffice" / "program" / "soffice.exe"
    installed.parent.mkdir(parents=True)
    installed.write_bytes(b"binary")
    monkeypatch.setattr(formatter.shutil, "which", lambda name: str(executable) if name == "soffice" else None)
    monkeypatch.setenv("PROGRAMFILES", str(program_files))
    monkeypatch.delenv("PROGRAMFILES(X86)", raising=False)
    candidates = formatter._office_candidates()
    assert executable in candidates and installed in candidates
    with monkeypatch.context() as patcher:
        patcher.setattr(formatter.os, "name", "posix")
        patcher.setattr(formatter.shutil, "which", lambda _name: None)
        assert formatter._office_candidates() == []

    source = tmp_path / "legacy.doc"
    source.write_bytes(b"legacy")
    workspace = tmp_path / "com"
    workspace.mkdir()

    class FakeDocument:
        def __init__(self, *, fallback: bool = False, close_error: bool = False) -> None:
            self.fallback = fallback
            self.close_error = close_error

        def SaveAs2(self, target: str, **_kwargs: object) -> None:
            if self.fallback:
                raise AttributeError
            Path(target).write_bytes(_docx_bytes(tmp_path, "标题"))

        def SaveAs(self, target: str, **_kwargs: object) -> None:
            Path(target).write_bytes(_docx_bytes(tmp_path, "标题"))

        def Close(self, _save: bool) -> None:
            if self.close_error:
                raise RuntimeError

    class FakeApplication:
        def __init__(self, document: FakeDocument | None, *, quit_error: bool = False) -> None:
            self.Visible = True
            self.DisplayAlerts = 1
            self.document = document
            self.quit_error = quit_error
            self.Documents = self

        def Open(self, *_args: object, **_kwargs: object) -> FakeDocument:
            if self.document is None:
                raise RuntimeError("open failed")
            return self.document

        def Quit(self) -> None:
            if self.quit_error:
                raise RuntimeError

    dispatches = iter(
        [
            RuntimeError("word unavailable"),
            FakeApplication(FakeDocument()),
        ]
    )

    def dispatch(_program_id: str) -> FakeApplication:
        value = next(dispatches)
        if isinstance(value, BaseException):
            raise value
        return value

    client_module = types.ModuleType("win32com.client")
    client_module.DispatchEx = dispatch  # type: ignore[attr-defined]
    package_module = types.ModuleType("win32com")
    package_module.client = client_module  # type: ignore[attr-defined]
    monkeypatch.setitem(sys.modules, "win32com", package_module)
    monkeypatch.setitem(sys.modules, "win32com.client", client_module)
    target = formatter._convert_with_windows_office(source, workspace)
    assert target is not None and target.is_file()

    target.unlink()
    client_module.DispatchEx = lambda _program_id: FakeApplication(FakeDocument(fallback=True))  # type: ignore[attr-defined]
    assert formatter._convert_with_windows_office(source, workspace) == target

    target.unlink()
    client_module.DispatchEx = lambda _program_id: FakeApplication(None, quit_error=True)  # type: ignore[attr-defined]
    assert formatter._convert_with_windows_office(source, workspace) is None
    with monkeypatch.context() as patcher:
        patcher.setattr(formatter.os, "name", "posix")
        assert formatter._convert_with_windows_office(source, workspace) is None


def test_posix_font_inventory_and_private_write_cleanup(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePath:
        def __init__(self, value: str) -> None:
            self.value = value
            self.name = value.rstrip("/").rsplit("/", 1)[-1]

        @classmethod
        def home(cls) -> "FakePath":
            return cls("/home/user")

        def __truediv__(self, child: str) -> "FakePath":
            return FakePath(f"{self.value}/{child}")

        def is_dir(self) -> bool:
            return self.value.endswith("Fonts")

        def rglob(self, _pattern: str) -> list["FakePath"]:
            return [FakePath(f"{self.value}/Example.ttf")]

    with monkeypatch.context() as patcher:
        patcher.setattr(formatter.os, "name", "posix")
        patcher.setattr(formatter, "Path", FakePath)
        patcher.setattr(formatter.shutil, "which", lambda _name: None)
        inventory = formatter._font_inventory()
    assert "example.ttf" in inventory

    path = tmp_path / "failed.bin"

    class BrokenStream:
        def __enter__(self) -> "BrokenStream":
            return self

        def __exit__(self, *_args: object) -> None:
            os.close(descriptor)
            return None

        def write(self, _payload: bytes) -> None:
            raise OSError("write failed")

    descriptor = os.open(path, os.O_WRONLY | os.O_CREAT | os.O_EXCL, 0o600)
    monkeypatch.setattr(formatter.os, "open", lambda *_args, **_kwargs: descriptor)
    monkeypatch.setattr(formatter.os, "fdopen", lambda *_args, **_kwargs: BrokenStream())
    with pytest.raises(OSError, match="write failed"):
        formatter._private_write(path, b"content")
    assert not path.exists()


def test_footer_reuse_rejects_unsafe_relationships_and_avoids_duplicates(tmp_path: Path) -> None:
    document = etree.fromstring(
        f'<w:document xmlns:w="{formatter.W}" xmlns:r="{formatter.R}"><w:body><w:sectPr>'
        '<w:footerReference w:type="default" r:id="rBad"/>'
        '<w:footerReference w:type="default" r:id="rMissing"/>'
        '<w:footerReference w:type="default" r:id="rGood"/>'
        '</w:sectPr></w:body></w:document>'.encode(),
        parser=formatter.XML_PARSER,
    )
    settings = etree.Element(formatter._qn("settings"), nsmap={"w": formatter.W})
    relationships = etree.Element(f"{{{formatter.PR}}}Relationships", nsmap={None: formatter.PR})
    for rel_id, target in (
        ("rIdPartyOpsFooter1", "occupied.xml"),
        ("rBad", "../outside.xml"),
        ("rMissing", "word/missing.xml"),
        ("rGood", "footer-existing.xml"),
    ):
        node = etree.SubElement(relationships, f"{{{formatter.PR}}}Relationship")
        node.set("Id", rel_id)
        node.set("Type", f"{formatter.R}/footer")
        node.set("Target", target)
    content_types = etree.Element(f"{{{formatter.CT}}}Types", nsmap={None: formatter.CT})
    override = etree.SubElement(content_types, f"{{{formatter.CT}}}Override")
    override.set("PartName", "/word/footer-partyops-even.xml")
    override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml")
    package_path = tmp_path / "footers.zip"
    with zipfile.ZipFile(package_path, "w") as output:
        output.writestr(
            "word/footer-existing.xml",
            f'<w:ftr xmlns:w="{formatter.W}"><w:p><w:r><w:instrText> PAGE </w:instrText></w:r></w:p></w:ftr>',
        )
    with zipfile.ZipFile(package_path) as package:
        outputs = formatter._add_page_footers(document, settings, relationships, content_types, package)
    assert "word/footer-existing.xml" in outputs
    assert "word/footer-partyops-even.xml" in outputs
    normalized = formatter._safe_xml(outputs["word/footer-existing.xml"], "word/footer-existing.xml")
    assert formatter._standard_page_footer_paragraph(formatter._page_field_paragraphs(normalized)[0])
    assert len(content_types.findall(f"{{{formatter.CT}}}Override")) == 1
    relationship_ids = {item.get("Id") for item in relationships}
    assert "rIdPartyOpsFooter2" in relationship_ids


def test_prepare_docx_and_private_write_branches(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    docx = tmp_path / "source.docx"
    docx.write_bytes(_docx_bytes(tmp_path, "标题"))
    assert formatter.prepare_docx(docx, tmp_path) == (docx, False)
    legacy = tmp_path / "source.doc"
    legacy.write_bytes(b"doc")
    converted = tmp_path / "converted.docx"
    converted.write_bytes(docx.read_bytes())
    monkeypatch.setattr(formatter, "_convert_with_libreoffice", lambda *_args: converted)
    assert formatter.prepare_docx(legacy, tmp_path) == (converted, True)
    monkeypatch.setattr(formatter, "_convert_with_libreoffice", lambda *_args: None)
    monkeypatch.setattr(formatter, "_convert_with_windows_office", lambda *_args: converted)
    assert formatter.prepare_docx(legacy, tmp_path) == (converted, True)
    monkeypatch.setattr(formatter, "_convert_with_windows_office", lambda *_args: None)
    assert _error_code(lambda: formatter.prepare_docx(legacy, tmp_path)) == "OFFICE_SUITE_REQUIRED"

    private = tmp_path / "private.bin"
    formatter._private_write(private, b"content")
    assert private.read_bytes() == b"content"
    assert formatter._safe_stem("../<>bad name?.docx") == "bad name"
    assert formatter._safe_stem("...docx") == "公文"


class _UploadHandler:
    def __init__(self, content_type: str, body: bytes, length: str | None = None) -> None:
        self.headers = {
            "Content-Type": content_type,
            "Content-Length": length if length is not None else str(len(body)),
        }
        self.rfile = io.BytesIO(body)


def test_upload_parser_error_and_success_matrix(monkeypatch: pytest.MonkeyPatch) -> None:
    boundary = "PartyOpsBoundary"
    valid = (
        f"--{boundary}\r\nContent-Disposition: form-data; name=\"ignored\"\r\n\r\nx\r\n"
        f"--{boundary}\r\nContent-Disposition: form-data; name=\"document\"; filename=\"文件.docx\"\r\n"
        "Content-Type: application/octet-stream\r\n\r\npayload\r\n"
        f"--{boundary}--\r\n"
    ).encode()
    assert formatter._extract_upload(_UploadHandler(f'multipart/form-data; boundary="{boundary}"', valid)) == ("文件.docx", b"payload")
    assert _error_code(lambda: formatter._extract_upload(_UploadHandler("multipart/form-data", b"x"))) == "UPLOAD_FORMAT_INVALID"
    assert _error_code(lambda: formatter._extract_upload(_UploadHandler("multipart/form-data; boundary=x", b"", "bad"))) == "UPLOAD_LENGTH_INVALID"
    assert _error_code(lambda: formatter._extract_upload(_UploadHandler("multipart/form-data; boundary=x", b"", "0"))) == "FILE_SIZE_LIMIT"
    missing = b"--x\r\nContent-Disposition: form-data; name=\"other\"\r\n\r\nx\r\n--x--\r\n"
    assert _error_code(lambda: formatter._extract_upload(_UploadHandler("multipart/form-data; boundary=x", missing))) == "UPLOAD_FILE_MISSING"
    unsupported = b"--x\r\nContent-Disposition: form-data; name=\"document\"; filename=\"bad.xlsm\"\r\n\r\nx\r\n--x--\r\n"
    assert _error_code(lambda: formatter._extract_upload(_UploadHandler("multipart/form-data; boundary=x", unsupported))) == "FORMAT_UNSUPPORTED"
    missing_separator = b"--x\r\nContent-Disposition: form-data; name=\"document\"; filename=\"bad.docx\"\r\n--x--\r\n"
    assert _error_code(lambda: formatter._extract_upload(_UploadHandler("multipart/form-data; boundary=x", missing_separator))) == "UPLOAD_FILE_MISSING"
    monkeypatch.setattr(formatter, "MAX_FILE_BYTES", 1)
    assert _error_code(lambda: formatter._extract_upload(_UploadHandler(f"multipart/form-data; boundary={boundary}", valid))) == "FILE_SIZE_LIMIT"


def test_stage_log_rotation_and_html_are_content_safe(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    log = tmp_path / "official-format.log"
    log.write_bytes(b"x" * (512 * 1024 + 1))
    formatter._append_stage_log(tmp_path, "diagnose", time.monotonic(), "OK")
    assert (tmp_path / "official-format.log.1").is_file()
    record = json.loads(log.read_text(encoding="utf-8"))
    assert set(record) == {"version", "stage", "duration_ms", "result_code"}
    issue = formatter.FormatIssue("<code>", "warning", "<标题>", "<细节>", "5.1")
    escaped = formatter._escape_issue(issue)
    assert "&lt;标题&gt;" in escaped and "<标题>" not in escaped
    body = formatter._report_body("<token>", "a" * 32, formatter.FormatReport(True, 1, 0, 0, ()), formatted=False)
    assert "按 GB/T 9704-2012 一键排版" in body
    formatted_body = formatter._report_body("token", "a" * 32, formatter.FormatReport(False, 1, 0, 2, (issue,)), formatted=True)
    assert "公文规范版" in formatted_body and "存在阻断项" in formatted_body

    monkeypatch.setattr(formatter.Path, "open", lambda *_args, **_kwargs: (_ for _ in ()).throw(OSError()))
    formatter._append_stage_log(tmp_path, "failure", time.monotonic(), "IO")


def _wait_formatter_url(config_dir: Path) -> str:
    marker = config_dir / "official-format.url"
    for _ in range(100):
        if marker.is_file():
            return marker.read_text(encoding="utf-8").strip()
        time.sleep(0.05)
    raise AssertionError("本机排版服务未启动")


def _request(url: str, method: str, path: str, body: bytes | None = None, headers: dict[str, str] | None = None) -> tuple[int, bytes]:
    parsed = formatter.urllib.parse.urlsplit(url)
    connection = http.client.HTTPConnection(parsed.hostname, parsed.port, timeout=10)
    connection.request(method, path, body=body, headers=headers or {})
    response = connection.getresponse()
    payload = response.read()
    status = response.status
    connection.close()
    return status, payload


def test_formatter_loopback_end_to_end(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    config_dir = tmp_path / "config"
    token = str(uuid.uuid4())
    errors: list[BaseException] = []
    monkeypatch.setattr(formatter, "_font_inventory", lambda: "方正小标宋 仿宋 楷体 黑体")

    def target() -> None:
        try:
            formatter.run_official_format_tool(token, open_browser=False, config_dir=config_dir)
        except BaseException as exc:  # noqa: BLE001 - 测试线程需要回传异常。
            errors.append(exc)

    thread = threading.Thread(target=target)
    thread.start()
    url = _wait_formatter_url(config_dir)
    status, _ = _request(url, "GET", "/")
    assert status == 403
    query = formatter.urllib.parse.urlsplit(url).query
    status, page = _request(url, "GET", f"/?{query}")
    assert status == 200 and "公文规范排版".encode() in page
    assert _request(url, "GET", f"/missing?{query}")[0] == 404
    assert _request(url, "POST", "/unknown", b"")[0] == 403
    assert _request(url, "POST", f"/unknown?{query}", b"")[0] == 404

    payload = _docx_bytes(tmp_path, "关于开展工作的通知", "请认真落实,按期完成。")
    boundary = "PartyOpsE2E"
    body = (
        f"--{boundary}\r\nContent-Disposition: form-data; name=\"document\"; filename=\"sample.docx\"\r\n"
        "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
    ).encode() + payload + f"\r\n--{boundary}--\r\n".encode()
    status, diagnose_page = _request(
        url,
        "POST",
        f"/diagnose?{query}",
        body,
        {"Content-Type": f"multipart/form-data; boundary={boundary}", "Content-Length": str(len(body))},
    )
    assert status == 200
    match = re.search(br"/format/([0-9a-f]{32})", diagnose_page)
    assert match
    document_id = match.group(1).decode()
    assert _request(url, "GET", f"/download/{document_id}?{query}")[0] == 410
    status, formatted_page = _request(url, "POST", f"/format/{document_id}?{query}", b"")
    assert status == 200 and "排版后复核".encode() in formatted_page
    status, download = _request(url, "GET", f"/download/{document_id}?{query}")
    assert status == 200 and download.startswith(b"PK")
    thread.join(8)
    assert not thread.is_alive() and errors == []
    assert not (config_dir / "official-format.url").exists()
    log = (config_dir / "official-format.log").read_text(encoding="utf-8")
    assert "sample.docx" not in log and "diagnose" in log and "download" in log


def test_formatter_cancel_invalid_transaction_and_browser_failure(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    assert _error_code(lambda: formatter.run_official_format_tool("bad", open_browser=False, config_dir=tmp_path)) == "FORMAT_TRANSACTION_INVALID"
    token = str(uuid.uuid4())
    thread = threading.Thread(
        target=lambda: formatter.run_official_format_tool(token, open_browser=False, config_dir=tmp_path / "cancel")
    )
    thread.start()
    url = _wait_formatter_url(tmp_path / "cancel")
    query = formatter.urllib.parse.urlsplit(url).query
    payload = _docx_bytes(tmp_path, "标题", "正文内容。")
    boundary = "PartyOpsCancel"
    upload = (
        f"--{boundary}\r\nContent-Disposition: form-data; name=\"document\"; filename=\"cancel.docx\"\r\n\r\n"
    ).encode() + payload + f"\r\n--{boundary}--\r\n".encode()
    status, diagnose_page = _request(
        url,
        "POST",
        f"/diagnose?{query}",
        upload,
        {"Content-Type": f"multipart/form-data; boundary={boundary}", "Content-Length": str(len(upload))},
    )
    assert status == 200
    matched = re.search(br"/format/([0-9a-f]{32})", diagnose_page)
    assert matched
    document_id = matched.group(1).decode()
    assert _request(url, "POST", f"/format/{document_id}?{query}", b"")[0] == 200
    status, body = _request(url, "POST", f"/cancel?{query}", b"")
    assert status == 200 and "临时文件已清理".encode() in body
    thread.join(8)
    assert not thread.is_alive()

    browser_config = tmp_path / "browser"

    def fail_browser(_url: str, **_kwargs: object) -> bool:
        (browser_config / "official-format.url").write_text("different\n", encoding="utf-8")
        return False

    monkeypatch.setattr(formatter.webbrowser, "open", fail_browser)
    assert _error_code(
        lambda: formatter.run_official_format_tool(
            str(uuid.uuid4()), open_browser=True, config_dir=browser_config
        )
    ) == "BROWSER_OPEN_FAILED"
    assert (browser_config / "official-format.url").read_text(encoding="utf-8").strip() == "different"


def test_formatter_idle_timeout_stops_without_opening_browser(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(formatter, "IDLE_TIMEOUT_SECONDS", 0)
    assert formatter.run_official_format_tool(
        str(uuid.uuid4()), open_browser=False, config_dir=tmp_path / "idle"
    ) == 0


def test_loopback_guard_covers_localhost_ipv6_and_raw_socket() -> None:
    restore = formatter._install_loopback_only_network_guard()
    try:
        assert formatter.socket.create_connection
        with pytest.raises(OSError, match="OFFICIAL_FORMAT_NETWORK_DENIED"):
            formatter.socket.create_connection("bad-address")
        with pytest.raises(OSError) as localhost:
            formatter.socket.create_connection(("localhost", 9), timeout=0.01)
        assert "OFFICIAL_FORMAT_NETWORK_DENIED" not in str(localhost.value)
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        try:
            with pytest.raises(OSError, match="OFFICIAL_FORMAT_NETWORK_DENIED"):
                sock.connect(("8.8.8.8", 443))
            with pytest.raises(OSError) as local:
                sock.connect(("127.0.0.1", 9))
            assert "OFFICIAL_FORMAT_NETWORK_DENIED" not in str(local.value)
        finally:
            sock.close()
        if hasattr(socket, "AF_UNIX"):
            unix_socket = socket.socket(socket.AF_UNIX, socket.SOCK_STREAM)
            try:
                with pytest.raises(OSError) as local_transport:
                    unix_socket.connect("partyops-nonexistent.sock")
                assert "OFFICIAL_FORMAT_NETWORK_DENIED" not in str(local_transport.value)
            finally:
                unix_socket.close()
    finally:
        restore()

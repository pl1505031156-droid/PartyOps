"""多格式正文提取、资源边界与恶意容器的发布回归。"""

from __future__ import annotations

import io
import tarfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from fastapi import UploadFile
from PIL import Image

from app import intake
from app.problems import ProblemException


def _zip_bytes(entries: dict[str, bytes]) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, value in entries.items():
            archive.writestr(name, value)
    return buffer.getvalue()


def test_docx_tables_and_office_container_guards(monkeypatch) -> None:
    document = Document()
    document.add_paragraph("关于开展年度复核")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "报送材料"
    table.cell(0, 1).text = "正式稿"
    buffer = io.BytesIO()
    document.save(buffer)
    text = intake._extract_docx(buffer.getvalue())
    assert "关于开展年度复核" in text
    assert "报送材料 | 正式稿" in text

    valid = _zip_bytes({"word/document.xml": b"<document/>"})
    monkeypatch.setattr(intake, "MAX_OFFICE_MEMBERS", 0)
    with pytest.raises(ProblemException) as complex_file:
        intake._validate_office_container(valid)
    assert complex_file.value.code == "INTAKE_OFFICE_TOO_COMPLEX"

    monkeypatch.setattr(intake, "MAX_OFFICE_MEMBERS", 10)
    monkeypatch.setattr(intake, "MAX_OFFICE_UNCOMPRESSED_BYTES", 0)
    with pytest.raises(ProblemException) as expansion:
        intake._validate_office_container(valid)
    assert expansion.value.code == "INTAKE_OFFICE_EXPANSION_LIMIT"

    monkeypatch.setattr(intake, "MAX_OFFICE_UNCOMPRESSED_BYTES", 1024 * 1024)
    unsafe = _zip_bytes({"word/document.xml": b"<!DOCTYPE x [<!ENTITY a 'b'>]><x/>"})
    with pytest.raises(ProblemException) as entity:
        intake._validate_office_container(unsafe)
    assert entity.value.code == "INTAKE_OFFICE_XML_UNSAFE"
    with pytest.raises(ProblemException) as invalid:
        intake._validate_office_container(b"not-a-zip")
    assert invalid.value.code == "INTAKE_OFFICE_INVALID"


class _FakePixmap:
    def tobytes(self, _format):
        image = Image.new("RGB", (2, 2), color="white")
        buffer = io.BytesIO()
        image.save(buffer, format="PNG")
        return buffer.getvalue()


class _FakePage:
    def __init__(self, text="", width=100, height=100):
        self._text = text
        self.rect = SimpleNamespace(width=width, height=height)

    def get_text(self, _kind):
        return self._text

    def get_pixmap(self, **_kwargs):
        return _FakePixmap()


class _FakePdf:
    def __init__(self, pages, *, encrypted=False, page_count=None):
        self.pages = pages
        self.needs_pass = encrypted
        self.page_count = len(pages) if page_count is None else page_count

    def __enter__(self):
        return self

    def __exit__(self, *_args):
        return False

    def __iter__(self):
        return iter(self.pages)


def test_pdf_ocr_limits_encryption_pixels_and_engine_errors(monkeypatch) -> None:
    monkeypatch.setattr(intake.fitz, "open", lambda **_kwargs: _FakePdf([], encrypted=True))
    with pytest.raises(ProblemException) as encrypted:
        intake._extract_pdf(b"pdf")
    assert encrypted.value.code == "INTAKE_PDF_ENCRYPTED"

    monkeypatch.setattr(intake, "MAX_PDF_OCR_PAGES", 0)
    monkeypatch.setattr(intake.fitz, "open", lambda **_kwargs: _FakePdf([_FakePage(), _FakePage()]))
    text, warnings = intake._extract_pdf(b"pdf")
    assert not text and len(warnings) == 1 and "停止继续 OCR" in warnings[0]

    monkeypatch.setattr(intake, "MAX_PDF_OCR_PAGES", 20)
    monkeypatch.setattr(intake, "MAX_IMAGE_PIXELS", 10)
    text, warnings = intake._extract_pdf(b"pdf")
    assert not text and "像素规模" in warnings[0]

    monkeypatch.setattr(intake, "MAX_IMAGE_PIXELS", 50_000_000)
    monkeypatch.setattr(intake.pytesseract, "image_to_string", lambda *_args, **_kwargs: "识别正文")
    text, warnings = intake._extract_pdf(b"pdf")
    assert "识别正文" in text and not warnings

    def unavailable(*_args, **_kwargs):
        raise intake.pytesseract.TesseractNotFoundError()

    monkeypatch.setattr(intake.pytesseract, "image_to_string", unavailable)
    text, warnings = intake._extract_pdf(b"pdf")
    assert not text and len(warnings) == 1 and "Tesseract" in warnings[0]


def test_image_ocr_engine_absence_and_runtime_failure(monkeypatch) -> None:
    image = Image.new("RGB", (4, 4), color="white")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")

    monkeypatch.setattr(
        intake.pytesseract,
        "image_to_string",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(intake.pytesseract.TesseractNotFoundError()),
    )
    assert "Tesseract" in intake._extract_image(buffer.getvalue())[1][0]
    monkeypatch.setattr(
        intake.pytesseract,
        "image_to_string",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(RuntimeError("timeout")),
    )
    assert "无法识别" in intake._extract_image(buffer.getvalue())[1][0]


def test_parse_upload_data_dispatch_fallback_and_filename(monkeypatch) -> None:
    monkeypatch.setattr(intake, "_extract_pdf", lambda _data: ("PDF 正文", ["需复核"]))
    pdf = intake._parse_upload_data(b"pdf", "Notice.PDF", "粘贴说明")
    assert pdf.source_kind == "pdf" and pdf.parser_label == "PDF 文本与中文 OCR"
    assert "需复核" in pdf.warnings and pdf.source_filename == "Notice.PDF"

    monkeypatch.setattr(intake, "_extract_image", lambda _data: ("图片正文", []))
    image = intake._parse_upload_data(b"png", "scan.PNG", "")
    assert image.source_kind == "image" and image.title == "图片正文"

    unknown = intake._parse_upload_data(b"\xff\xfe\xfd", "opaque.bin", "")
    assert unknown.title == "opaque"
    assert any("暂不支持" in warning for warning in unknown.warnings)

    monkeypatch.setattr(intake, "_extract_docx", lambda _data: (_ for _ in ()).throw(ValueError("bad")))
    broken = intake._parse_upload_data(b"bad", "broken.docx", "")
    assert broken.title == "broken"
    assert any("识别未完成" in warning for warning in broken.warnings)


@pytest.mark.asyncio
async def test_parse_upload_size_boundary(monkeypatch) -> None:
    monkeypatch.setattr(intake, "MAX_INTAKE_UPLOAD_BYTES", 3)
    upload = UploadFile(filename="large.txt", file=io.BytesIO(b"1234"))
    with pytest.raises(ProblemException) as large:
        await intake.parse_upload(upload)
    assert large.value.code == "INTAKE_FILE_TOO_LARGE"

    upload = UploadFile(filename="folder/notice.txt", file=io.BytesIO(b"notice"))
    monkeypatch.setattr(intake, "MAX_INTAKE_UPLOAD_BYTES", 10)
    parsed = await intake.parse_upload(upload, "请于明天反馈")
    assert parsed.source_filename == "notice.txt"


def test_type_detection_for_supported_signatures(tmp_path) -> None:
    signatures = {
        "archive.bin": (b"PK\x03\x04payload", "application/zip"),
        "document.bin": (b"%PDF-1.7", "application/pdf"),
        "legacy.bin": (b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1", "application/x-ole-storage"),
        "image.bin": (b"\x89PNG\r\n", "image/png"),
        "photo.bin": (b"\xFF\xD8\xFF", "image/jpeg"),
        "scan.bin": (b"II*\x00", "image/tiff"),
        "plain.txt": (b"hello", "text/plain"),
        "unknown.bin": (b"unknown", "application/octet-stream"),
    }
    for name, (payload, expected) in signatures.items():
        path = tmp_path / name
        path.write_bytes(payload)
        assert intake._detect_type(path, path.suffix) == expected


def test_spreadsheet_and_presentation_extractors(monkeypatch, tmp_path) -> None:
    class _Sheet:
        title = "清单"

        @staticmethod
        def iter_rows(values_only=True):
            assert values_only
            return iter([("标题", "内容"), ("年度任务", "A" * 100_001)])

    closed: list[bool] = []
    workbook = SimpleNamespace(worksheets=[_Sheet()], close=lambda: closed.append(True))
    monkeypatch.setattr(intake, "_validate_office_container", lambda _data: None)
    monkeypatch.setattr(intake, "load_workbook", lambda *_args, **_kwargs: workbook)
    sheet_path = tmp_path / "sheet.xlsx"
    sheet_path.write_bytes(b"placeholder")
    text, warnings = intake._extract_spreadsheet(sheet_path)
    assert "[清单]" in text and "10 万字" in warnings[0] and closed == [True]

    presentation = tmp_path / "slides.pptx"
    xml = (
        '<p:sld xmlns:p="urn:p" xmlns:a="urn:a"><p:cSld><a:t>第一页</a:t>'
        '<a:t>部署说明</a:t></p:cSld></p:sld>'
    ).encode("utf-8")
    with zipfile.ZipFile(presentation, "w") as archive:
        archive.writestr("ppt/slides/slide1.xml", xml)
        archive.writestr("ppt/slides/slide2.xml", b'<p:sld xmlns:p="urn:p"/>')
    value, warnings = intake._extract_presentation(presentation)
    assert value == "第一页 部署说明" and not warnings


def test_zip_and_tar_listing_apply_path_and_capacity_guards(monkeypatch, tmp_path) -> None:
    archive_path = tmp_path / "team.zip"
    with zipfile.ZipFile(archive_path, "w") as archive:
        archive.writestr("docs/report.txt", "正文")
        archive.writestr("../escape.txt", "拒绝")
    value, warnings, count = intake._extract_zip_listing(archive_path)
    assert value == "docs/report.txt" and not warnings and count == 1

    monkeypatch.setattr(intake, "MAX_ARCHIVE_MEMBERS", 0)
    assert "数量" in intake._extract_zip_listing(archive_path)[1][0]
    monkeypatch.setattr(intake, "MAX_ARCHIVE_MEMBERS", 5_000)
    monkeypatch.setattr(intake, "MAX_ARCHIVE_UNCOMPRESSED_BYTES", 0)
    assert "容量" in intake._extract_zip_listing(archive_path)[1][0]
    monkeypatch.setattr(intake, "MAX_ARCHIVE_UNCOMPRESSED_BYTES", 2 * 1024 * 1024 * 1024)
    monkeypatch.setattr(intake, "MAX_ARCHIVE_COMPRESSION_RATIO", 0)
    assert "压缩比" in intake._extract_zip_listing(archive_path)[1][0]

    tar_path = tmp_path / "team.tar"
    with tarfile.open(tar_path, "w") as archive:
        info = tarfile.TarInfo("docs/manual.txt")
        payload = b"manual"
        info.size = len(payload)
        archive.addfile(info, io.BytesIO(payload))
        escape = tarfile.TarInfo("../escape.txt")
        escape.size = 0
        archive.addfile(escape, io.BytesIO())
    monkeypatch.setattr(intake, "MAX_ARCHIVE_COMPRESSION_RATIO", 200)
    value, warnings, count = intake._extract_tar_listing(tar_path)
    assert value == "docs/manual.txt" and not warnings and count == 1
    monkeypatch.setattr(intake, "MAX_ARCHIVE_MEMBERS", 0)
    assert "数量" in intake._extract_tar_listing(tar_path)[1][0]
    monkeypatch.setattr(intake, "MAX_ARCHIVE_MEMBERS", 5_000)
    monkeypatch.setattr(intake, "MAX_ARCHIVE_UNCOMPRESSED_BYTES", 0)
    assert "容量" in intake._extract_tar_listing(tar_path)[1][0]


def test_extract_path_content_routes_formats_and_statuses(monkeypatch, tmp_path) -> None:
    txt = tmp_path / "gb.txt"
    txt.write_bytes("协同正文".encode("gb18030"))
    assert intake.extract_path_content(txt).text == "协同正文"

    binary_text = tmp_path / "broken.txt"
    binary_text.write_bytes(b"\x81")
    metadata = intake.extract_path_content(binary_text)
    assert metadata.content_status == "metadata_only"

    zip_path = tmp_path / "folder.zip"
    with zipfile.ZipFile(zip_path, "w") as archive:
        archive.writestr("one.txt", "1")
    zipped = intake.extract_path_content(zip_path)
    assert zipped.content_status == "indexed" and zipped.archive_member_count == 1

    tar_path = tmp_path / "folder.tar"
    with tarfile.open(tar_path, "w") as archive:
        info = tarfile.TarInfo("one.txt")
        info.size = 1
        archive.addfile(info, io.BytesIO(b"1"))
    tarred = intake.extract_path_content(tar_path)
    assert tarred.detected_type == "application/x-tar" and tarred.archive_member_count == 1

    ocr = tmp_path / "scan.png"
    ocr.write_bytes(b"\x89PNG\r\n")
    monkeypatch.setattr(intake, "_extract_image", lambda _data: ("", ["Tesseract OCR 待安装"]))
    assert intake.extract_path_content(ocr).content_status == "pending_ocr"

    monkeypatch.setattr(intake, "_extract_image", lambda _data: ("", []))
    assert intake.extract_path_content(ocr).content_status == "metadata_only"
    monkeypatch.setattr(intake, "_extract_image", lambda _data: (_ for _ in ()).throw(TypeError("bad parser")))
    assert intake.extract_path_content(ocr).error_code == "CONTENT_PARSE_FAILED"

    assert intake.extract_path_text(txt)[0] == "协同正文"

"""rc.6 党委会文件导入：真实容器识别、人工确认和正式入账。"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document

from app import meeting_imports
from app.problems import ProblemException


def _agenda_docx() -> bytes:
    document = Document()
    document.add_paragraph("中共测试镇委员会2026年第8次党委会议程")
    document.add_paragraph("会议时间：2026年8月28日下午18:00")
    document.add_paragraph("会议地点：第一会议室")
    document.add_paragraph("主持人：张三")
    document.add_paragraph("参会人员：张三、李四")
    document.add_paragraph("一、研究基层党建重点任务")
    document.add_paragraph("二、审议专项资金使用方案")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def test_meeting_import_requires_real_office_container() -> None:
    with pytest.raises(ProblemException) as captured:
        meeting_imports.detect_meeting_container(b"not-an-office-document")
    assert captured.value.code == "MEETING_IMPORT_CONTAINER_UNSUPPORTED"
    assert meeting_imports.detect_meeting_container(meeting_imports.OLE_HEADER + b"\0" * 16) == "ole"
    assert meeting_imports.detect_meeting_container(_agenda_docx()) == "ooxml"


def test_meeting_import_rejects_empty_and_oversized_payload(monkeypatch) -> None:
    with pytest.raises(ProblemException) as empty:
        meeting_imports.extract_meeting_text(b"")
    assert empty.value.code == "MEETING_IMPORT_SIZE_LIMIT"
    monkeypatch.setattr(meeting_imports, "MAX_MEETING_IMPORT_BYTES", 4)
    with pytest.raises(ProblemException) as oversized:
        meeting_imports.extract_meeting_text(b"PK\x03\x04too-large")
    assert oversized.value.code == "MEETING_IMPORT_SIZE_LIMIT"


def test_ole_conversion_error_is_preserved(monkeypatch) -> None:
    def fail_prepare(*_args):
        raise meeting_imports.OfficialFormatError("OFFICE_SUITE_REQUIRED", "缺少办公套件", "需要本机转换器")

    monkeypatch.setattr(meeting_imports, "prepare_docx", fail_prepare)
    with pytest.raises(ProblemException) as captured:
        meeting_imports.extract_meeting_text(meeting_imports.OLE_HEADER + b"\0" * 32)
    assert captured.value.code == "OFFICE_SUITE_REQUIRED"


def test_ole_with_docx_extension_uses_converter_instead_of_extension_guess(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    converted = tmp_path / "converted.docx"
    converted.write_bytes(_agenda_docx())

    def fake_prepare(source: Path, _workspace: Path) -> tuple[Path, bool]:
        assert source.suffix == ".doc"
        assert source.read_bytes().startswith(meeting_imports.OLE_HEADER)
        return converted, True

    monkeypatch.setattr(meeting_imports, "prepare_docx", fake_prepare)
    text, kind = meeting_imports.extract_meeting_text(meeting_imports.OLE_HEADER + b"\0" * 32)
    assert kind == "ole"
    assert "党委会议程" in text


def test_meeting_candidates_do_not_guess_ambiguous_facts() -> None:
    candidate = meeting_imports.propose_meeting(
        """关于事项的说明\n地点：第二会议室\n参会人：王一、赵二\n普通正文没有明确序号\n""",
        "普通文件.docx",
    )
    assert candidate["title"] == "普通文件"
    assert candidate["scheduled_at"] is None
    assert candidate["topics"] == []
    assert len(candidate["warnings"]) == 2
    assert candidate["venue"] == "第二会议室"
    assert candidate["attendee_text"] == "王一、赵二"


def test_meeting_candidate_date_time_and_title_scoring_edges() -> None:
    candidate = meeting_imports.propose_meeting(
        """中共测试委员会党委会议程\n会议时间：2026年2月30日上午25:61\n主持：李甲\n一、第一议题\n一、第一议题\n（2）第二议题\n""",
        "没有日期.docx",
    )
    assert candidate["title"] == "中共测试委员会党委会议程"
    assert candidate["date_candidate"] is None
    assert candidate["scheduled_at"] is None
    assert [item["title"] for item in candidate["topics"]] == ["第一议题", "第二议题"]
    assert candidate["host_name"] == "李甲"

    afternoon = meeting_imports.propose_meeting("会议时间：2026年8月1日下午6:05", "党委会记录.docx")
    assert afternoon["scheduled_at"] == "2026-08-01T18:05:00+08:00"

    no_time = meeting_imports.propose_meeting("会议日期：2026年8月2日", "党委会记录.docx")
    assert no_time["date_candidate"] == "2026-08-02" and no_time["scheduled_at"] is None
    assert any("未猜测会议时间" in item for item in no_time["warnings"])

    invalid_time = meeting_imports.propose_meeting("会议时间：2026年8月3日上午25:61", "党委会记录.docx")
    assert invalid_time["date_candidate"] == "2026-08-03" and invalid_time["scheduled_at"] is None

    penalized = meeting_imports.propose_meeting(
        "建议以上事项现提请党委会审议。\n其他普通正文",
        "",
    )
    assert penalized["title"] == "建议以上事项现提请党委会审议。"


def test_meeting_import_is_draft_first_and_requires_exact_confirmation(client, admin: dict) -> None:
    inspected = client.post(
        "/api/v1/business-meetings/imports/inspect",
        files={
            "file": (
                "20260828党委会议程.docx",
                _agenda_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert inspected.status_code == 201, inspected.text
    draft = inspected.json()
    assert draft["source_kind"] == "ooxml"
    assert draft["meeting"]["requires_confirmation"] is True
    assert draft["meeting"]["date_candidate"] == "2026-08-28"
    assert len(draft["topics"]) == 2
    assert all(item["confirmed"] is False for item in draft["topics"])
    assert "参会人员：张三、李四" not in inspected.text
    assert "20260828党委会议程.docx" not in inspected.text

    premature = client.post(
        f"/api/v1/business-meetings/imports/{draft['id']}/commit",
        headers={"If-Match": str(draft["version"])},
    )
    assert premature.status_code == 409
    assert premature.json()["code"] == "MEETING_IMPORT_CONFIRMATION_REQUIRED"

    confirmed = client.patch(
        f"/api/v1/business-meetings/imports/{draft['id']}",
        headers={"If-Match": str(draft["version"])},
        json={
            "meeting_type": "party_committee",
            "organization": "测试镇党委",
            "title": "2026年第8次党委会",
            "scheduled_at": "2026-08-28T18:00:00+08:00",
            "venue": "第一会议室",
            "host_name": "张三",
            "attendee_text": "张三、李四",
            "topics": ["研究基层党建重点任务", "审议专项资金使用方案"],
        },
    )
    assert confirmed.status_code == 200, confirmed.text
    updated = confirmed.json()
    assert updated["meeting"]["scheduled_at"] == "2026-08-28T10:00:00+00:00"
    assert all(item["source"] == "用户人工确认" for item in updated["topics"])

    committed = client.post(
        f"/api/v1/business-meetings/imports/{draft['id']}/commit",
        headers={"If-Match": str(updated["version"])},
    )
    assert committed.status_code == 201, committed.text
    meeting = committed.json()["meeting"]
    assert meeting["organization"] == "测试镇党委"
    assert [item["title"] for item in meeting["topics"]] == [
        "研究基层党建重点任务",
        "审议专项资金使用方案",
    ]
    assert meeting["progress"]["total"] == 6
    assert committed.json()["draft"]["status"] == "committed"

    repeated = client.post(
        f"/api/v1/business-meetings/imports/{draft['id']}/commit",
        headers={"If-Match": str(committed.json()["draft"]["version"])},
    )
    assert repeated.status_code == 201
    assert repeated.json()["meeting"]["id"] == meeting["id"]


def test_cancel_meeting_import_clears_candidates(client, admin: dict) -> None:
    inspected = client.post(
        "/api/v1/business-meetings/imports/inspect",
        files={"file": ("agenda.docx", _agenda_docx())},
    ).json()
    cancelled = client.delete(
        f"/api/v1/business-meetings/imports/{inspected['id']}",
        headers={"If-Match": str(inspected["version"])},
    )
    assert cancelled.status_code == 200, cancelled.text
    assert cancelled.json()["cancelled"] is True

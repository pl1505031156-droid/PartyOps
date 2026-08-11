"""1.4.3 党员发展确定性规则、权限、补充材料和 Word 导出测试。"""

from __future__ import annotations

from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from fastapi.testclient import TestClient
from alembic import command
from alembic.config import Config
from sqlalchemy import create_engine, inspect

from app.database import Base
from app.models import WorkCalendarEntry
from app.party_development import WorkdayCalendar, add_months, calculate_party_development
from app.party_development import safe_person_filename
from app.routers import party_development as party_router
from app.schemas import PartyDevelopmentCalculateRequest


def login(client: TestClient, username: str) -> None:
    response = client.post(
        "/api/v1/auth/login",
        json={"username": username, "password": "PartyOps@2026"},
    )
    assert response.status_code == 200, response.text


def test_natural_month_leap_year_and_workday_fallback() -> None:
    assert add_months(date(2026, 1, 31), 1) == date(2026, 2, 28)
    assert add_months(date(2024, 2, 29), 12) == date(2025, 2, 28)
    provisional = WorkdayCalendar([])
    end, is_provisional = provisional.add_inclusive(date(2026, 5, 1), 5)
    assert end == date(2026, 5, 7)
    assert is_provisional is True

    configured = WorkdayCalendar([
        WorkCalendarEntry(date_key="2026-05-04", title="节假日", kind="holiday", is_workday=False, owner_id="test"),
        WorkCalendarEntry(date_key="2026-05-09", title="调休", kind="adjusted_workday", is_workday=True, owner_id="test"),
    ])
    end, is_provisional = configured.add_inclusive(date(2026, 5, 1), 5)
    assert end == date(2026, 5, 8)
    assert is_provisional is False
    invalid_entry = WorkCalendarEntry(date_key="not-a-date", title="坏数据", kind="holiday", is_workday=False, owner_id="test")
    assert WorkdayCalendar([invalid_entry]).configured_years == set()
    try:
        provisional.add_inclusive(date(2026, 5, 1), 0)
    except ValueError as exc:
        assert "必须大于零" in str(exc)
    else:  # pragma: no cover
        raise AssertionError("零工作日必须被拒绝")
    assert safe_person_filename("../非法:姓名") == "非法_姓名"


def test_rules_cover_deadlines_manual_nodes_and_risk_warnings() -> None:
    payload = PartyDevelopmentCalculateRequest.model_validate({
        "name": "张三",
        "application_date": "2026-01-31",
        "actual_dates": {
            "conversation_date": "2026-03-05",
            "activist_date": "2026-03-01",
            "publicity_start_date": "2027-02-22",
            "development_object_date": "2027-02-20",
            "training_completed_date": "2027-03-01",
            "training_days": 2,
            "training_hours": 12,
            "pre_review_approved_date": "2027-03-31",
            "branch_acceptance_date": "2027-05-02",
            "committee_approval_date": "2027-11-03",
            "transition_application_date": "2028-04-01",
            "transition_branch_meeting_date": "2028-05-03",
            "transition_approval_date": "2028-08-04",
        },
    })
    result = calculate_party_development(payload, [], {}, today=date(2026, 2, 1))
    nodes = {node.key: node for node in result.nodes}
    assert nodes["conversation_deadline"].date == date(2026, 2, 28)
    assert nodes["development_object_earliest"].date == date(2027, 3, 1)
    assert nodes["development_object_publicity"].end_date == date(2027, 2, 26)
    assert nodes["branch_acceptance_deadline"].date == date(2027, 4, 30)
    assert nodes["committee_approval"].date == date(2027, 8, 2)
    assert nodes["committee_approval"].end_date == date(2027, 11, 2)
    assert nodes["probation_end"].date == date(2028, 5, 2)
    assert nodes["transition_approval_deadline"].date == date(2028, 8, 3)
    codes = {warning["code"] for warning in result.warnings}
    assert {
        "CONVERSATION_OVERDUE",
        "WORK_CALENDAR_INCOMPLETE",
        "DEVELOPMENT_OBJECT_TOO_EARLY",
        "PUBLICITY_NOT_COMPLETE",
        "TRAINING_INSUFFICIENT",
        "BRANCH_MEETING_OVERDUE",
        "COMMITTEE_APPROVAL_OVERDUE",
        "TRANSITION_APPLICATION_EARLY",
    }.issubset(codes)
    assert result.provisional is True
    assert "政治审查完成" in result.manual_confirmation_items


def test_public_calculation_and_word_export_do_not_require_admin(
    client: TestClient,
    admin: dict,
    staff: dict,
) -> None:
    login(client, "staff")
    current = client.get("/api/v1/party-development/rules/current")
    assert current.status_code == 200
    assert current.json()["version"] == "2026.05"

    payload = {
        "name": "李四/测试",
        "application_date": "2026-05-20",
        "actual_dates": {"activist_date": "2026-06-30"},
    }
    calculated = client.post("/api/v1/party-development/calculate", json=payload)
    assert calculated.status_code == 200, calculated.text
    assert calculated.json()["name"] == "李四/测试"
    assert len(calculated.json()["nodes"]) >= 15

    exported = client.post("/api/v1/party-development/export.docx", json=payload)
    assert exported.status_code == 200, exported.text
    assert exported.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    document = Document(BytesIO(exported.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "李四/测试" in text
    assert "规则版本为2026.05" in text
    assert "本材料仅供党务工作参考" in text
    assert "2026年5月20日" in text

    section = document.sections[0]
    assert round(section.page_width.mm) == 210
    assert round(section.page_height.mm) == 297
    assert round(section.top_margin.mm) == 37
    assert round(section.bottom_margin.mm) == 35
    assert round(section.left_margin.mm) == 28
    assert round(section.right_margin.mm) == 26
    assert document.paragraphs[0].text == "党员发展时间节点与材料提示"
    assert document.paragraphs[0].runs[0].font.size.pt == 22
    assert document.paragraphs[0].runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == "方正小标宋简体"

    assert len(document.tables) == 6
    assert "（一）申请入党" in text
    assert "（六）材料归档" in text
    for timeline in document.tables:
        headers = [cell.text for cell in timeline.rows[0].cells]
        assert headers == ["工作节点", "时间安排", "办理状态", "条款依据"]
    table_text = "\n".join(
        cell.text
        for timeline in document.tables
        for row in timeline.rows
        for cell in row.cells
    )
    assert "日期类型" not in table_text
    assert "阶段" not in {cell.text for table in document.tables for cell in table.rows[0].cells}
    assert not {"completed", "overdue", "waiting_manual", "planned", "actual", "deadline"}.intersection(
        table_text.split()
    )
    assert "已办理" in table_text
    assert "待组织确认" in table_text
    assert "2026年5月20日" in table_text
    expected_widths = [1984, 1701, 1134, 4025]
    for timeline in document.tables:
        table_width = timeline._tbl.tblPr.find(qn("w:tblW"))
        assert table_width is not None
        assert int(table_width.get(qn("w:w"))) == sum(expected_widths)
        assert [int(column.get(qn("w:w"))) for column in timeline._tbl.tblGrid] == expected_widths
        for row in timeline.rows:
            assert [int(cell._tc.tcPr.tcW.get(qn("w:w"))) for cell in row.cells] == expected_widths
    assert "制度来源：共产党员网《2026年新版细则全文》（点击查看）。" in text
    assert any(
        relationship.is_external
        and relationship.target_ref == "https://www.12371.cn/2026/05/18/ARTI1779102179030620.shtml"
        for relationship in document.part.rels.values()
    )
    assert " PAGE " in document.sections[0].footer._element.xml
    assert client.get("/api/v1/admin/party-development/profiles").status_code == 403
    login(client, "admin")


def test_admin_profiles_are_versioned_audited_and_only_add_materials(
    client: TestClient,
    admin: dict,
) -> None:
    login(client, "admin")
    seeded = client.get("/api/v1/admin/party-development/profiles")
    assert seeded.status_code == 200, seeded.text
    reference = next(item for item in seeded.json() if "待管理员确认" in item["name"])
    assert reference["active"] is False
    assert {item["name"] for item in reference["items"]} >= {"三考材料", "季度思想汇报", "个人自传"}
    assert client.get("/api/v1/admin/party-development/profiles").status_code == 200

    created = client.post("/api/v1/admin/party-development/profiles", json={
        "name": "测试单位材料模板",
        "description": "仅追加单位材料",
        "active": True,
        "items": [{
            "phase": "activist",
            "name": "单位培养联系人记录",
            "responsible_party": "培养联系人",
            "guidance": "单位补充，不改变一年培养期限",
            "required": True,
            "sort_order": 1,
        }],
    })
    assert created.status_code == 201, created.text
    profile = created.json()
    assert client.patch(
        f"/api/v1/admin/party-development/profiles/{profile['id']}",
        json={"description": "过期更新"},
        headers={"If-Match": "99"},
    ).status_code == 409

    replaced = client.put(
        f"/api/v1/admin/party-development/profiles/{profile['id']}/items",
        json=[{
            "phase": "development_object",
            "name": "本单位三考复核单",
            "responsible_party": "党支部",
            "guidance": "只作为材料提示",
            "required": False,
            "enabled": True,
            "sort_order": 5,
        }],
        headers={"If-Match": str(profile["version"])},
    )
    assert replaced.status_code == 200, replaced.text
    updated = replaced.json()
    assert updated["version"] == profile["version"] + 1

    calculation = client.post("/api/v1/party-development/calculate", json={
        "name": "王五",
        "application_date": "2026-06-01",
        "actual_dates": {},
        "profile_ids": [profile["id"]],
    })
    assert calculation.status_code == 200
    development_materials = [
        material
        for node in calculation.json()["nodes"] if node["phase"] == "development_object"
        for material in node["materials"]
    ]
    assert any(item["name"] == "本单位三考复核单" and item["national"] is False for item in development_materials)

    duplicate = client.put(
        f"/api/v1/admin/party-development/profiles/{profile['id']}/items",
        json=[
            {"phase": "activist", "name": "重复材料"},
            {"phase": "activist", "name": "重复材料"},
        ],
        headers={"If-Match": str(updated["version"])},
    )
    assert duplicate.status_code == 422

    deleted = client.delete(
        f"/api/v1/admin/party-development/profiles/{profile['id']}",
        headers={"If-Match": str(updated["version"])},
    )
    assert deleted.status_code == 204, deleted.text


def test_profile_error_paths_and_provisional_word(
    client: TestClient,
    admin: dict,
) -> None:
    login(client, "admin")
    first = client.post("/api/v1/admin/party-development/profiles", json={
        "name": "错误路径模板甲", "source_label": " 来源甲 ", "items": [],
    })
    assert first.status_code == 201, first.text
    duplicate = client.post("/api/v1/admin/party-development/profiles", json={
        "name": "错误路径模板甲", "items": [],
    })
    assert duplicate.status_code == 409
    second = client.post("/api/v1/admin/party-development/profiles", json={
        "name": "错误路径模板乙", "items": [],
    }).json()
    patched = client.patch(
        f"/api/v1/admin/party-development/profiles/{first.json()['id']}",
        json={"description": "  已去空格  ", "source_label": "  本地来源  "},
        headers={"If-Match": str(first.json()["version"])},
    )
    assert patched.status_code == 200
    assert patched.json()["description"] == "已去空格"
    rename_conflict = client.patch(
        f"/api/v1/admin/party-development/profiles/{first.json()['id']}",
        json={"name": second["name"]},
        headers={"If-Match": str(patched.json()["version"])},
    )
    assert rename_conflict.status_code == 409
    assert client.patch(
        "/api/v1/admin/party-development/profiles/not-found",
        json={"active": True}, headers={"If-Match": "1"},
    ).status_code == 404
    too_many = client.put(
        f"/api/v1/admin/party-development/profiles/{second['id']}/items",
        json=[{"phase": "activist", "name": f"材料{index}"} for index in range(201)],
        headers={"If-Match": str(second["version"])},
    )
    assert too_many.status_code == 422

    exported = client.post("/api/v1/party-development/export.docx", json={
        "name": "暂算导出",
        "application_date": "2026-01-01",
        "actual_dates": {
            "conversation_date": "2026-02-10",
            "publicity_start_date": "2027-02-22",
        },
    })
    assert exported.status_code == 200
    document = Document(BytesIO(exported.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "工作日结果按周一至周五暂算" in text
    assert "实际谈话日期超过" in text

    class BrokenPath:
        def unlink(self, *, missing_ok: bool) -> None:
            assert missing_ok is True
            raise OSError("locked")

    party_router._delete_temporary_export(BrokenPath())
    for item in (patched.json(), second):
        current = client.get("/api/v1/admin/party-development/profiles").json()
        target = next(profile for profile in current if profile["id"] == item["id"])
        response = client.delete(
            f"/api/v1/admin/party-development/profiles/{item['id']}",
            headers={"If-Match": str(target["version"])},
        )
        assert response.status_code == 204


def test_0018_migration_round_trip(tmp_path) -> None:
    database = tmp_path / "upgrade-0018.sqlite3"
    engine = create_engine(f"sqlite:///{database.as_posix()}")
    backend_root = Path(__file__).resolve().parents[1]
    config = Config(str(backend_root / "alembic.ini"))
    config.set_main_option("script_location", str(backend_root / "alembic"))
    with engine.begin() as connection:
        Base.metadata.create_all(connection)
        config.attributes["connection"] = connection
        command.stamp(config, "0018")
        command.downgrade(config, "0017")
        command.upgrade(config, "0018")
    inspector = inspect(engine)
    assert {"party_development_profiles", "party_development_materials"}.issubset(inspector.get_table_names())
    assert {"source_label", "active", "version", "created_by"}.issubset(
        {column["name"] for column in inspector.get_columns("party_development_profiles")}
    )

"""1.4.5-rc.3 内嵌公文排版票据、回环服务与隐私边界。"""

from __future__ import annotations

import http.client
import json
import re
from datetime import UTC, datetime, timedelta
from io import BytesIO
from pathlib import Path

import pytest
from app.official_format import OfficialFormatError
from app.official_format_service import (
    OfficialFormatLocalService,
    issue_local_format_ticket,
    verify_local_format_ticket,
)
from docx import Document
from fastapi.testclient import TestClient

ORIGIN = "https://partyops.local"
SECRET = "a" * 64


def _docx_bytes() -> bytes:
    stream = BytesIO()
    document = Document()
    document.add_heading("关于开展基层党建工作的通知", level=0)
    document.add_paragraph("请认真落实,按期完成。")
    document.save(stream)
    return stream.getvalue()


def _request(
    service: OfficialFormatLocalService,
    method: str,
    path: str,
    *,
    body: bytes | None = None,
    headers: dict[str, str] | None = None,
) -> tuple[int, dict[str, str], bytes]:
    connection = http.client.HTTPConnection("127.0.0.1", service.port, timeout=10)
    request_headers = {"Host": f"127.0.0.1:{service.port}", "Origin": ORIGIN}
    request_headers.update(headers or {})
    connection.request(method, path, body=body, headers=request_headers)
    response = connection.getresponse()
    payload = response.read()
    result_headers = {key.lower(): value for key, value in response.getheaders()}
    status = response.status
    connection.close()
    return status, result_headers, payload


def _json(payload: bytes) -> dict:
    return json.loads(payload.decode("utf-8"))


def test_local_format_ticket_is_origin_bound_one_purpose_and_expires() -> None:
    now = datetime(2026, 8, 25, 8, 0, tzinfo=UTC)
    ticket, expires_at = issue_local_format_ticket(
        SECRET,
        origin=ORIGIN,
        user_id="user-1",
        device_id="device-1",
        now=now,
    )
    assert expires_at == now + timedelta(seconds=120)
    claims = verify_local_format_ticket(SECRET, ticket, origin=ORIGIN, now=now)
    assert claims["user_id"] == "user-1"
    assert claims["device_id"] == "device-1"
    with pytest.raises(OfficialFormatError, match="页面会话可能已过期"):
        verify_local_format_ticket(SECRET, ticket, origin="https://evil.invalid", now=now)
    with pytest.raises(OfficialFormatError):
        verify_local_format_ticket(SECRET, ticket + "tampered", origin=ORIGIN, now=now)
    with pytest.raises(OfficialFormatError):
        verify_local_format_ticket(
            SECRET,
            ticket,
            origin=ORIGIN,
            now=now + timedelta(minutes=3),
        )


def test_embedded_local_service_diagnose_format_download_and_cleanup(tmp_path: Path) -> None:
    service = OfficialFormatLocalService(
        secret=SECRET,
        config_dir=tmp_path / "logs",
        port=0,
        idle_timeout=60,
    ).start()
    try:
        status, headers, payload = _request(service, "GET", "/health")
        assert status == 200
        assert headers["access-control-allow-origin"] == ORIGIN
        assert _json(payload)["service"] == "official-format"

        status, headers, _ = _request(
            service,
            "OPTIONS",
            "/v1/sessions",
            headers={
                "Access-Control-Request-Method": "POST",
                "Access-Control-Request-Private-Network": "true",
            },
        )
        assert status == 204
        assert headers["access-control-allow-private-network"] == "true"

        ticket, _ = issue_local_format_ticket(
            SECRET,
            origin=ORIGIN,
            user_id="user-1",
            device_id="device-1",
        )
        status, _, payload = _request(
            service,
            "POST",
            "/v1/sessions",
            headers={"Authorization": f"Bearer {ticket}"},
        )
        assert status == 201
        session = _json(payload)
        session_id = session["session_id"]
        local_token = session["session_token"]

        # 启动票据只能兑换一次。
        status, _, payload = _request(
            service,
            "POST",
            "/v1/sessions",
            headers={"Authorization": f"Bearer {ticket}"},
        )
        assert status == 422 and _json(payload)["code"] == "LOCAL_TICKET_USED"

        boundary = "PartyOpsEmbedded"
        document = _docx_bytes()
        body = (
            f"--{boundary}\r\nContent-Disposition: form-data; name=\"document\"; filename=\"sample.docx\"\r\n"
            "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
        ).encode() + document + f"\r\n--{boundary}--\r\n".encode()
        status, _, payload = _request(
            service,
            "POST",
            f"/v1/sessions/{session_id}/diagnose",
            body=body,
            headers={
                "Content-Type": f"multipart/form-data; boundary={boundary}",
                "Content-Length": str(len(body)),
                "X-PartyOps-Local-Token": local_token,
            },
        )
        assert status == 200
        diagnosis = _json(payload)
        assert diagnosis["report"]["paragraph_count"] >= 1
        document_id = diagnosis["document_id"]

        status, _, payload = _request(
            service,
            "POST",
            f"/v1/sessions/{session_id}/documents/{document_id}/format",
            headers={"X-PartyOps-Local-Token": local_token},
        )
        assert status == 200
        assert "report" in _json(payload)

        status, headers, payload = _request(
            service,
            "GET",
            f"/v1/sessions/{session_id}/documents/{document_id}/download",
            headers={"X-PartyOps-Local-Token": local_token},
        )
        assert status == 200 and payload.startswith(b"PK")
        assert "%E5%85%AC%E6%96%87%E8%A7%84%E8%8C%83%E7%89%88" in headers["content-disposition"]
        assert session_id not in service.sessions

        log = (tmp_path / "logs" / "official-format.log").read_text(encoding="utf-8")
        assert "sample.docx" not in log
        assert "关于开展基层党建工作的通知" not in log
        assert re.search(r'"stage":"diagnose"', log)
    finally:
        service.close()


def test_local_service_rejects_wrong_origin_and_session_token(tmp_path: Path) -> None:
    service = OfficialFormatLocalService(
        secret=SECRET,
        config_dir=tmp_path,
        port=0,
    ).start()
    try:
        ticket, _ = issue_local_format_ticket(
            SECRET,
            origin=ORIGIN,
            user_id="user-1",
            device_id="device-1",
        )
        status, _, payload = _request(
            service,
            "POST",
            "/v1/sessions",
            headers={"Authorization": f"Bearer {ticket}"},
        )
        session = _json(payload)
        assert status == 201
        status, _, payload = _request(
            service,
            "DELETE",
            f"/v1/sessions/{session['session_id']}",
            headers={"X-PartyOps-Local-Token": "wrong"},
        )
        assert status == 422 and _json(payload)["code"] == "LOCAL_SESSION_INVALID"
    finally:
        service.close()


def test_authenticated_api_only_returns_ticket_metadata(
    client: TestClient,
    admin: dict,
) -> None:
    response = client.post(
        "/api/v1/official-format/local-ticket",
        json={"origin": "http://testserver"},
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert set(payload) == {"expires_at", "local_base_url", "ticket"}
    assert payload["local_base_url"].startswith("http://127.0.0.1:")
    assert "filename" not in response.text.lower()
    assert "path" not in response.text.lower()

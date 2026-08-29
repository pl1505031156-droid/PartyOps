"""新公文工具六类功能与 25 项能力契约的跨平台回归。"""

from __future__ import annotations

import hashlib
import subprocess
import zipfile
from pathlib import Path
from types import SimpleNamespace

import fitz
import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app import official_format_features as features
from app.official_format import OfficialFormatError
from app.official_format_features import capabilities_payload, execute_feature


def _source_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("关于开展年度党建工作的通知")
    document.add_paragraph("中共测试委员会〔2026〕8号")
    document.add_paragraph("基层支部工作安排")
    document.add_paragraph("旧称于2026年执行。第十二条需要格式调整。")
    document.add_paragraph("中共测试委员会")
    document.add_paragraph("2026年8月29日")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = "党建"
    table.cell(1, 1).text = "完成"
    document.save(path)


def _source_pdf(path: Path, pages: int = 2) -> None:
    document = fitz.open()
    for index in range(pages):
        page = document.new_page(width=595, height=842)
        page.insert_text((72, 96), f"PartyOps PDF page {index + 1} local conversion text")
    document.save(path)
    document.close()


def _hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def test_capability_catalog_is_exact_and_office_independent() -> None:
    payload = capabilities_payload()
    assert payload["engine"] == "partyops-bundled"
    assert payload["external_office_required"] is False
    assert payload["capability_count"] == 25
    assert [item["id"] for item in payload["features"]] == [
        "format",
        "replace",
        "redheader",
        "rename",
        "convert",
        "pdf-to-word",
    ]
    assert sum(len(item["capabilities"]) for item in payload["features"]) == 25
    convert_dpi = next(
        option
        for feature in payload["features"] if feature["id"] == "convert"
        for option in feature["option_schema"] if option["id"] == "dpi"
    )
    assert convert_dpi == {
        "id": "dpi",
        "label": "图片清晰度",
        "type": "number",
        "default": 200,
        "min": 72,
        "max": 600,
    }


def test_format_full_and_selection_scope_preserve_source(tmp_path: Path, monkeypatch) -> None:
    source = tmp_path / "中文 空格.docx"
    _source_docx(source)
    original_hash = _hash(source)
    monkeypatch.setattr("app.official_format._font_inventory", lambda: "方正小标宋 仿宋 楷体 黑体")

    full = execute_feature("format", source, tmp_path / "full", {"scope": "full"})
    assert full.report and full.report.changed_count > 0
    full_document = Document(full.outputs[0].path)
    assert full_document.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert full_document.tables[0].cell(1, 1).text == "完成"

    selection = execute_feature(
        "format",
        source,
        tmp_path / "selection",
        {"scope": "selection", "start_paragraph": 4, "end_paragraph": 4},
    )
    selected_document = Document(selection.outputs[0].path)
    assert selected_document.paragraphs[0].alignment is None
    assert selected_document.paragraphs[3].runs[0]._element.rPr is not None
    assert _hash(source) == original_hash

    with pytest.raises(OfficialFormatError, match="可排版段落"):
        execute_feature(
            "format",
            source,
            tmp_path / "invalid-selection",
            {"scope": "selection", "start_paragraph": 999, "end_paragraph": 1000},
        )


def test_replace_text_regex_wildcard_format_and_batch_rules(tmp_path: Path) -> None:
    source = tmp_path / "replace.docx"
    _source_docx(source)
    result = execute_feature(
        "replace",
        source,
        tmp_path / "replace-output",
        {
            "plan_name": "完整替换方案",
            "rules": [
                {"mode": "text", "find": "旧称", "replace": "新称"},
                {"mode": "regex", "find": r"(2026)年", "replace": "$1年度"},
                {"mode": "wildcard", "find": "第*条", "replace": "条款"},
                {"mode": "format", "find": "格式", "font_name": "黑体", "font_size": 16, "alignment": "left"},
            ],
        },
    )
    output = Document(result.outputs[0].path)
    body = output.paragraphs[3]
    assert "新称于2026年度执行。条款需要格式调整。" == body.text
    assert body.runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"
    assert "共应用" in result.message

    with pytest.raises(OfficialFormatError, match="unterminated"):
        execute_feature(
            "replace",
            source,
            tmp_path / "replace-invalid",
            {"rules": [{"mode": "regex", "find": "(", "replace": ""}]},
        )


def test_redheader_three_document_types_and_top_marks(tmp_path: Path) -> None:
    source = tmp_path / "redheader.docx"
    _source_docx(source)
    common = {
        "copy_number": "8",
        "security": "机密 10年",
        "urgency": "特急",
        "agency": "中共测试委员会",
        "document_number": "测试党发〔2026〕8号",
        "signatory": "张三",
        "imprint": "中共测试委员会办公室",
    }
    for document_type in ("down", "up", "letter"):
        result = execute_feature(
            "redheader",
            source,
            tmp_path / document_type,
            {**common, "document_type": document_type},
        )
        texts = [paragraph.text for paragraph in Document(result.outputs[0].path).paragraphs]
        assert "000008" in texts
        assert "中共测试委员会" in texts
        assert any("测试党发〔2026〕8号" in text for text in texts)
        if document_type == "up":
            assert any("签发人：张三" in text for text in texts)
        if document_type == "letter":
            assert texts[-1] != "中共测试委员会办公室"
        else:
            assert texts[-1] == "中共测试委员会办公室"


def test_rename_content_parts_rotation_and_safe_copy(tmp_path: Path) -> None:
    source = tmp_path / "待命名 文档.docx"
    _source_docx(source)
    original_hash = _hash(source)
    result = execute_feature(
        "rename",
        source,
        tmp_path / "rename",
        {
            "parts": ["mainTitle", "docNumber", "custom", "rotation"],
            "custom_text": "归档/件",
            "rotation_words": "甲、乙、丙",
            "separator": "-",
        },
    )
    assert result.outputs[0].filename.endswith(".docx")
    assert "关于开展年度党建工作的通知" in result.outputs[0].filename
    assert "中共测试委员会〔2026〕8号" in result.outputs[0].filename
    assert "归档件" in result.outputs[0].filename
    assert "甲" in result.outputs[0].filename
    assert _hash(source) == original_hash


def test_convert_document_text_page_images_long_image_and_page_validation(tmp_path: Path) -> None:
    source = tmp_path / "convert.docx"
    pdf = tmp_path / "convert.pdf"
    _source_docx(source)
    _source_pdf(pdf)

    docx_result = execute_feature("convert", source, tmp_path / "docx", {"target_format": "docx"})
    assert docx_result.outputs[0].path.read_bytes().startswith(b"PK")
    text_result = execute_feature("convert", source, tmp_path / "txt", {"target_format": "txt"})
    assert "关于开展年度党建工作的通知" in text_result.outputs[0].path.read_text(encoding="utf-8")
    pdf_result = execute_feature("convert", pdf, tmp_path / "pdf", {"target_format": "pdf"})
    assert pdf_result.outputs[0].path.read_bytes().startswith(b"%PDF")

    page_images = execute_feature(
        "convert",
        pdf,
        tmp_path / "pages",
        {"target_format": "png", "image_mode": "pages", "page_selection": "1，2", "dpi": 72},
    )
    assert len(page_images.outputs) == 2
    assert all(output.path.read_bytes().startswith(b"\x89PNG") for output in page_images.outputs)
    long_image = execute_feature(
        "convert",
        pdf,
        tmp_path / "long",
        {"target_format": "jpg", "image_mode": "long", "page_selection": "1至2", "dpi": 72},
    )
    assert len(long_image.outputs) == 1
    assert long_image.outputs[0].path.read_bytes().startswith(b"\xff\xd8")

    with pytest.raises(OfficialFormatError, match="起始页"):
        execute_feature(
            "convert",
            pdf,
            tmp_path / "invalid-pages",
            {"target_format": "png", "page_selection": "2-1"},
        )


def test_pdf_to_word_local_reconstruction_and_verification(tmp_path: Path) -> None:
    source = tmp_path / "local.pdf"
    _source_pdf(source, pages=1)
    result = execute_feature(
        "pdf-to-word",
        source,
        tmp_path / "word",
        {"normalize_punctuation": True, "reconstruct_tables": True},
    )
    output = Document(result.outputs[0].path)
    assert "PartyOps PDF page 1 local conversion text" in "\n".join(
        paragraph.text for paragraph in output.paragraphs
    )
    assert result.outputs[0].path.read_bytes().startswith(b"PK")


def test_cancel_and_unknown_feature_are_deterministic(tmp_path: Path) -> None:
    source = tmp_path / "cancel.docx"
    _source_docx(source)
    with pytest.raises(OfficialFormatError, match="源文件未改变"):
        execute_feature("format", source, tmp_path / "cancelled", cancelled=lambda: True)
    with pytest.raises(OfficialFormatError, match="受支持的排版功能"):
        execute_feature("missing", source, tmp_path / "missing")


def test_output_policy_and_invalid_word_packages_are_rejected(tmp_path: Path) -> None:
    """同名策略和损坏输入必须在写入结果前给出确定错误。"""

    workspace = tmp_path / "outputs"
    workspace.mkdir()
    base = workspace / "公文-结果.docx"
    assert features._unique_output(workspace, "公文", "-结果.docx", {}) == base
    base.write_bytes(b"occupied")
    assert features._unique_output(
        workspace,
        "公文",
        "-结果.docx",
        {"same_name_policy": "overwrite"},
    ) == base
    with pytest.raises(OfficialFormatError, match="同名策略为跳过"):
        features._unique_output(workspace, "公文", "-结果.docx", {"same_name_policy": "skip"})

    second = workspace / "公文-结果 (2).docx"
    second.write_bytes(b"occupied")
    assert features._unique_output(workspace, "公文", "-结果.docx", {}).name == "公文-结果 (3).docx"

    invalid_zip = tmp_path / "invalid.docx"
    invalid_zip.write_bytes(b"not-a-zip")
    with pytest.raises(OfficialFormatError, match="结构不完整"):
        features._validate_docx(invalid_zip)

    missing_document = tmp_path / "missing-document.docx"
    with zipfile.ZipFile(missing_document, "w") as package:
        package.writestr("[Content_Types].xml", "<Types/>")
    with pytest.raises(OfficialFormatError, match="结构不完整"):
        features._validate_docx(missing_document)

    unsupported = tmp_path / "unsupported.txt"
    unsupported.write_text("文本", encoding="utf-8")
    with pytest.raises(OfficialFormatError, match="只接受 DOCX"):
        features._prepare_word_source(unsupported, workspace)


def test_replace_rule_edge_cases_and_noop_formatting(tmp_path: Path) -> None:
    """替换规则覆盖空值、上限、非法结构、通配符和无操作格式分支。"""

    document = Document()
    blank = document.add_paragraph()
    body = document.add_paragraph("保持原文")
    assert features._replace_paragraph_text(blank, features.re.compile("原文"), "新文") == 0
    assert features._replace_paragraph_text(body, features.re.compile("不存在"), "新文") == 0
    assert features._compile_replace_pattern({"mode": "wildcard", "find": "保*文"}).search(body.text)
    with pytest.raises(OfficialFormatError, match="查找内容不能为空"):
        features._compile_replace_pattern({"find": ""})
    with pytest.raises(OfficialFormatError, match="不得超过 1000"):
        features._compile_replace_pattern({"mode": "regex", "find": "x" * 1001})

    assert features._apply_format_rule(document, {}) == 2
    assert features._apply_format_rule(document, {"find": "不匹配"}) == 0

    source = tmp_path / "replace-validation.docx"
    _source_docx(source)
    with pytest.raises(OfficialFormatError, match="至少添加一条"):
        execute_feature("replace", source, tmp_path / "no-rules", {"rules": []})
    with pytest.raises(OfficialFormatError, match="最多执行 100"):
        execute_feature(
            "replace",
            source,
            tmp_path / "too-many-rules",
            {"rules": [{"find": "x"}] * 101},
        )
    with pytest.raises(OfficialFormatError, match="结构化对象"):
        execute_feature("replace", source, tmp_path / "bad-rule", {"rules": ["invalid"]})


def test_page_selection_validation_covers_all_supported_notations() -> None:
    assert features._parse_page_selection("全部", 3) == [0, 1, 2]
    assert features._parse_page_selection("1,,3", 3) == [0, 2]
    assert features._parse_page_selection("2", 3) == [1]
    with pytest.raises(OfficialFormatError, match="1-3,5"):
        features._parse_page_selection("一-2", 3)
    with pytest.raises(OfficialFormatError, match="1-3,5"):
        features._parse_page_selection("首页", 3)
    with pytest.raises(OfficialFormatError, match="超出文档范围"):
        features._parse_page_selection("99", 3)


def test_feature_validation_and_pdf_conversion_alternatives(tmp_path: Path) -> None:
    """覆盖非法能力参数，以及 PDF 到 DOCX/TXT 的本地替代路径。"""

    source = tmp_path / "format-validation.docx"
    pdf = tmp_path / "two-pages.pdf"
    _source_docx(source)
    _source_pdf(pdf, pages=2)

    with pytest.raises(OfficialFormatError, match="一键排版 支持"):
        execute_feature("format", pdf, tmp_path / "wrong-input")
    with pytest.raises(OfficialFormatError, match="请选择全文"):
        execute_feature("format", source, tmp_path / "wrong-scope", {"scope": "page"})
    with pytest.raises(OfficialFormatError, match="必须为整数"):
        execute_feature(
            "format",
            source,
            tmp_path / "wrong-range",
            {"scope": "selection", "start_paragraph": "第一段"},
        )
    with pytest.raises(OfficialFormatError, match="请选择 DOCX"):
        execute_feature("convert", source, tmp_path / "wrong-format", {"target_format": "html"})

    text_result = execute_feature("convert", pdf, tmp_path / "pdf-text", {"target_format": "txt"})
    assert "PartyOps PDF page 2" in text_result.outputs[0].path.read_text(encoding="utf-8")
    docx_result = execute_feature(
        "convert",
        pdf,
        tmp_path / "pdf-docx",
        {"target_format": "docx", "normalize_punctuation": False, "reconstruct_tables": False},
    )
    assert len(Document(docx_result.outputs[0].path).sections) == 2


def test_redheader_empty_document_and_existing_border(tmp_path: Path) -> None:
    """空文档仍能套红，重复写红线时复用已有边框容器。"""

    source = tmp_path / "empty.docx"
    Document().save(source)
    result = execute_feature(
        "redheader",
        source,
        tmp_path / "redheader-empty",
        {"agency": "", "document_number": "", "imprint": ""},
    )
    output = Document(result.outputs[0].path)
    assert any(paragraph.text == "中共××委员会" for paragraph in output.paragraphs)

    paragraph = output.add_paragraph()
    features._red_line(paragraph)
    features._red_line(paragraph)
    borders = paragraph._p.get_or_add_pPr().find(qn("w:pBdr"))
    assert borders is not None and len(borders) == 2


def test_bundled_conversion_runtime_failures_and_success(tmp_path: Path, monkeypatch) -> None:
    """固定运行时缺失、启动失败、超时和成功结果都具有稳定行为。"""

    source = tmp_path / "runtime.docx"
    _source_docx(source)
    monkeypatch.setattr(features, "_office_candidates", lambda: [])
    with pytest.raises(OfficialFormatError, match="运行时缺失"):
        features._run_libreoffice_conversion(source, tmp_path, "pdf")

    runtime = tmp_path / "soffice"
    runtime.write_text("bundled", encoding="utf-8")
    monkeypatch.setattr(features, "_office_candidates", lambda: [runtime])
    monkeypatch.setattr(
        features.subprocess,
        "run",
        lambda *_args, **_kwargs: SimpleNamespace(returncode=1),
    )
    with pytest.raises(OfficialFormatError, match="没有生成"):
        features._run_libreoffice_conversion(source, tmp_path, "pdf")

    def raise_timeout(*_args, **_kwargs):
        raise subprocess.TimeoutExpired("soffice", 120)

    monkeypatch.setattr(features.subprocess, "run", raise_timeout)
    with pytest.raises(OfficialFormatError, match="120 秒"):
        features._run_libreoffice_conversion(source, tmp_path, "pdf")

    def create_result(command, **_kwargs):
        output_dir = Path(command[command.index("--outdir") + 1])
        (output_dir / f"{source.stem}.pdf").write_bytes(b"%PDF-local")
        return SimpleNamespace(returncode=0)

    monkeypatch.setattr(features.subprocess, "run", create_result)
    converted = features._run_libreoffice_conversion(source, tmp_path, "pdf")
    assert converted.read_bytes() == b"%PDF-local"


def test_image_memory_budgets_are_enforced_before_allocation(tmp_path: Path, monkeypatch) -> None:
    pdf = tmp_path / "memory.pdf"
    _source_pdf(pdf, pages=2)

    class FakeImage:
        def __init__(self, width: int, height: int):
            self.width = width
            self.height = height

        def convert(self, _mode: str):
            return self

        def copy(self):
            return self

    monkeypatch.setattr(features.Image, "open", lambda _stream: FakeImage(10_000, 9_000))
    with pytest.raises(OfficialFormatError, match="降低 DPI"):
        features._render_pdf_images(pdf, tmp_path / "page-budget", "memory", {"dpi": 72}, lambda *_: None, lambda: False)

    monkeypatch.setattr(features.Image, "open", lambda _stream: FakeImage(10_000, 7_000))
    with pytest.raises(OfficialFormatError, match="减少页面"):
        features._render_pdf_images(
            pdf,
            tmp_path / "long-budget",
            "memory",
            {"dpi": 72, "image_mode": "long"},
            lambda *_: None,
            lambda: False,
        )


def test_defensive_replace_and_image_only_pdf_fallbacks(tmp_path: Path, monkeypatch) -> None:
    """覆盖第三方对象异常形态和无文字 PDF 的图像保真降级。"""

    class ChangingRunsParagraph:
        def __init__(self):
            self.access_count = 0
            self.added_text = ""

        @property
        def runs(self):
            self.access_count += 1
            if self.access_count == 1:
                return [SimpleNamespace(text="旧值")]
            return []

        def add_run(self, text: str):
            self.added_text = text

    paragraph = ChangingRunsParagraph()
    assert features._replace_paragraph_text(paragraph, features.re.compile("旧值"), "新值") == 1
    assert paragraph.added_text == "新值"

    monkeypatch.setattr(features.fnmatch, "translate", lambda _pattern: "literal")
    assert features._compile_replace_pattern({"mode": "wildcard", "find": "literal"}).fullmatch("literal")

    image_only_pdf = tmp_path / "image-only.pdf"
    pdf = fitz.open()
    pdf.new_page(width=200, height=200)
    pdf.save(image_only_pdf)
    pdf.close()
    result = execute_feature(
        "pdf-to-word",
        image_only_pdf,
        tmp_path / "image-only-word",
        {"reconstruct_tables": False},
    )
    assert len(Document(result.outputs[0].path).inline_shapes) == 1

    wrong_input = tmp_path / "wrong.docx"
    _source_docx(wrong_input)
    with pytest.raises(OfficialFormatError, match="PDF 转 Word 只接受 PDF"):
        features._execute_pdf_to_word(wrong_input, tmp_path, {}, lambda *_: None, lambda: False)

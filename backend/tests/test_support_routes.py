"""模板、周期、知识库、联系人和辅助接口测试。"""

from __future__ import annotations

import io
from datetime import datetime, timedelta, timezone

import fitz
from docx import Document
from fastapi.testclient import TestClient

from .conftest import create_task


def test_template_create_duplicate_and_instantiate(
    client: TestClient, admin: dict
) -> None:
    payload = {
        "name": "测试月度台账模板",
        "category": "报送",
        "task_type": "standard",
        "description": "复用上期口径。",
        "steps": ["汇总数据", "审核报送"],
        "materials": [
            {"category": "final", "name": "实际报送稿", "required": True}
        ],
    }
    created = client.post("/api/v1/templates", json=payload)
    assert created.status_code == 201, created.text
    template = created.json()
    assert template["steps"] == payload["steps"]
    assert template["materials"][0]["required"] is True

    duplicate = client.post("/api/v1/templates", json=payload)
    assert duplicate.status_code == 409
    assert duplicate.json()["code"] == "TEMPLATE_EXISTS"

    instantiated = client.post(
        f"/api/v1/templates/{template['id']}/instantiate",
        json={
            "owner_id": admin["id"],
            "title": "八月月度台账",
            "formal_due_at": "2026-08-28T18:00:00+08:00",
        },
    )
    assert instantiated.status_code == 201
    body = instantiated.json()
    assert len(body["steps"]) == 2
    assert len(body["materials"]) == 1
    assert client.post(
        "/api/v1/templates/not-found/instantiate",
        json={"owner_id": admin["id"]},
    ).status_code == 404


def test_recurrence_create_and_run_due(client: TestClient, admin: dict) -> None:
    template = next(
        (
            item
            for item in client.get("/api/v1/templates").json()
            if item["name"] == "测试月度台账模板"
        ),
        None,
    )
    if template is None:
        response = client.post(
            "/api/v1/templates",
            json={
                "name": "测试月度台账模板",
                "category": "报送",
                "task_type": "standard",
                "description": "复用上期口径。",
                "steps": ["汇总数据", "审核报送"],
                "materials": [
                    {"category": "final", "name": "实际报送稿", "required": True}
                ],
            },
        )
        assert response.status_code == 201, response.text
        template = response.json()
    past = datetime.now(timezone.utc) - timedelta(minutes=5)
    response = client.post(
        "/api/v1/recurrences",
        json={
            "name": "月度台账自动生成",
            "template_id": template["id"],
            "owner_id": admin["id"],
            "kind": "monthly",
            "custom_days": None,
            "next_run_at": past.isoformat(),
        },
    )
    assert response.status_code == 201, response.text
    rule = response.json()
    generated = client.post("/api/v1/recurrences/run-due")
    assert generated.status_code == 200
    assert generated.json()
    updated = next(
        item
        for item in client.get("/api/v1/recurrences").json()
        if item["id"] == rule["id"]
    )
    assert updated["last_run_at"] is not None
    next_run = datetime.fromisoformat(updated["next_run_at"])
    if next_run.tzinfo is None:
        next_run = next_run.replace(tzinfo=timezone.utc)
    assert next_run > past
    missing = client.post(
        "/api/v1/recurrences",
        json={
            "name": "无效周期",
            "template_id": "missing",
            "owner_id": admin["id"],
            "kind": "monthly",
            "next_run_at": past.isoformat(),
        },
    )
    assert missing.status_code == 422


def test_knowledge_and_contacts(client: TestClient, admin: dict) -> None:
    entry = client.post(
        "/api/v1/knowledge",
        json={
            "title": "主题党日材料口径",
            "category": "经验",
            "body": "照片应包含活动主题和日期。",
        },
    )
    assert entry.status_code == 201
    results = client.get("/api/v1/knowledge", params={"keyword": "照片"}).json()
    assert any(item["id"] == entry.json()["id"] for item in results)

    contact = client.post(
        "/api/v1/contacts",
        json={
            "name": "李老师",
            "organization": "上级党建办",
            "phone": "内部短号",
            "note": "负责台账口径",
        },
    )
    assert contact.status_code == 201
    assert any(
        item["id"] == contact.json()["id"]
        for item in client.get("/api/v1/contacts").json()
    )


def test_docx_pdf_and_unknown_intake(client: TestClient, admin: dict) -> None:
    document = Document()
    document.add_paragraph("关于报送组织生活材料的通知")
    document.add_paragraph("请于2026年9月2日前提交材料。")
    buffer = io.BytesIO()
    document.save(buffer)
    parsed_docx = client.post(
        "/api/v1/intake/parse",
        data={"pasted_text": ""},
        files={
            "file": (
                "通知.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert parsed_docx.status_code == 200
    assert parsed_docx.json()["source_kind"] == "word"
    assert parsed_docx.json()["source_filename"] == "通知.docx"
    assert parsed_docx.json()["parser_label"] == "Word 文档本地提取"

    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "PartyOps PDF notice 2026-09-03")
    pdf_bytes = pdf.tobytes()
    pdf.close()
    parsed_pdf = client.post(
        "/api/v1/intake/parse",
        data={"pasted_text": "PDF 补充说明"},
        files={"file": ("notice.pdf", pdf_bytes, "application/pdf")},
    )
    assert parsed_pdf.status_code == 200
    assert parsed_pdf.json()["source_kind"] == "pdf"

    unsupported = client.post(
        "/api/v1/intake/parse",
        data={"pasted_text": ""},
        files={"file": ("binary.bin", b"\xff\xfe\x00", "application/octet-stream")},
    )
    assert unsupported.status_code == 200
    assert unsupported.json()["warnings"]
    assert unsupported.json()["source_filename"] == "binary.bin"
    assert unsupported.json()["parser_label"] == "仅保存原始附件"


def test_task_subresources_download_list_and_delete(
    client: TestClient, admin: dict, staff: dict
) -> None:
    task = create_task(client, admin["id"], steps=[], materials=[])
    task_id = task["id"]
    participant = client.post(
        f"/api/v1/tasks/{task_id}/participants",
        json={"user_id": staff["id"], "role": "collaborator"},
    )
    assert participant.status_code == 200
    assert client.post(
        f"/api/v1/tasks/{task_id}/participants",
        json={"user_id": "missing", "role": "collaborator"},
    ).status_code == 422

    step = client.post(
        f"/api/v1/tasks/{task_id}/steps",
        json={"title": "形成报送清单", "assignee_id": staff["id"]},
    )
    assert step.status_code == 201
    step_body = step.json()
    completed = client.patch(
        f"/api/v1/tasks/{task_id}/steps/{step_body['id']}",
        json={"done": True, "version": step_body["version"]},
    )
    assert completed.status_code == 200
    assert completed.json()["done"] is True
    assert client.patch(
        f"/api/v1/tasks/{task_id}/steps/{step_body['id']}",
        json={"done": False, "version": step_body["version"]},
    ).status_code == 409

    comment = client.post(
        f"/api/v1/tasks/{task_id}/comments",
        json={"body": "已与上级确认口径。"},
    )
    assert comment.status_code == 201
    material = client.post(
        f"/api/v1/tasks/{task_id}/materials",
        json={"category": "receipt", "name": "报送回执", "required": False},
    )
    assert material.status_code == 201
    uploaded = client.post(
        f"/api/v1/tasks/{task_id}/materials/{material.json()['id']}/versions",
        data={"stage": "submitted", "is_final": "true", "note": "回执"},
        files={"file": ("回执.txt", b"received", "text/plain")},
    )
    assert uploaded.status_code == 201
    version_id = uploaded.json()["materials"][0]["versions"][0]["id"]
    downloaded = client.get(f"/api/v1/attachments/{version_id}/download")
    assert downloaded.status_code == 200
    assert downloaded.content == b"received"

    listing = client.get(
        "/api/v1/tasks",
        params={"owner_id": admin["id"], "keyword": "基层党建", "page_size": 5},
    )
    assert listing.status_code == 200
    assert any(item["id"] == task_id for item in listing.json()["items"])
    assert client.get(f"/api/v1/tasks/{task_id}").status_code == 200
    assert client.get("/api/v1/tasks/missing").status_code == 404

    missing_header = client.delete(f"/api/v1/tasks/{task_id}")
    assert missing_header.status_code == 428
    current = client.get(f"/api/v1/tasks/{task_id}").json()
    deleted = client.delete(
        f"/api/v1/tasks/{task_id}",
        headers={"If-Match": str(current["version"])},
    )
    assert deleted.status_code == 200
    assert client.get(f"/api/v1/tasks/{task_id}").status_code == 404


def test_if_match_validation_and_invalid_action(client: TestClient, admin: dict) -> None:
    task = create_task(client, admin["id"], materials=[], steps=[])
    invalid_header = client.patch(
        f"/api/v1/tasks/{task['id']}",
        headers={"If-Match": "not-an-integer"},
        json={"title": "不会保存"},
    )
    assert invalid_header.status_code == 400
    invalid_action = client.post(
        f"/api/v1/tasks/{task['id']}/actions",
        json={"action": "archive", "note": ""},
    )
    assert invalid_action.status_code == 409

"""1.4.5-rc.2 公文本机排版的标准、保真与安全边界。"""

from __future__ import annotations

import base64
import zipfile
from io import BytesIO

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.official_format import (
    GRID_CHARACTER_SPACE,
    PAGE_FOOTER_DISTANCE,
    OfficialFormatError,
    diagnose_docx,
    format_docx,
    normalize_chinese_punctuation,
)

ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def test_punctuation_normalization_is_conservative() -> None:
    source = '请核对,然后访问https://example.test/a,b，邮箱ops@example.test，金额3.14万元，参见1.2.3。'
    normalized, changes = normalize_chinese_punctuation(source)
    assert "请核对，然后" in normalized
    assert "https://example.test/a,b" in normalized
    assert "ops@example.test" in normalized
    assert "3.14" in normalized
    assert "1.2.3" in normalized
    assert changes >= 1


def test_format_docx_preserves_content_images_and_merged_tables(tmp_path, monkeypatch) -> None:
    source = tmp_path / "source.docx"
    target = tmp_path / "formatted.docx"
    document = Document()
    document.add_paragraph("关于开展年度工作的通知")
    document.add_paragraph("这是正文,请访问https://example.test/a,b并核对金额3.14万元。")
    document.add_paragraph("一、工作要求")
    document.add_picture(BytesIO(ONE_PIXEL_PNG))
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "金额"
    table.cell(1, 0).merge(table.cell(1, 1)).text = "合并内容"
    document.save(source)

    monkeypatch.setattr("app.official_format._font_inventory", lambda: "方正小标宋 仿宋 楷体 黑体")
    report = format_docx(source, target)
    assert report.changed_count > 0
    assert report.compliant is True

    output = Document(target)
    section = output.sections[0]
    assert round(section.page_width.mm) == 210
    assert round(section.page_height.mm) == 297
    assert round(section.top_margin.mm) == 37
    assert round(section.bottom_margin.mm) == 35
    assert round(section.left_margin.mm) == 28
    assert round(section.right_margin.mm) == 26
    assert output.paragraphs[0].alignment == 1
    assert output.paragraphs[0].runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == "方正小标宋简体"
    assert output.paragraphs[1].paragraph_format.first_line_indent.twips == 640
    assert output.paragraphs[1].paragraph_format.line_spacing.pt == 28
    assert "这是正文，请访问https://example.test/a,b" in output.paragraphs[1].text
    assert len(output.inline_shapes) == 1
    assert output.tables[0].cell(1, 0).text == "合并内容"
    assert output.tables[0].cell(1, 1).text == "合并内容"

    with zipfile.ZipFile(target) as package:
        names = set(package.namelist())
        assert "word/media/image1.png" in names
        assert "word/footer-partyops-odd.xml" in names
        assert "word/footer-partyops-even.xml" in names
        for name in ("word/footer-partyops-odd.xml", "word/footer-partyops-even.xml"):
            footer_xml = package.read(name).decode("utf-8")
            assert "— " in footer_xml and " —" in footer_xml and " PAGE " in footer_xml
        document_xml = package.read("word/document.xml").decode("utf-8")
        assert 'w:line="560"' in document_xml
        assert 'w:firstLine="640"' in document_xml


def test_formatter_preserves_footer_emphasis_and_table_alignment(tmp_path, monkeypatch) -> None:
    """已有页脚和语义强调不能被页码或表格规范化静默覆盖。"""

    source = tmp_path / "with-footer.docx"
    target = tmp_path / "with-footer-formatted.docx"
    document = Document()
    document.add_paragraph("关于规范排版的通知")
    body = document.add_paragraph()
    body.add_run("请重点核对：")
    emphasized = body.add_run("不得改变原意")
    emphasized.bold = True
    document.sections[0].footer.paragraphs[0].text = "内部流转标识"
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "保持右对齐"
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.save(source)

    monkeypatch.setattr("app.official_format._font_inventory", lambda: "方正小标宋 仿宋 楷体 黑体")
    report = format_docx(source, target)
    assert report.compliant is True

    output = Document(target)
    assert output.paragraphs[1].runs[1].bold is True
    assert output.tables[0].cell(0, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    with zipfile.ZipFile(target) as package:
        footer_payloads = [
            package.read(name).decode("utf-8")
            for name in package.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        ]
        assert any("内部流转标识" in payload and "PAGE" in payload for payload in footer_payloads)
        settings_xml = package.read("word/settings.xml").decode("utf-8")
        assert "w:kinsoku" in settings_xml
        assert 'w:val="compressPunctuation"' in settings_xml
        document_xml = package.read("word/document.xml").decode("utf-8")
        assert "w:tcMar" in document_xml
        assert 'w:w="72"' in document_xml


def test_formatter_applies_header_roles_exact_grid_and_page_number_position(tmp_path, monkeypatch) -> None:
    """公开项目暴露过版头误判、网格和页码问题，这里用标准要素一次锁死。"""

    source = tmp_path / "official-header.docx"
    target = tmp_path / "official-header-formatted.docx"
    document = Document()
    for text in (
        "000001",
        "机密★1年",
        "特急",
        "某某委员会文件",
        "某办〔2026〕12号",
        "关于开展工作的通知",
        "某某单位：",
        "请认真落实。",
        "附件：1.有关材料",
        "某某委员会",
        "2026年8月24日",
        "（此件公开发布）",
        "抄送：有关单位。",
        "2026年8月24日印发",
    ):
        document.add_paragraph(text)
    document.save(source)

    monkeypatch.setattr("app.official_format._font_inventory", lambda: "方正小标宋 仿宋 楷体 黑体")
    report = format_docx(source, target)
    assert report.compliant is True
    assert "SPECIAL_LAYOUT_VISUAL_REVIEW_REQUIRED" in {item.code for item in report.issues}

    output = Document(target)
    assert output.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert output.paragraphs[3].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert str(output.paragraphs[3].runs[0]._element.rPr.color.val) == "FF0000"
    assert output.paragraphs[5].runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == "方正小标宋简体"
    assert output.paragraphs[6].paragraph_format.first_line_indent.twips == 0
    assert output.paragraphs[7].paragraph_format.first_line_indent.twips == 640
    assert output.paragraphs[9].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert output.paragraphs[10].alignment == WD_ALIGN_PARAGRAPH.RIGHT

    with zipfile.ZipFile(target) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        styles_xml = package.read("word/styles.xml").decode("utf-8")
        assert f'w:charSpace="{GRID_CHARACTER_SPACE}"' in document_xml
        assert f'w:footer="{PAGE_FOOTER_DISTANCE}"' in document_xml
        assert 'w:styleId="Normal"' in styles_xml
        assert 'w:sz w:val="32"' in styles_xml


def test_diagnosis_rejects_zip_slip_and_explicitly_blocks_missing_fonts(tmp_path, monkeypatch) -> None:
    unsafe = tmp_path / "unsafe.docx"
    with zipfile.ZipFile(unsafe, "w") as package:
        package.writestr("[Content_Types].xml", "<Types/>")
        package.writestr("word/document.xml", "<w:document xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'/>")
        package.writestr("../outside", "bad")
    with pytest.raises(OfficialFormatError) as caught:
        diagnose_docx(unsafe)
    assert caught.value.code == "OOXML_PATH_UNSAFE"

    valid = tmp_path / "valid.docx"
    Document().save(valid)
    monkeypatch.setattr("app.official_format._font_inventory", lambda: "")
    report = diagnose_docx(valid)
    assert report.compliant is False
    assert "FONT_CHECK_UNAVAILABLE" in {item.code for item in report.issues}


def test_official_formatter_never_imports_remote_http_clients() -> None:
    import inspect

    from app import official_format

    source = inspect.getsource(official_format)
    assert "requests." not in source
    assert "urllib.request" not in source
    assert "httpx" not in source
    assert "partyops-client://official-format" not in source
    assert not hasattr(official_format, "run_official_format_tool")

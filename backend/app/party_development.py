"""2026 年 5 月新版《中国共产党发展党员工作细则》确定性日期引擎。

国家规则保存在代码中并带版本、条款和来源；单位补充材料只附加到清单，
绝不进入日期运算。所有组织研究、政治审查和审批结论仍须人工录入。
"""

from __future__ import annotations

import calendar
import re
import uuid
from collections import defaultdict
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Mm, Pt, Twips
from sqlalchemy import select
from sqlalchemy.orm import Session

from .models import PartyDevelopmentMaterial, PartyDevelopmentProfile, User, WorkCalendarEntry, utcnow
from .schemas import (
    PartyDevelopmentCalculateRequest,
    PartyDevelopmentNodeOut,
    PartyDevelopmentResultOut,
)
from .compat import strict_zip


RULE_VERSION = "2026.05"
RULE_PUBLISHED_AT = date(2026, 5, 18)
RULE_TITLE = "中国共产党发展党员工作细则（2026年5月修订）"
RULE_SOURCE_URL = "https://www.12371.cn/2026/05/18/ARTI1779102179030620.shtml"

PHASE_LABELS = {
    "application": "申请入党",
    "activist": "入党积极分子培养考察",
    "development_object": "发展对象确定与审查",
    "probationary": "预备党员接收与教育",
    "transition": "预备党员转正",
    "archive": "材料归档",
}

STATUS_LABELS = {
    "completed": "已办理",
    "overdue": "已超期",
    "waiting_manual": "待组织确认",
    "planned": "计划安排",
}

WARNING_LEVEL_LABELS = {
    "high": "高风险",
    "medium": "提醒",
    "low": "提示",
}

CHINESE_SECTION_ORDINALS = ("一", "二", "三", "四", "五", "六", "七", "八", "九", "十")

NATIONAL_MATERIALS: dict[str, list[str]] = {
    "application": ["入党申请书", "谈话记录"],
    "activist": ["入党积极分子培养考察材料", "思想汇报"],
    "development_object": ["发展对象公示材料", "政治审查材料", "集中培训证明", "预审材料"],
    "probationary": ["入党志愿书", "支部大会决议", "党委审批材料", "入党宣誓记录"],
    "transition": ["转正申请书", "预备期考察材料", "转正支部大会决议", "党委审批材料"],
    "archive": ["发展党员全套材料归档目录"],
}

REFERENCE_PROFILE_NAME = "基层补充材料参考（待管理员确认）"
REFERENCE_MATERIALS = [
    ("activist", "季度思想汇报", "发展对象本人", "按本单位实际要求确认频次，不改变国家规则期限", 10),
    ("development_object", "个人自传", "发展对象本人", "参考上传表格，启用前须由管理员核对", 20),
    ("development_object", "三考材料", "基层党组织", "参考上传表格中的基层考察材料，不作为国家法定期限", 30),
    ("development_object", "备案报告及批复", "基层党组织", "本地流程补充，启用前确认责任主体", 40),
    ("probationary", "接收预备党员公示材料", "基层党组织", "如本单位制度要求，可作为补充材料启用", 50),
    ("transition", "转正公示材料", "基层党组织", "如本单位制度要求，可作为补充材料启用", 60),
]


def rule_metadata() -> dict[str, Any]:
    return {
        "version": RULE_VERSION,
        "published_at": RULE_PUBLISHED_AT,
        "title": RULE_TITLE,
        "source_url": RULE_SOURCE_URL,
        "principles": [
            "国家规则节点不可删除、改写或缩短",
            "组织研究、政治审查、预审和审批结论必须人工确认",
            "工作日历未配置完整时仅提供暂算结果",
            "计算结果是工作辅助，不代替组织程序和人工复核",
        ],
        "phase_labels": PHASE_LABELS,
    }


def add_months(value: date, months: int) -> date:
    """按自然月平移；目标月缺少同日时落在月末。"""

    month_index = value.year * 12 + value.month - 1 + months
    year, zero_based_month = divmod(month_index, 12)
    month = zero_based_month + 1
    day = min(value.day, calendar.monthrange(year, month)[1])
    return date(year, month, day)


class WorkdayCalendar:
    def __init__(self, entries: Iterable[WorkCalendarEntry]) -> None:
        self.overrides: dict[date, bool] = {}
        self.configured_years: set[int] = set()
        for entry in entries:
            try:
                key = date.fromisoformat(entry.date_key)
            except ValueError:
                continue
            self.overrides[key] = bool(entry.is_workday)
            self.configured_years.add(key.year)

    def is_workday(self, value: date) -> bool:
        if value in self.overrides:
            return self.overrides[value]
        return value.weekday() < 5

    def add_inclusive(self, start: date, count: int) -> tuple[date, bool]:
        if count < 1:
            raise ValueError("工作日数量必须大于零")
        current = start
        completed = 0
        years: set[int] = set()
        while True:
            years.add(current.year)
            if self.is_workday(current):
                completed += 1
                if completed == count:
                    break
            current += timedelta(days=1)
        provisional = any(year not in self.configured_years for year in years)
        return current, provisional


def ensure_reference_profile(db: Session, admin: User) -> PartyDevelopmentProfile:
    existing = db.scalar(
        select(PartyDevelopmentProfile).where(PartyDevelopmentProfile.name == REFERENCE_PROFILE_NAME)
    )
    if existing:
        return existing
    profile = PartyDevelopmentProfile(
        name=REFERENCE_PROFILE_NAME,
        description="由用户上传的旧工作表提取，仅作为本单位补充材料候选；默认停用，不参与法定日期计算。",
        source_label="入党、转正所需材料及时间顺序模板.xlsx（仅提取材料）",
        active=False,
        created_by=admin.id,
    )
    db.add(profile)
    db.flush()
    for phase, name, responsible, guidance, order in REFERENCE_MATERIALS:
        db.add(PartyDevelopmentMaterial(
            profile_id=profile.id,
            phase=phase,
            name=name,
            responsible_party=responsible,
            guidance=guidance,
            required=False,
            enabled=True,
            sort_order=order,
            created_by=admin.id,
        ))
    return profile


def profile_to_dict(db: Session, profile: PartyDevelopmentProfile) -> dict[str, Any]:
    items = db.scalars(
        select(PartyDevelopmentMaterial)
        .where(PartyDevelopmentMaterial.profile_id == profile.id)
        .order_by(PartyDevelopmentMaterial.sort_order, PartyDevelopmentMaterial.created_at)
    ).all()
    return {
        "id": profile.id,
        "name": profile.name,
        "description": profile.description,
        "source_label": profile.source_label,
        "active": profile.active,
        "version": profile.version,
        "created_by": profile.created_by,
        "created_at": profile.created_at,
        "updated_at": profile.updated_at,
        "items": items,
    }


def supplemental_materials(
    db: Session,
    profile_ids: list[str],
) -> dict[str, list[dict[str, Any]]]:
    query = select(PartyDevelopmentProfile).where(PartyDevelopmentProfile.active.is_(True))
    if profile_ids:
        query = query.where(PartyDevelopmentProfile.id.in_(profile_ids))
    profiles = db.scalars(query).all()
    if not profiles:
        return {}
    profile_map = {profile.id: profile for profile in profiles}
    items = db.scalars(
        select(PartyDevelopmentMaterial)
        .where(
            PartyDevelopmentMaterial.profile_id.in_(profile_map),
            PartyDevelopmentMaterial.enabled.is_(True),
        )
        .order_by(PartyDevelopmentMaterial.sort_order, PartyDevelopmentMaterial.created_at)
    ).all()
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for item in items:
        profile = profile_map[item.profile_id]
        grouped[item.phase].append({
            "name": item.name,
            "source": profile.name,
            "responsible_party": item.responsible_party,
            "guidance": item.guidance,
            "required": item.required,
            "national": False,
        })
    return dict(grouped)


def _materials(
    phase: str,
    supplemental: dict[str, list[dict[str, Any]]],
) -> list[dict[str, Any]]:
    national = [
        {
            "name": name,
            "source": RULE_TITLE,
            "responsible_party": "按组织程序确认",
            "guidance": "国家规则材料，具体表单和签署要求由组织部门确认",
            "required": True,
            "national": True,
        }
        for name in NATIONAL_MATERIALS.get(phase, [])
    ]
    return national + supplemental.get(phase, [])


def _node(
    *,
    key: str,
    title: str,
    phase: str,
    date_kind: str,
    value: date | None,
    end: date | None = None,
    provisional: bool = False,
    article: str,
    basis: str,
    manual: bool = False,
    completed: bool = False,
    overdue: bool = False,
    supplemental: dict[str, list[dict[str, Any]]],
) -> PartyDevelopmentNodeOut:
    if completed:
        status = "completed"
    elif overdue:
        status = "overdue"
    elif manual and value is None:
        status = "waiting_manual"
    else:
        status = "planned"
    return PartyDevelopmentNodeOut(
        key=key,
        title=title,
        phase=phase,
        date_kind=date_kind,
        date=value,
        end_date=end,
        provisional=provisional,
        status=status,
        article=article,
        basis=basis,
        requires_manual_confirmation=manual,
        materials=_materials(phase, supplemental),
    )


def calculate_party_development(
    payload: PartyDevelopmentCalculateRequest,
    calendar_entries: Iterable[WorkCalendarEntry],
    supplemental: dict[str, list[dict[str, Any]]],
    *,
    today: date | None = None,
) -> PartyDevelopmentResultOut:
    today = today or datetime.now(timezone.utc).date()
    actual = payload.actual_dates
    workdays = WorkdayCalendar(calendar_entries)
    warnings: list[dict[str, str]] = []
    nodes: list[PartyDevelopmentNodeOut] = []
    application = payload.application_date

    def warn(code: str, level: str, message: str) -> None:
        warnings.append({"code": code, "level": level, "message": message})

    talk_deadline = add_months(application, 1)
    nodes.append(_node(
        key="application", title="提交入党申请书", phase="application", date_kind="actual",
        value=application, article="第六条", basis="申请人自愿向党组织提出入党申请。",
        completed=True, supplemental=supplemental,
    ))
    nodes.append(_node(
        key="conversation_window", title="谈话建议办理窗口", phase="application", date_kind="window",
        value=application + timedelta(days=7), end=min(application + timedelta(days=20), talk_deadline),
        article="第七条", basis="系统建议窗口，不是另行增加的法定期限。",
        completed=bool(actual.conversation_date), supplemental=supplemental,
    ))
    nodes.append(_node(
        key="conversation_deadline", title="收到申请书后谈话截止日", phase="application", date_kind="deadline",
        value=talk_deadline, article="第七条", basis="党组织应当在一个月内派人同入党申请人谈话。",
        completed=bool(actual.conversation_date and actual.conversation_date <= talk_deadline),
        overdue=bool((actual.conversation_date and actual.conversation_date > talk_deadline) or (not actual.conversation_date and today > talk_deadline)),
        supplemental=supplemental,
    ))
    if actual.conversation_date and actual.conversation_date > talk_deadline:
        warn("CONVERSATION_OVERDUE", "high", "实际谈话日期超过申请书提交后一个月，请核查并说明原因。")

    nodes.append(_node(
        key="activist_date", title="确定入党积极分子", phase="activist", date_kind="manual",
        value=actual.activist_date, article="第八条", basis="须经党员推荐、群团组织推优等并由支部委员会研究决定，不能从申请日期机械推算。",
        manual=True, completed=bool(actual.activist_date), supplemental=supplemental,
    ))
    eligible_date = add_months(actual.activist_date, 12) if actual.activist_date else None
    if actual.activist_date:
        nodes.append(_node(
            key="first_half_year_assessment", title="首次半年培养考察提醒", phase="activist", date_kind="window",
            value=add_months(actual.activist_date, 6), article="第十一条", basis="党支部每半年对入党积极分子进行一次考察。",
            overdue=today > add_months(actual.activist_date, 6) and not actual.development_object_date,
            supplemental=supplemental,
        ))
    nodes.append(_node(
        key="development_object_earliest", title="列为发展对象最早日期", phase="activist", date_kind="earliest",
        value=eligible_date, article="第十三条", basis="经过一年以上培养教育和考察，基本具备党员条件后方可列为发展对象。",
        manual=eligible_date is None, supplemental=supplemental,
    ))

    publicity_end = None
    publicity_provisional = False
    if actual.publicity_start_date:
        publicity_end, publicity_provisional = workdays.add_inclusive(actual.publicity_start_date, 5)
    nodes.append(_node(
        key="development_object_publicity", title="发展对象公示", phase="development_object", date_kind="workday_window",
        value=actual.publicity_start_date, end=publicity_end, provisional=publicity_provisional,
        article="第十三条", basis="列为发展对象前应当公示，公示期不少于五个工作日。",
        manual=True, completed=bool(actual.publicity_start_date and publicity_end and today >= publicity_end),
        supplemental=supplemental,
    ))
    if publicity_provisional:
        warn("WORK_CALENDAR_INCOMPLETE", "medium", "公示相关年份未配置节假日和调休，当前按周一至周五暂算，须人工复核。")

    nodes.append(_node(
        key="development_object_date", title="确定发展对象", phase="development_object", date_kind="manual",
        value=actual.development_object_date, article="第十三条", basis="支部委员会讨论同意并报上级党委备案后确定，须人工录入。",
        manual=True, completed=bool(actual.development_object_date), supplemental=supplemental,
    ))
    if actual.development_object_date and eligible_date and actual.development_object_date < eligible_date:
        warn("DEVELOPMENT_OBJECT_TOO_EARLY", "high", f"实际确定发展对象日期早于培养考察满一年的最早日期 {eligible_date.isoformat()}。")
    if actual.development_object_date and publicity_end and actual.development_object_date < publicity_end:
        warn("PUBLICITY_NOT_COMPLETE", "high", "实际确定发展对象日期早于五个工作日公示期结束日期。")

    nodes.append(_node(
        key="political_review", title="政治审查完成", phase="development_object", date_kind="manual",
        value=actual.political_review_completed_date, article="第十六条", basis="凡未经政治审查或政治审查不合格的，不能发展入党。",
        manual=True, completed=bool(actual.political_review_completed_date), supplemental=supplemental,
    ))
    nodes.append(_node(
        key="training", title="发展对象集中培训完成", phase="development_object", date_kind="manual",
        value=actual.training_completed_date, article="第十七条", basis="基层党委或县级党委组织部门应进行不少于三天（或二十四学时）的集中培训。",
        manual=True, completed=bool(actual.training_completed_date), supplemental=supplemental,
    ))
    if (actual.training_days is not None or actual.training_hours is not None) and not (
        (actual.training_days or 0) >= 3 or (actual.training_hours or 0) >= 24
    ):
        warn("TRAINING_INSUFFICIENT", "high", "录入的集中培训天数和学时均未达到三天或二十四学时。")

    nodes.append(_node(
        key="pre_review_approved", title="上级党委预审合格", phase="development_object", date_kind="manual",
        value=actual.pre_review_approved_date, article="第十八条", basis="预审结论属于组织决定，系统不能推定。",
        manual=True, completed=bool(actual.pre_review_approved_date), supplemental=supplemental,
    ))
    branch_deadline = add_months(actual.pre_review_approved_date, 1) if actual.pre_review_approved_date else None
    nodes.append(_node(
        key="branch_acceptance_deadline", title="提交支部大会讨论截止日", phase="probationary", date_kind="deadline",
        value=branch_deadline, article="第十九条", basis="预审合格后，党支部应当在一个月内提交支部大会讨论。",
        manual=branch_deadline is None,
        completed=bool(actual.branch_acceptance_date and branch_deadline and actual.branch_acceptance_date <= branch_deadline),
        overdue=bool(branch_deadline and ((actual.branch_acceptance_date and actual.branch_acceptance_date > branch_deadline) or (not actual.branch_acceptance_date and today > branch_deadline))),
        supplemental=supplemental,
    ))
    if branch_deadline and actual.branch_acceptance_date and actual.branch_acceptance_date > branch_deadline:
        warn("BRANCH_MEETING_OVERDUE", "high", "实际接收预备党员支部大会日期超过预审合格后一个月。")

    nodes.append(_node(
        key="branch_acceptance", title="支部大会通过接收预备党员", phase="probationary", date_kind="manual",
        value=actual.branch_acceptance_date, article="第二十二条", basis="支部大会决议须人工录入；预备期从该日算起。",
        manual=True, completed=bool(actual.branch_acceptance_date), supplemental=supplemental,
    ))
    approval_normal = add_months(actual.branch_acceptance_date, 3) if actual.branch_acceptance_date else None
    approval_max = add_months(actual.branch_acceptance_date, 6) if actual.branch_acceptance_date else None
    nodes.append(_node(
        key="committee_approval", title="党委审批期限", phase="probationary", date_kind="deadline",
        value=approval_normal, end=approval_max, article="第二十五条", basis="党委一般应在三个月内审批；遇特殊情况可适当延长，但不得超过六个月。",
        manual=approval_normal is None,
        completed=bool(actual.committee_approval_date),
        overdue=bool(approval_max and ((actual.committee_approval_date and actual.committee_approval_date > approval_max) or (not actual.committee_approval_date and today > approval_max))),
        supplemental=supplemental,
    ))
    if approval_max and actual.committee_approval_date and actual.committee_approval_date > approval_max:
        warn("COMMITTEE_APPROVAL_OVERDUE", "high", "党委实际审批日期超过支部大会决议后六个月的最长时限。")

    oath_deadline = add_months(actual.committee_approval_date, 1) if actual.committee_approval_date else None
    nodes.append(_node(
        key="oath_deadline", title="组织入党宣誓建议截止日", phase="probationary", date_kind="deadline",
        value=oath_deadline, article="第二十九条", basis="上级党委批准接收后，一般应在一个月内组织入党宣誓。",
        manual=oath_deadline is None,
        completed=bool(actual.oath_date and oath_deadline and actual.oath_date <= oath_deadline),
        overdue=bool(oath_deadline and ((actual.oath_date and actual.oath_date > oath_deadline) or (not actual.oath_date and today > oath_deadline))),
        supplemental=supplemental,
    ))
    probation_end = add_months(actual.branch_acceptance_date, 12) if actual.branch_acceptance_date else None
    nodes.append(_node(
        key="probation_end", title="预备期满日期", phase="transition", date_kind="earliest",
        value=probation_end, article="第三十一条", basis="预备党员的预备期为一年，从支部大会通过其为预备党员之日算起。",
        manual=probation_end is None, supplemental=supplemental,
    ))
    nodes.append(_node(
        key="transition_application", title="提交转正申请", phase="transition", date_kind="manual",
        value=actual.transition_application_date, article="第三十二条", basis="预备期满后由本人提出书面转正申请，具体提交时间须结合实际进度确认。",
        manual=True, completed=bool(actual.transition_application_date), supplemental=supplemental,
    ))
    if probation_end and actual.transition_application_date and actual.transition_application_date < probation_end:
        warn("TRANSITION_APPLICATION_EARLY", "medium", "转正申请日期早于预备期满日期，请核对本单位办理要求和实际记录。")
    nodes.append(_node(
        key="transition_branch_meeting", title="支部大会通过转正决议", phase="transition", date_kind="manual",
        value=actual.transition_branch_meeting_date, article="第三十二条", basis="须经支部大会讨论表决，系统不能推定。",
        manual=True, completed=bool(actual.transition_branch_meeting_date), supplemental=supplemental,
    ))
    transition_deadline = add_months(actual.transition_branch_meeting_date, 3) if actual.transition_branch_meeting_date else None
    nodes.append(_node(
        key="transition_approval_deadline", title="转正决议审批截止日", phase="transition", date_kind="deadline",
        value=transition_deadline, article="第三十三条", basis="党委对党支部上报的预备党员转正决议，应当在三个月内审批。",
        manual=transition_deadline is None,
        completed=bool(actual.transition_approval_date and transition_deadline and actual.transition_approval_date <= transition_deadline),
        overdue=bool(transition_deadline and ((actual.transition_approval_date and actual.transition_approval_date > transition_deadline) or (not actual.transition_approval_date and today > transition_deadline))),
        supplemental=supplemental,
    ))
    nodes.append(_node(
        key="archive", title="党员档案材料移交归档", phase="archive", date_kind="manual",
        value=None, article="第三十六条", basis="预备党员转正后，党支部应及时将有关材料交党委存入本人档案。",
        manual=True, supplemental=supplemental,
    ))

    order_checks = [
        ("谈话", actual.conversation_date, "申请书", application),
        ("积极分子确定", actual.activist_date, "谈话", actual.conversation_date),
        ("发展对象确定", actual.development_object_date, "积极分子确定", actual.activist_date),
        ("预审合格", actual.pre_review_approved_date, "发展对象确定", actual.development_object_date),
        ("接收预备党员支部大会", actual.branch_acceptance_date, "预审合格", actual.pre_review_approved_date),
        ("党委审批", actual.committee_approval_date, "接收预备党员支部大会", actual.branch_acceptance_date),
        ("转正支部大会", actual.transition_branch_meeting_date, "接收预备党员支部大会", actual.branch_acceptance_date),
        ("转正审批", actual.transition_approval_date, "转正支部大会", actual.transition_branch_meeting_date),
    ]
    for later_label, later, earlier_label, earlier in order_checks:
        if later and earlier and later < earlier:
            warn("DATE_ORDER_INVALID", "high", f"{later_label}日期早于{earlier_label}日期，请核对实际记录。")

    manual_items = [node.title for node in nodes if node.requires_manual_confirmation and node.status == "waiting_manual"]
    return PartyDevelopmentResultOut(
        name=payload.name,
        application_date=application,
        rule_version=RULE_VERSION,
        rule_published_at=RULE_PUBLISHED_AT,
        rule_title=RULE_TITLE,
        source_url=RULE_SOURCE_URL,
        generated_at=utcnow(),
        provisional=any(node.provisional for node in nodes),
        nodes=nodes,
        warnings=warnings,
        manual_confirmation_items=manual_items,
    )


def safe_person_filename(name: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|\x00-\x1f]", "_", name).strip(" ._")
    return (cleaned or "党员发展").strip()[:60]


def _set_east_asia_font(target: Any, name: str, size: float, *, bold: bool | None = None) -> None:
    """同时写入 Word 东亚字体，避免在国产办公套件中回退为西文字体。"""

    target.font.name = name
    target.font.size = Pt(size)
    if bold is not None:
        target.font.bold = bold
    properties = target._element.get_or_add_rPr()
    properties.rFonts.set(qn("w:eastAsia"), name)


def _configure_official_document(document: Any) -> None:
    """按 GB/T 9704-2012 的 A4 版心、字号和黑白公文风配置工作参考。"""

    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(28)

    normal = document.styles["Normal"]
    _set_east_asia_font(normal, "仿宋_GB2312", 16)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(28)
    normal.paragraph_format.first_line_indent = Pt(32)

    document.settings.odd_and_even_pages_header_footer = True
    for footer, alignment in (
        (section.footer, WD_ALIGN_PARAGRAPH.RIGHT),
        (section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT),
    ):
        paragraph = footer.paragraphs[0]
        paragraph.alignment = alignment
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        start = paragraph.add_run("— ")
        _set_east_asia_font(start, "宋体", 14)
        field_run = paragraph.add_run()
        _set_east_asia_font(field_run, "宋体", 14)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        value = OxmlElement("w:t")
        value.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        field_run._r.extend([begin, instruction, separate, value, end])
        finish = paragraph.add_run(" —")
        _set_east_asia_font(finish, "宋体", 14)


def _add_official_heading(document: Any, text: str, *, level: int = 1) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 0)
    paragraph.paragraph_format.space_after = Pt(0)
    # 一级标题需要与正文相连；阶段标题紧邻固定表格，避免 Word 将整表强制上提越过页边距。
    paragraph.paragraph_format.keep_with_next = level == 1
    run = paragraph.add_run(text)
    _set_east_asia_font(run, "黑体" if level == 1 else "楷体_GB2312", 16, bold=False)
    return paragraph


def _add_official_body(document: Any, text: str, *, indent: bool = True) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(32 if indent else 0)
    run = paragraph.add_run(text)
    _set_east_asia_font(run, "仿宋_GB2312", 16)
    return paragraph


def _add_official_source_line(document: Any, source_url: str) -> Any:
    """用黑色中文标题承载制度来源链接，避免公文正文直接堆放长网址。"""

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(32)
    prefix = paragraph.add_run("制度来源：")
    _set_east_asia_font(prefix, "仿宋_GB2312", 16)

    relationship_id = paragraph.part.relate_to(source_url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), "仿宋_GB2312")
    run_properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    run_properties.append(underline)
    for tag in ("w:sz", "w:szCs"):
        size = OxmlElement(tag)
        size.set(qn("w:val"), "32")
        run_properties.append(size)
    run.append(run_properties)
    text = OxmlElement("w:t")
    text.text = "共产党员网《2026年新版细则全文》（点击查看）"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    suffix = paragraph.add_run("。")
    _set_east_asia_font(suffix, "仿宋_GB2312", 16)
    return paragraph


def _format_chinese_date(value: date | None) -> str:
    if value is None:
        return "待组织研究或人工录入"
    return f"{value.year}年{value.month}月{value.day}日"


def _format_node_date(node: PartyDevelopmentNodeOut) -> str:
    text = _format_chinese_date(node.date)
    if node.end_date:
        text = f"{text}至{_format_chinese_date(node.end_date)}"
    if node.provisional:
        text += "（暂算）"
    return text


def _format_status(status: str) -> str:
    return STATUS_LABELS.get(status, "待核实")


def _set_table_cell_text(
    cell: Any,
    text: str,
    *,
    header: bool = False,
    alignment: Any = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else alignment
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(20)
    run = paragraph.add_run(text)
    _set_east_asia_font(run, "黑体" if header else "仿宋_GB2312", 12, bold=False)


def _prevent_row_split(row: Any, *, repeat_header: bool = False) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))
    if repeat_header:
        properties.append(OxmlElement("w:tblHeader"))


def _ensure_ooxml_child(parent: Any, tag: str) -> Any:
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_ooxml_width(parent: Any, tag: str, width_dxa: int) -> None:
    width = _ensure_ooxml_child(parent, tag)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(width_dxa))


def _apply_official_table_geometry(table: Any, widths_dxa: list[int]) -> None:
    """把表宽、网格与每个单元格宽度固定为 156 mm 公文版心。"""

    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    _set_ooxml_width(properties, "w:tblW", total)
    indent = _ensure_ooxml_child(properties, "w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "0")
    layout = _ensure_ooxml_child(properties, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for column, width in strict_zip(table.columns, widths_dxa):
        column.width = Twips(width)
    for row in table.rows:
        row.height = None
        for cell, width in strict_zip(row.cells, widths_dxa):
            cell.width = Twips(width)
            cell_properties = cell._tc.get_or_add_tcPr()
            _set_ooxml_width(cell_properties, "w:tcW", width)
            margins = _ensure_ooxml_child(cell_properties, "w:tcMar")
            for side, margin_width in (("top", 80), ("bottom", 80), ("start", 100), ("end", 100)):
                margin = _ensure_ooxml_child(margins, f"w:{side}")
                margin.set(qn("w:type"), "dxa")
                margin.set(qn("w:w"), str(margin_width))


def export_result_docx(result: PartyDevelopmentResultOut, exports_dir: Path) -> Path:
    """生成 GB/T 9704-2012 风格工作参考；调用方应在响应完成后删除。"""

    exports_dir.mkdir(parents=True, exist_ok=True)
    path = exports_dir / f"{safe_person_filename(result.name)}-党员发展时间节点-{uuid.uuid4().hex[:10]}.docx"
    document = Document()
    _configure_official_document(document)
    document.core_properties.title = "党员发展时间节点与材料提示"
    document.core_properties.subject = "依据2026年5月修订的《中国共产党发展党员工作细则》生成的工作参考"
    document.core_properties.keywords = "党员发展,时间节点,材料清单,工作参考"

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    title_run = title.add_run("党员发展时间节点与材料提示")
    _set_east_asia_font(title_run, "方正小标宋简体", 22, bold=False)
    reference = document.add_paragraph()
    reference.alignment = WD_ALIGN_PARAGRAPH.CENTER
    reference.paragraph_format.first_line_indent = Pt(0)
    reference.paragraph_format.space_after = Pt(8)
    reference_run = reference.add_run("（工作参考）")
    _set_east_asia_font(reference_run, "楷体_GB2312", 16, bold=False)

    _add_official_heading(document, "一、基本信息")
    _add_official_body(document, f"人员姓名：{result.name}。")
    _add_official_body(document, f"入党申请书提交日期：{_format_chinese_date(result.application_date)}。")
    _add_official_body(document, f"制度依据：{result.rule_title}，规则版本为{result.rule_version}。")
    _add_official_source_line(document, result.source_url)
    if result.provisional:
        _add_official_body(
            document,
            "提示：部分工作日结果按周一至周五暂算，须结合法定节假日和调休安排人工复核。",
        )

    _add_official_heading(document, "二、办理时间轴")
    grouped_nodes: dict[str, list[PartyDevelopmentNodeOut]] = defaultdict(list)
    for node in result.nodes:
        grouped_nodes[node.phase].append(node)
    phase_order = [phase for phase in PHASE_LABELS if grouped_nodes.get(phase)]
    phase_order.extend(phase for phase in grouped_nodes if phase not in PHASE_LABELS)
    for phase_index, phase in enumerate(phase_order):
        ordinal = CHINESE_SECTION_ORDINALS[phase_index] if phase_index < len(CHINESE_SECTION_ORDINALS) else str(phase_index + 1)
        phase_heading = _add_official_heading(
            document,
            f"（{ordinal}）{PHASE_LABELS.get(phase, '其他待核实阶段')}",
            level=2,
        )
        # 国家规则阶段及节点数量固定；在两个自然转折点分页，避免表头孤悬和阶段跨页断裂。
        if phase in {"activist", "transition"}:
            phase_heading.paragraph_format.page_break_before = True
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        column_widths_dxa = [1984, 1701, 1134, 4025]
        headers = ["工作节点", "时间安排", "办理状态", "条款依据"]
        for cell, text in strict_zip(table.rows[0].cells, headers):
            _set_table_cell_text(cell, text, header=True)
        _prevent_row_split(table.rows[0], repeat_header=True)
        for node in grouped_nodes[phase]:
            row = table.add_row()
            values = [
                node.title,
                _format_node_date(node),
                _format_status(node.status),
                f"{node.article}：{node.basis}",
            ]
            for index, (cell, value) in enumerate(strict_zip(row.cells, values)):
                alignment = WD_ALIGN_PARAGRAPH.CENTER if index in {1, 2} else WD_ALIGN_PARAGRAPH.LEFT
                _set_table_cell_text(cell, value, alignment=alignment)
            _prevent_row_split(row)
        _apply_official_table_geometry(table, column_widths_dxa)

    _add_official_heading(document, "三、材料清单")
    seen: set[tuple[str, str]] = set()
    material_index = 0
    for node in result.nodes:
        for material in node.materials:
            identity = (node.phase, str(material.get("name", "")))
            if identity in seen:
                continue
            seen.add(identity)
            source = "国家规则" if material.get("national") else f"单位补充：{material.get('source', '')}"
            material_index += 1
            _add_official_body(
                document,
                f"{material_index}．{PHASE_LABELS.get(node.phase, '待核实阶段')}：{material.get('name')}；"
                f"来源为{source}；责任主体为{material.get('responsible_party') or '待确认'}。",
            )

    _add_official_heading(document, "四、风险提示")
    if result.warnings:
        for index, warning in enumerate(result.warnings, start=1):
            level = WARNING_LEVEL_LABELS.get(str(warning.get("level", "")), "提示")
            _add_official_body(document, f"{index}．【{level}】{warning['message']}")
    else:
        _add_official_body(document, "当前输入未发现日期顺序或期限异常，仍须由党务工作人员人工复核。")

    _add_official_heading(document, "五、待人工确认事项")
    for index, item in enumerate(result.manual_confirmation_items, start=1):
        _add_official_body(document, f"{index}．{item}。")
    if not result.manual_confirmation_items:
        _add_official_body(document, "当前没有待补录的人工确认事项。")

    generated_at = result.generated_at.astimezone(timezone(timedelta(hours=8)))
    _add_official_body(
        document,
        "说明：本材料仅供党务工作参考，不是党组织决定、审批意见或正式批复，"
        "不能替代组织研究、政治审查、会议表决和档案审核。",
    )
    generated = _add_official_body(
        document,
        f"生成时间：{generated_at.year}年{generated_at.month}月{generated_at.day}日"
        f"{generated_at.hour:02d}时{generated_at.minute:02d}分（北京时间）。",
        indent=False,
    )
    generated.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.save(path)
    return path

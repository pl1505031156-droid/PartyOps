"""Word、Excel 和迎检归档包导出。"""

from __future__ import annotations

import io
import json
import re
import zipfile
from datetime import datetime, timedelta, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from sqlalchemy import select
from sqlalchemy.orm import Session

from .config import get_settings
from .models import AttachmentVersion, FileBlob, MaterialItem, Task, User
from .schemas import serialize_api_datetime
from .storage import resolve_blob_path
from .spreadsheet_security import safe_spreadsheet_row
from .task_service import can_view_task, task_to_out, visible_tasks
from .compat import strict_zip


STATUS_LABELS = {
    "pending_receipt": "待接收",
    "pending_breakdown": "待拆解",
    "in_progress": "办理中",
    "waiting_feedback": "等待反馈",
    "pending_review": "待审核",
    "returned": "退回修改",
    "completed": "已完成",
    "archived": "已归档",
}


def _stamp() -> str:
    return datetime.now().strftime("%Y%m%d-%H%M%S")


def _safe_name(value: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]+', "_", value).strip(" .")
    return cleaned[:80] or "未命名"


def _tasks_for_kind(db: Session, user: User, kind: str) -> list[Task]:
    tasks = visible_tasks(db, user)
    now = datetime.now(timezone.utc)

    def aware(value: datetime | None) -> datetime | None:
        if value is None:
            return None
        return value if value.tzinfo else value.replace(tzinfo=timezone.utc)

    if "周" in kind:
        end = now + timedelta(days=7)
        return [
            task
            for task in tasks
            if task.status.value != "archived"
            and (due := aware(task.internal_due_at or task.formal_due_at)) is not None
            and due <= end
        ]
    if "催报" in kind:
        return [
            task
            for task in tasks
            if task.status.value
            in {"pending_receipt", "waiting_feedback", "pending_review", "returned"}
            or (
                (due := aware(task.internal_due_at or task.formal_due_at)) is not None
                and due < now
                and task.status.value not in {"completed", "archived"}
            )
        ]
    if "交接" in kind:
        return [
            task
            for task in tasks
            if task.status.value not in {"completed", "archived"}
        ]
    return tasks


def export_tasks_xlsx(db: Session, user: User, kind: str = "台账") -> Path:
    settings = get_settings()
    path = settings.exports_dir / f"党建智办-{_safe_name(kind)}-{_stamp()}.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = _safe_name(kind)[:31]
    users = {item.id: item.display_name for item in db.scalars(select(User)).all()}
    tasks = _tasks_for_kind(db, user, kind)
    if "材料目录" in kind or "缺项" in kind:
        headers = [
            "事项名称",
            "工作类别",
            "主办人",
            "材料类别",
            "材料名称",
            "是否必备",
            "状态",
            "最终文件名",
            "最终文件哈希",
        ]
        widths = [34, 16, 14, 16, 28, 12, 12, 34, 66]
    else:
        headers = [
            "事项名称",
            "工作类别",
            "标签",
            "类型",
            "状态",
            "主办人",
            "正式截止",
            "内部完成",
            "来源",
            "必备材料缺项",
            "经验与交接说明",
            "更新时间",
        ]
        widths = [36, 16, 20, 12, 14, 14, 20, 20, 28, 16, 34, 20]
    sheet.append(headers)
    header_fill = PatternFill("solid", fgColor="B42318")
    for cell in sheet[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
    if "材料目录" in kind or "缺项" in kind:
        for task in tasks:
            materials = db.scalars(
                select(MaterialItem)
                .where(MaterialItem.task_id == task.id)
                .order_by(MaterialItem.category, MaterialItem.created_at)
            ).all()
            for material in materials:
                final = db.execute(
                    select(AttachmentVersion, FileBlob)
                    .join(FileBlob, FileBlob.sha256 == AttachmentVersion.blob_sha256)
                    .where(
                        AttachmentVersion.material_item_id == material.id,
                        AttachmentVersion.is_final.is_(True),
                    )
                ).one_or_none()
                complete = material.not_applicable or final is not None
                if "缺项" in kind and (not material.required or complete):
                    continue
                sheet.append(
                    safe_spreadsheet_row([
                        task.title,
                        task.category,
                        users.get(task.owner_id, "未知"),
                        material.category,
                        material.name,
                        "是" if material.required else "否",
                        "不适用" if material.not_applicable else "齐全" if complete else "缺项",
                        (final[0].display_name or final[1].original_name) if final else "",
                        final[1].sha256 if final else "",
                    ])
                )
    else:
        for task in tasks:
            output = task_to_out(db, task, include_detail=False)
            sheet.append(
                safe_spreadsheet_row([
                    task.title,
                    task.category,
                    "、".join(task.tags or []),
                    task.task_type.value,
                    STATUS_LABELS[task.status.value],
                    users.get(task.owner_id, "未知"),
                    task.formal_due_at.isoformat(sep=" ", timespec="minutes")
                    if task.formal_due_at
                    else "",
                    task.internal_due_at.isoformat(sep=" ", timespec="minutes")
                    if task.internal_due_at
                    else "",
                    task.source,
                    output.missing_required_materials,
                    task.experience_notes,
                    task.updated_at.isoformat(sep=" ", timespec="minutes"),
                ])
            )
    for index, width in enumerate(widths, 1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="center", wrap_text=True)
    workbook.save(path)
    return path


def export_tasks_docx(db: Session, user: User, kind: str = "周工作清单") -> Path:
    settings = get_settings()
    path = settings.exports_dir / f"党建智办-{_safe_name(kind)}-{_stamp()}.docx"
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"党建智办 · {kind}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Noto Sans CJK SC"
    meta = document.add_paragraph(f"生成时间：{datetime.now():%Y-%m-%d %H:%M}")
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tasks = _tasks_for_kind(db, user, kind)
    users = {item.id: item.display_name for item in db.scalars(select(User)).all()}
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["事项", "类别", "主办", "状态", "内部节点", "正式截止", "材料"]
    for cell, label in strict_zip(table.rows[0].cells, headers):
        cell.text = label
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for cell_run in paragraph.runs:
                cell_run.bold = True
    for task in tasks:
        output = task_to_out(db, task, include_detail=False)
        cells = table.add_row().cells
        values = [
            task.title,
            task.category,
            users.get(task.owner_id, "未知"),
            STATUS_LABELS[task.status.value],
            task.internal_due_at.strftime("%m-%d %H:%M") if task.internal_due_at else "",
            task.formal_due_at.strftime("%m-%d %H:%M") if task.formal_due_at else "",
            "齐全" if output.missing_required_materials == 0 else f"缺 {output.missing_required_materials} 项",
        ]
        for cell, value in strict_zip(cells, values):
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    document.save(path)
    return path


def export_inspection_package(
    db: Session, user: User, task_ids: list[str] | None = None
) -> Path:
    settings = get_settings()
    path = settings.exports_dir / f"党建智办-迎检归档包-{_stamp()}.zip"
    tasks = visible_tasks(db, user)
    if task_ids:
        requested = set(task_ids)
        tasks = [task for task in tasks if task.id in requested]
    manifest: dict[str, object] = {
        "format": "partyops-inspection",
        "version": 1,
        "generated_at": serialize_api_datetime(datetime.now(timezone.utc)),
        "tasks": [],
    }
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED, allowZip64=True) as archive:
        directory_rows: list[list[str]] = []
        for task in tasks:
            if not can_view_task(db, task, user):
                continue
            task_dir = f"{_safe_name(task.title)}-{task.id[:8]}"
            task_entry: dict[str, object] = {
                "id": task.id,
                "title": task.title,
                "status": task.status.value,
                "materials": [],
            }
            materials = db.scalars(
                select(MaterialItem).where(MaterialItem.task_id == task.id)
            ).all()
            for material in materials:
                versions = db.execute(
                    select(AttachmentVersion, FileBlob)
                    .join(FileBlob, FileBlob.sha256 == AttachmentVersion.blob_sha256)
                    .where(AttachmentVersion.material_item_id == material.id)
                    .order_by(AttachmentVersion.version_no)
                ).all()
                material_entry: dict[str, object] = {
                    "name": material.name,
                    "required": material.required,
                    "not_applicable": material.not_applicable,
                    "files": [],
                }
                for version, blob in versions:
                    source = resolve_blob_path(blob.relative_path)
                    archive_name = (
                        f"{task_dir}/{_safe_name(material.category)}/"
                        f"v{version.version_no}-{_safe_name(version.display_name or blob.original_name)}"
                    )
                    if source.exists():
                        archive.write(source, archive_name)
                    material_entry["files"].append(
                        {
                            "path": archive_name,
                            "sha256": blob.sha256,
                            "final": version.is_final,
                            "stage": version.stage.value,
                        }
                    )
                task_entry["materials"].append(material_entry)
                complete = material.not_applicable or any(
                    version.is_final for version, _blob in versions
                )
                directory_rows.append(
                    [
                        task.title,
                        material.category,
                        material.name,
                        "是" if material.required else "否",
                        "齐全" if complete else "缺项",
                    ]
                )
            manifest["tasks"].append(task_entry)
        archive.writestr(
            "manifest.json",
            json.dumps(manifest, ensure_ascii=False, indent=2).encode("utf-8"),
        )
        archive.writestr(
            "校验清单.json",
            json.dumps(manifest, ensure_ascii=False, indent=2).encode("utf-8"),
        )
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "材料目录"
        sheet.append(["事项", "类别", "材料", "必备", "状态"])
        for row in directory_rows:
            sheet.append(safe_spreadsheet_row(row))
        sheet.freeze_panes = "A2"
        for column, width in {"A": 34, "B": 18, "C": 28, "D": 10, "E": 12}.items():
            sheet.column_dimensions[column].width = width
        workbook_buffer = io.BytesIO()
        workbook.save(workbook_buffer)
        archive.writestr("材料目录.xlsx", workbook_buffer.getvalue())
    return path

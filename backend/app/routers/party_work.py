"""三会一课、中心组学习及会议闭环专属接口。"""

from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Any, Dict, List, Literal
from urllib.parse import quote

from docx import Document
from fastapi import APIRouter, Depends, Header, Query, Request, Response
from fastapi.responses import StreamingResponse
from openpyxl import Workbook
from pydantic import BaseModel, Field
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from ..audit import write_audit
from ..database import get_session
from ..enums import Priority, Sensitivity, TaskStatus, TaskType, UserRole
from ..models import (
    BusinessDocument,
    BusinessMeeting,
    MeetingAction,
    MeetingAttendee,
    StudyPlan,
    StudyPlanTopic,
    Task,
    User,
    WorkflowTemplate,
    utcnow,
)
from ..problems import ProblemException
from ..security import get_current_user
from .business import _create_meeting_record, meeting_out
from .router_utils import client_ip, parse_if_match

router = APIRouter(tags=["party-work"])

PARTY_LIFE_TYPES = {
    "party_member_meeting": "支部党员大会",
    "branch_members": "支委会",
    "party_group": "党小组会",
    "party_class": "党课",
}
STUDY_CENTER_TYPE = "study_group"
LEDGER_STATES = {"完整", "需补充", "逾期", "待人工确认", "不适用"}

SYSTEM_TEMPLATES: tuple[dict[str, Any], ...] = (
    {
        "system_key": "party-life.party-member-meeting.v1",
        "name": "支部党员大会标准流程",
        "business_type": "party_member_meeting",
        "description": "按季度计划，记录议题、出席、表决和决议落实；系统提示不替代党组织判断。",
        "recurrence": {"kind": "quarterly", "rule_hint": "一般每季度召开一次"},
        "steps": [
            {"title": "季度计划与议题准备", "offset_days": -14, "responsible_role": "负责人", "required_document": "agenda"},
            {"title": "会议通知与材料检查", "offset_days": -7, "responsible_role": "会务", "required_document": "notice"},
            {"title": "出席与表决资格核对", "offset_days": 0, "responsible_role": "记录人", "required_document": "attendance"},
            {"title": "会议记录与表决结果", "offset_days": 0, "responsible_role": "记录人", "required_document": "minutes"},
            {"title": "决议落实与归档", "offset_days": 7, "responsible_role": "负责人", "required_document": "summary"},
        ],
    },
    {
        "system_key": "party-life.branch-members.v1",
        "name": "支委会标准流程",
        "business_type": "branch_members",
        "description": "按月计划，围绕议题准备、讨论决定、责任分工和后续任务形成闭环。",
        "recurrence": {"kind": "monthly", "rule_hint": "一般每月召开一次"},
        "steps": [
            {"title": "月度计划和议题征集", "offset_days": -10, "responsible_role": "负责人", "required_document": "agenda"},
            {"title": "议题材料准备", "offset_days": -5, "responsible_role": "承办人", "required_document": "materials"},
            {"title": "讨论决定与会议记录", "offset_days": 0, "responsible_role": "记录人", "required_document": "minutes"},
            {"title": "责任分工与后续任务", "offset_days": 2, "responsible_role": "负责人", "required_document": "summary"},
        ],
    },
    {
        "system_key": "party-life.party-group.v1",
        "name": "党小组会标准流程",
        "business_type": "party_group",
        "description": "按月计划，记录学习内容、出席、讨论情况和向支部反馈事项。",
        "recurrence": {"kind": "monthly", "rule_hint": "一般每月召开一次"},
        "steps": [
            {"title": "学习内容与通知", "offset_days": -5, "responsible_role": "负责人", "required_document": "notice"},
            {"title": "出席和讨论记录", "offset_days": 0, "responsible_role": "记录人", "required_document": "minutes"},
            {"title": "支部反馈与材料归档", "offset_days": 3, "responsible_role": "负责人", "required_document": "summary"},
        ],
    },
    {
        "system_key": "party-life.party-class.v1",
        "name": "党课标准流程",
        "business_type": "party_class",
        "description": "记录授课人、课题、教案材料、参加人员、学习反馈和归档，不虚构统一季度频率。",
        "recurrence": {"kind": "periodic", "rule_hint": "按期开展；党委（党组）书记每年至少讲一次党课"},
        "steps": [
            {"title": "确定授课人和课题", "offset_days": -14, "responsible_role": "负责人", "required_document": "agenda"},
            {"title": "教案和学习材料", "offset_days": -7, "responsible_role": "授课人", "required_document": "materials"},
            {"title": "参加人员和学习反馈", "offset_days": 0, "responsible_role": "记录人", "required_document": "minutes"},
            {"title": "成果归档", "offset_days": 5, "responsible_role": "记录人", "required_document": "summary"},
        ],
    },
    {
        "system_key": "study-center.session.v1",
        "name": "理论学习中心组集体学习研讨流程",
        "business_type": STUDY_CENTER_TYPE,
        "description": "年度计划、会前自学、集体研讨、发言考勤、成果转化和备案归档。",
        "recurrence": {"kind": "quarterly", "rule_hint": "集体学习研讨每季度不少于一次"},
        "steps": [
            {"title": "专题与会前自学材料", "offset_days": -14, "responsible_role": "学习秘书", "required_document": "materials"},
            {"title": "集体学习研讨准备", "offset_days": -5, "responsible_role": "主持人", "required_document": "agenda"},
            {"title": "发言、考勤与会议记录", "offset_days": 0, "responsible_role": "记录人", "required_document": "minutes"},
            {"title": "专题调研与成果转化", "offset_days": 15, "responsible_role": "负责人", "required_document": "summary"},
            {"title": "备案归档", "offset_days": 30, "responsible_role": "学习秘书", "required_document": "archive"},
        ],
    },
)


class PartyMeetingInput(BaseModel):
    meeting_type: str
    organization: str = Field(min_length=1, max_length=160)
    title: str = Field(min_length=1, max_length=240)
    scheduled_at: datetime | None = None
    owner_id: str | None = None
    host_id: str | None = None
    recorder_id: str | None = None
    venue: str = Field(default="", max_length=240)
    study_plan_id: str | None = None
    business_data: dict[str, Any] = Field(default_factory=dict)
    assignees: dict[str, str] = Field(default_factory=dict)


class AttendeeInput(BaseModel):
    user_id: str | None = None
    display_name: str = Field(min_length=1, max_length=120)
    role: str = Field(default="member", max_length=48)
    attendance_status: Literal["expected", "present", "absent", "leave"] = "expected"
    voting_eligible: bool = False
    note: str = Field(default="", max_length=2000)


class AttendeePatch(BaseModel):
    role: str | None = Field(default=None, max_length=48)
    attendance_status: Literal["expected", "present", "absent", "leave"] | None = None
    voting_eligible: bool | None = None
    note: str | None = Field(default=None, max_length=2000)


class ActionInput(BaseModel):
    title: str = Field(min_length=1, max_length=240)
    responsible_user_id: str | None = None
    due_at: datetime | None = None
    note: str = Field(default="", max_length=4000)
    create_task: bool = True


class ActionPatch(BaseModel):
    title: str | None = Field(default=None, min_length=1, max_length=240)
    responsible_user_id: str | None = None
    due_at: datetime | None = None
    note: str | None = Field(default=None, max_length=4000)
    status: Literal["pending", "in_progress", "completed", "cancelled"] | None = None


class StudyPlanInput(BaseModel):
    organization: str = Field(min_length=1, max_length=160)
    year: int = Field(ge=2000, le=2200)
    title: str = Field(min_length=1, max_length=240)
    group_leader_id: str | None = None
    secretary_id: str | None = None
    notes: str = Field(default="", max_length=10_000)


class StudyPlanPatch(BaseModel):
    organization: str | None = Field(default=None, min_length=1, max_length=160)
    year: int | None = Field(default=None, ge=2000, le=2200)
    title: str | None = Field(default=None, min_length=1, max_length=240)
    group_leader_id: str | None = None
    secretary_id: str | None = None
    status: Literal["draft", "active", "archived"] | None = None
    notes: str | None = Field(default=None, max_length=10_000)


class StudyTopicInput(BaseModel):
    quarter: int = Field(ge=1, le=4)
    title: str = Field(min_length=1, max_length=240)
    learning_materials: list[str] = Field(default_factory=list, max_length=100)
    research_topic: str = Field(default="", max_length=4000)
    conversion_goal: str = Field(default="", max_length=4000)
    sort_order: int = Field(default=0, ge=0, le=1000)


class StudyTopicPatch(BaseModel):
    quarter: int | None = Field(default=None, ge=1, le=4)
    title: str | None = Field(default=None, min_length=1, max_length=240)
    learning_materials: list[str] | None = Field(default=None, max_length=100)
    research_topic: str | None = Field(default=None, max_length=4000)
    conversion_goal: str | None = Field(default=None, max_length=4000)
    sort_order: int | None = Field(default=None, ge=0, le=1000)


class LifecycleReason(BaseModel):
    reason: str = Field(min_length=2, max_length=1000)


def ensure_party_work_templates(db: Session, user: User) -> list[WorkflowTemplate]:
    """幂等补齐内置模板；只升级系统模板，不覆盖用户自建模板。"""

    templates: list[WorkflowTemplate] = []
    for definition in SYSTEM_TEMPLATES:
        item = db.scalar(
            select(WorkflowTemplate).where(
                WorkflowTemplate.system_key == definition["system_key"]
            )
        )
        if item is None:
            item = WorkflowTemplate(
                **definition,
                built_in=True,
                template_version="1.0",
                created_by=user.id,
            )
            db.add(item)
            db.flush()
        templates.append(item)
    return templates


def _is_admin(user: User) -> bool:
    return user.role == UserRole.ADMIN


def _study_plan(db: Session, plan_id: str) -> StudyPlan:
    item = db.get(StudyPlan, plan_id)
    if not item:
        raise ProblemException(404, "STUDY_PLAN_NOT_FOUND", "年度学习计划不存在", "请刷新后重试。")
    return item


def _require_study_plan_modify(plan: StudyPlan, user: User) -> None:
    if not (_is_admin(user) or user.id in {plan.created_by, plan.secretary_id}):
        raise ProblemException(403, "STUDY_PLAN_MODIFY_FORBIDDEN", "没有修改年度计划的权限", "仅管理员、创建人或学习秘书可以修改。")


def _validate_study_plan_users(
    db: Session,
    *,
    group_leader_id: str | None,
    secretary_id: str | None,
) -> None:
    for user_id, label in ((group_leader_id, "组长"), (secretary_id, "学习秘书")):
        if user_id:
            target = db.get(User, user_id)
            if not target or not target.active:
                raise ProblemException(422, "STUDY_PLAN_USER_INVALID", f"{label}账号不可用", "请选择当前实例内已启用的账号。")


def can_modify_meeting(db: Session, meeting: BusinessMeeting, user: User) -> bool:
    if _is_admin(user) or user.id in {
        meeting.created_by,
        meeting.host_id,
        meeting.recorder_id,
    }:
        return True
    study = db.get(StudyPlan, meeting.study_plan_id) if meeting.study_plan_id else None
    return bool(study and user.id == study.secretary_id)


def require_meeting_modify(db: Session, meeting: BusinessMeeting, user: User) -> None:
    if not can_modify_meeting(db, meeting, user):
        raise ProblemException(
            403,
            "MEETING_MODIFY_FORBIDDEN",
            "没有修改此台账的权限",
            "仅管理员、创建人、主持人、记录人或指定学习秘书可以修改。",
        )


def require_meeting_export(db: Session, meeting: BusinessMeeting, user: User) -> None:
    if _is_admin(user) or user.id == meeting.recorder_id:
        return
    study = db.get(StudyPlan, meeting.study_plan_id) if meeting.study_plan_id else None
    if study and user.id == study.secretary_id:
        return
    raise ProblemException(
        403,
        "MEETING_EXPORT_FORBIDDEN",
        "没有导出台账的权限",
        "仅管理员、记录负责人或指定学习秘书可以导出。",
    )


def _meeting(db: Session, meeting_id: str) -> BusinessMeeting:
    item = db.get(BusinessMeeting, meeting_id)
    if not item:
        raise ProblemException(404, "MEETING_NOT_FOUND", "会议不存在", "请刷新后重试。")
    return item


def _attendee_out(item: MeetingAttendee) -> dict[str, Any]:
    return {
        "id": item.id,
        "meeting_id": item.meeting_id,
        "user_id": item.user_id,
        "display_name": item.display_name,
        "role": item.role,
        "attendance_status": item.attendance_status,
        "voting_eligible": item.voting_eligible,
        "note": item.note,
        "archived_at": item.archived_at,
        "archive_reason": item.archive_reason,
        "version": item.version,
    }


def _action_out(item: MeetingAction) -> dict[str, Any]:
    return {
        "id": item.id,
        "meeting_id": item.meeting_id,
        "title": item.title,
        "responsible_user_id": item.responsible_user_id,
        "due_at": item.due_at,
        "task_id": item.task_id,
        "status": item.status,
        "note": item.note,
        "archived_at": item.archived_at,
        "archive_reason": item.archive_reason,
        "version": item.version,
    }


def _ledger_row(db: Session, meeting: BusinessMeeting) -> dict[str, Any]:
    attendees = db.scalars(
        select(MeetingAttendee).where(
            MeetingAttendee.meeting_id == meeting.id,
            MeetingAttendee.archived_at.is_(None),
        )
    ).all()
    documents = db.scalar(
        select(func.count(BusinessDocument.id)).where(BusinessDocument.meeting_id == meeting.id)
    ) or 0
    actions = db.scalars(
        select(MeetingAction).where(
            MeetingAction.meeting_id == meeting.id,
            MeetingAction.archived_at.is_(None),
        )
    ).all()
    now = utcnow()
    scheduled = meeting.scheduled_at
    if scheduled and scheduled.tzinfo is None:
        scheduled = scheduled.replace(tzinfo=timezone.utc)
    missing: list[str] = []
    if meeting.status not in {"cancelled"}:
        if not attendees:
            missing.append("出席记录")
        if not documents:
            missing.append("会议材料")
        if meeting.status == "completed" and meeting.completed_at is None:
            missing.append("完成日期")
    overdue_actions = sum(
        1
        for action in actions
        if action.status not in {"completed", "cancelled"}
        and action.due_at is not None
        and (action.due_at if action.due_at.tzinfo else action.due_at.replace(tzinfo=timezone.utc)) < now
    )
    if meeting.status == "cancelled":
        ledger_state = "不适用"
    elif overdue_actions:
        ledger_state = "逾期"
    elif missing:
        ledger_state = "需补充"
    elif meeting.status == "completed":
        ledger_state = "完整"
    elif scheduled and scheduled < now:
        ledger_state = "待人工确认"
    else:
        ledger_state = "需补充"
    assert ledger_state in LEDGER_STATES
    return {
        **meeting_out(db, meeting),
        "ledger_state": ledger_state,
        "missing_items": missing,
        "attendee_count": len(attendees),
        "present_count": sum(1 for item in attendees if item.attendance_status == "present"),
        "document_count": int(documents),
        "action_count": len(actions),
        "overdue_action_count": overdue_actions,
    }


def _meeting_query(
    types: set[str], year: int | None, organization: str, *, include_archived: bool = False
):
    query = select(BusinessMeeting).where(BusinessMeeting.meeting_type.in_(types))
    if not include_archived:
        query = query.where(BusinessMeeting.archived_at.is_(None))
    if organization:
        query = query.where(BusinessMeeting.organization == organization)
    if year:
        query = query.where(
            BusinessMeeting.scheduled_at >= datetime(year, 1, 1, tzinfo=timezone.utc),
            BusinessMeeting.scheduled_at < datetime(year + 1, 1, 1, tzinfo=timezone.utc),
        )
    return query.order_by(BusinessMeeting.scheduled_at.desc(), BusinessMeeting.created_at.desc())


def _create_specialized_meeting(
    db: Session, payload: PartyMeetingInput, user: User, allowed: set[str]
) -> BusinessMeeting:
    if payload.meeting_type not in allowed:
        raise ProblemException(422, "MEETING_TYPE_INVALID", "会议类型无效", "请选择本模块支持的会议类型。")
    templates = ensure_party_work_templates(db, user)
    template = next(item for item in templates if item.business_type == payload.meeting_type)
    owner = db.get(User, payload.owner_id) if payload.owner_id else user
    if not owner or not owner.active:
        raise ProblemException(422, "MEETING_OWNER_INVALID", "负责人不可用", "请选择启用用户。")
    for candidate_id in (payload.host_id, payload.recorder_id):
        if candidate_id and not db.scalar(
            select(User.id).where(User.id == candidate_id, User.active.is_(True))
        ):
            raise ProblemException(422, "MEETING_ROLE_USER_INVALID", "会议角色人员不可用", "请选择仍在启用状态的用户。")
    if payload.study_plan_id:
        plan = db.get(StudyPlan, payload.study_plan_id)
        if not plan or payload.meeting_type != STUDY_CENTER_TYPE:
            raise ProblemException(422, "STUDY_PLAN_INVALID", "年度学习计划无效", "请选择当前中心组学习年度计划。")
    meeting = _create_meeting_record(
        db,
        template=template,
        meeting_type=payload.meeting_type,
        organization=payload.organization,
        title=payload.title,
        scheduled_at=payload.scheduled_at,
        owner=owner,
        created_by=user,
        assignees=payload.assignees,
    )
    meeting.host_id = payload.host_id
    meeting.recorder_id = payload.recorder_id
    meeting.venue = payload.venue.strip()
    meeting.study_plan_id = payload.study_plan_id
    meeting.business_data = payload.business_data
    return meeting


@router.get("/party-life/overview", response_model=dict)
def party_life_overview(
    year: int | None = Query(default=None, ge=2000, le=2200),
    organization: str = "",
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    year = year or utcnow().year
    ensure_party_work_templates(db, user)
    db.commit()
    rows = [_ledger_row(db, item) for item in db.scalars(_meeting_query(set(PARTY_LIFE_TYPES), year, organization)).all()]
    quarter = max(1, min(4, (utcnow().month - 1) // 3 + 1))
    return {
        "year": year,
        "current_quarter": quarter,
        "total": len(rows),
        "completed": sum(1 for row in rows if row["status"] == "completed"),
        "needs_completion": sum(1 for row in rows if row["ledger_state"] in {"需补充", "待人工确认"}),
        "overdue_actions": sum(int(row["overdue_action_count"]) for row in rows),
        "quarter_guidance": {
            "party_member_meeting": "党员大会一般每季度召开一次",
            "branch_members": "支委会一般每月召开一次",
            "party_group": "党小组会一般每月召开一次",
            "party_class": "党课按期开展，不以系统提示代替本单位计划",
        },
    }


@router.get("/party-life/meetings", response_model=List[Dict])
def party_life_meetings(
    year: int | None = Query(default=None, ge=2000, le=2200),
    organization: str = "",
    lifecycle: str = Query(default="active", pattern=r"^(active|archived|all)$"),
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> list[dict[str, Any]]:
    query = _meeting_query(
        set(PARTY_LIFE_TYPES),
        year,
        organization,
        include_archived=lifecycle != "active",
    )
    if lifecycle == "archived":
        query = query.where(BusinessMeeting.archived_at.is_not(None))
    return [_ledger_row(db, item) for item in db.scalars(query).all()]


@router.post("/party-life/meetings", response_model=dict, status_code=201)
def create_party_life_meeting(
    payload: PartyMeetingInput,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    item = _create_specialized_meeting(db, payload, user, set(PARTY_LIFE_TYPES))
    write_audit(db, user, "party_life.meeting_create", "business_meeting", item.id, {"meeting_type": item.meeting_type}, client_ip(request))
    db.commit()
    return _ledger_row(db, item)


@router.get("/party-life/ledger", response_model=List[Dict])
def party_life_ledger(
    year: int | None = Query(default=None, ge=2000, le=2200),
    organization: str = "",
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> list[dict[str, Any]]:
    return [_ledger_row(db, item) for item in db.scalars(_meeting_query(set(PARTY_LIFE_TYPES), year, organization)).all()]


@router.get("/study-center/plans", response_model=List[Dict])
def list_study_plans(
    year: int | None = Query(default=None, ge=2000, le=2200),
    organization: str = "",
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
    lifecycle: str = Query(default="active", pattern=r"^(active|archived|all)$"),
) -> list[dict[str, Any]]:
    query = select(StudyPlan)
    if lifecycle == "active":
        query = query.where(StudyPlan.archived_at.is_(None))
    elif lifecycle == "archived":
        query = query.where(StudyPlan.archived_at.is_not(None))
    if year:
        query = query.where(StudyPlan.year == year)
    if organization:
        query = query.where(StudyPlan.organization == organization)
    result: list[dict[str, Any]] = []
    for item in db.scalars(query.order_by(StudyPlan.year.desc(), StudyPlan.organization)).all():
        topics = db.scalars(select(StudyPlanTopic).where(StudyPlanTopic.plan_id == item.id).order_by(StudyPlanTopic.quarter, StudyPlanTopic.sort_order)).all()
        result.append({
            "id": item.id,
            "organization": item.organization,
            "year": item.year,
            "title": item.title,
            "group_leader_id": item.group_leader_id,
            "secretary_id": item.secretary_id,
            "status": item.status,
            "archived_at": item.archived_at,
            "archive_reason": item.archive_reason,
            "notes": item.notes,
            "version": item.version,
            "created_by": item.created_by,
            "topics": [{
                "id": topic.id,
                "quarter": topic.quarter,
                "title": topic.title,
                "learning_materials": topic.learning_materials,
                "research_topic": topic.research_topic,
                "conversion_goal": topic.conversion_goal,
                "sort_order": topic.sort_order,
                "version": topic.version,
            } for topic in topics],
        })
    return result


@router.get("/study-center/plans/{plan_id}/deletion-impact", response_model=dict)
def study_plan_deletion_impact(
    plan_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    plan = _study_plan(db, plan_id)
    _require_study_plan_modify(plan, user)
    topics = db.scalar(select(func.count(StudyPlanTopic.id)).where(StudyPlanTopic.plan_id == plan.id)) or 0
    meetings = db.scalar(select(func.count(BusinessMeeting.id)).where(BusinessMeeting.study_plan_id == plan.id)) or 0
    documents = db.scalar(select(func.count(BusinessDocument.id)).join(BusinessMeeting, BusinessDocument.meeting_id == BusinessMeeting.id).where(BusinessMeeting.study_plan_id == plan.id)) or 0
    return {"topics": topics, "meetings": meetings, "documents": documents, "recoverable": True, "physical_delete": False}


@router.delete("/study-center/plans/{plan_id}", response_model=dict)
def archive_study_plan(
    plan_id: str,
    payload: LifecycleReason,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    plan = _study_plan(db, plan_id)
    _require_study_plan_modify(plan, user)
    if plan.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "年度学习计划已更新", "请刷新后重试。")
    if not plan.archived_at:
        plan.status_before_archive = plan.status
        plan.status = "archived"
        plan.archived_at = utcnow()
        plan.archived_by = user.id
        plan.archive_reason = payload.reason.strip()
        plan.version += 1
        write_audit(db, user, "study_center.plan_archive", "study_plan", plan.id, {"reason": plan.archive_reason}, client_ip(request))
        db.commit()
    return {"id": plan.id, "archived": True, "version": plan.version}


@router.post("/study-center/plans/{plan_id}/restore", response_model=dict)
def restore_study_plan(
    plan_id: str,
    payload: LifecycleReason,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    plan = _study_plan(db, plan_id)
    _require_study_plan_modify(plan, user)
    if plan.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "年度学习计划已更新", "请刷新后重试。")
    if plan.archived_at:
        plan.status = plan.status_before_archive if plan.status_before_archive != "archived" else "draft"
        plan.archived_at = None
        plan.archived_by = None
        plan.archive_reason = ""
        plan.version += 1
        write_audit(db, user, "study_center.plan_restore", "study_plan", plan.id, {"reason": payload.reason.strip()}, client_ip(request))
        db.commit()
    return {"id": plan.id, "archived": False, "version": plan.version}


@router.post("/study-center/plans", response_model=dict, status_code=201)
def create_study_plan(
    payload: StudyPlanInput,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    if not _is_admin(user):
        raise ProblemException(403, "STUDY_PLAN_CREATE_FORBIDDEN", "没有创建年度计划的权限", "请联系系统管理员。")
    if db.scalar(select(StudyPlan.id).where(StudyPlan.organization == payload.organization.strip(), StudyPlan.year == payload.year)):
        raise ProblemException(409, "STUDY_PLAN_EXISTS", "本年度学习计划已存在", "请打开已有计划继续维护。")
    _validate_study_plan_users(
        db,
        group_leader_id=payload.group_leader_id,
        secretary_id=payload.secretary_id,
    )
    values = payload.model_dump()
    values["organization"] = payload.organization.strip()
    values["title"] = payload.title.strip()
    item = StudyPlan(**values, created_by=user.id)
    db.add(item)
    db.flush()
    write_audit(db, user, "study_center.plan_create", "study_plan", item.id, {"year": item.year}, client_ip(request))
    db.commit()
    return list_study_plans(
        year=item.year,
        organization=item.organization,
        lifecycle="active",
        _user=user,
        db=db,
    )[0]


@router.patch("/study-center/plans/{plan_id}", response_model=dict)
def patch_study_plan(
    plan_id: str,
    payload: StudyPlanPatch,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    plan = _study_plan(db, plan_id)
    _require_study_plan_modify(plan, user)
    if plan.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "年度学习计划已更新", "请刷新后核对最新内容。")
    changes = payload.model_dump(exclude_unset=True)
    organization = str(changes.get("organization", plan.organization)).strip()
    year = int(changes.get("year", plan.year))
    duplicate = db.scalar(
        select(StudyPlan.id).where(
            StudyPlan.organization == organization,
            StudyPlan.year == year,
            StudyPlan.id != plan.id,
        )
    )
    if duplicate:
        raise ProblemException(409, "STUDY_PLAN_EXISTS", "本年度学习计划已存在", "请合并到已有计划。")
    group_leader_id = changes.get("group_leader_id", plan.group_leader_id)
    secretary_id = changes.get("secretary_id", plan.secretary_id)
    _validate_study_plan_users(
        db,
        group_leader_id=group_leader_id,
        secretary_id=secretary_id,
    )
    for key, value in changes.items():
        if key in {"organization", "title"} and isinstance(value, str):
            value = value.strip()
        setattr(plan, key, value)
    plan.version += 1
    write_audit(
        db,
        user,
        "study_center.plan_update",
        "study_plan",
        plan.id,
        {"fields": sorted(changes)},
        client_ip(request),
    )
    db.commit()
    return list_study_plans(
        year=plan.year,
        organization=plan.organization,
        lifecycle="active",
        _user=user,
        db=db,
    )[0]


@router.post("/study-center/plans/{plan_id}/topics", response_model=dict, status_code=201)
def create_study_topic(
    plan_id: str,
    payload: StudyTopicInput,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    plan = _study_plan(db, plan_id)
    _require_study_plan_modify(plan, user)
    item = StudyPlanTopic(plan_id=plan.id, **payload.model_dump())
    db.add(item)
    plan.version += 1
    write_audit(db, user, "study_center.topic_create", "study_plan", plan.id, {"quarter": item.quarter}, client_ip(request))
    db.commit()
    return {"id": item.id, **payload.model_dump(), "version": item.version}


@router.patch("/study-center/plans/{plan_id}/topics/{topic_id}", response_model=dict)
def patch_study_topic(
    plan_id: str,
    topic_id: str,
    payload: StudyTopicPatch,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    plan = _study_plan(db, plan_id)
    _require_study_plan_modify(plan, user)
    topic = db.get(StudyPlanTopic, topic_id)
    if not topic or topic.plan_id != plan.id:
        raise ProblemException(404, "STUDY_TOPIC_NOT_FOUND", "学习专题不存在", "请刷新后重试。")
    if topic.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "学习专题已更新", "请刷新后核对最新内容。")
    changes = payload.model_dump(exclude_unset=True)
    for key, value in changes.items():
        setattr(topic, key, value.strip() if isinstance(value, str) else value)
    topic.version += 1
    plan.version += 1
    write_audit(db, user, "study_center.topic_update", "study_plan", plan.id, {"topic_id": topic.id, "fields": sorted(changes)}, client_ip(request))
    db.commit()
    return {
        "id": topic.id,
        "quarter": topic.quarter,
        "title": topic.title,
        "learning_materials": topic.learning_materials,
        "research_topic": topic.research_topic,
        "conversion_goal": topic.conversion_goal,
        "sort_order": topic.sort_order,
        "version": topic.version,
    }


@router.delete(
    "/study-center/plans/{plan_id}/topics/{topic_id}",
    status_code=204,
    response_class=Response,
)
def delete_study_topic(
    plan_id: str,
    topic_id: str,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> Response:
    plan = _study_plan(db, plan_id)
    _require_study_plan_modify(plan, user)
    topic = db.get(StudyPlanTopic, topic_id)
    if not topic or topic.plan_id != plan.id:
        raise ProblemException(404, "STUDY_TOPIC_NOT_FOUND", "学习专题不存在", "请刷新后重试。")
    if topic.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "学习专题已更新", "请刷新后核对最新内容。")
    detail = {"topic_id": topic.id, "quarter": topic.quarter, "title": topic.title}
    db.delete(topic)
    plan.version += 1
    write_audit(db, user, "study_center.topic_delete", "study_plan", plan.id, detail, client_ip(request))
    db.commit()
    return Response(status_code=204)


@router.get("/study-center/sessions", response_model=List[Dict])
def list_study_sessions(
    year: int | None = Query(default=None, ge=2000, le=2200),
    organization: str = "",
    lifecycle: str = Query(default="active", pattern=r"^(active|archived|all)$"),
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> list[dict[str, Any]]:
    query = _meeting_query(
        {STUDY_CENTER_TYPE},
        year,
        organization,
        include_archived=lifecycle != "active",
    )
    if lifecycle == "archived":
        query = query.where(BusinessMeeting.archived_at.is_not(None))
    return [_ledger_row(db, item) for item in db.scalars(query).all()]


@router.post("/study-center/sessions", response_model=dict, status_code=201)
def create_study_session(
    payload: PartyMeetingInput,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    item = _create_specialized_meeting(db, payload, user, {STUDY_CENTER_TYPE})
    write_audit(db, user, "study_center.session_create", "business_meeting", item.id, {"study_plan_id": item.study_plan_id}, client_ip(request))
    db.commit()
    return _ledger_row(db, item)


@router.get("/study-center/ledger", response_model=List[Dict])
def study_center_ledger(
    year: int | None = Query(default=None, ge=2000, le=2200),
    organization: str = "",
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> list[dict[str, Any]]:
    return [_ledger_row(db, item) for item in db.scalars(_meeting_query({STUDY_CENTER_TYPE}, year, organization)).all()]


@router.get("/business-meetings/{meeting_id}/attendees", response_model=List[Dict])
def list_attendees(
    meeting_id: str,
    lifecycle: str = Query(default="active", pattern=r"^(active|archived|all)$"),
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> list[dict[str, Any]]:
    _meeting(db, meeting_id)
    query = select(MeetingAttendee).where(MeetingAttendee.meeting_id == meeting_id)
    if lifecycle == "active":
        query = query.where(MeetingAttendee.archived_at.is_(None))
    elif lifecycle == "archived":
        query = query.where(MeetingAttendee.archived_at.is_not(None))
    return [_attendee_out(item) for item in db.scalars(query.order_by(MeetingAttendee.created_at)).all()]


@router.post("/business-meetings/{meeting_id}/attendees", response_model=dict, status_code=201)
def add_attendee(
    meeting_id: str,
    payload: AttendeeInput,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting = _meeting(db, meeting_id)
    require_meeting_modify(db, meeting, user)
    if meeting.archived_at:
        raise ProblemException(409, "MEETING_ARCHIVED", "会议已归档", "请先恢复会议，再添加出席记录。")
    if payload.user_id and not db.get(User, payload.user_id):
        raise ProblemException(422, "ATTENDEE_USER_INVALID", "参会账号不存在", "可以清空账号并只记录姓名。")
    item = MeetingAttendee(meeting_id=meeting.id, **payload.model_dump())
    db.add(item)
    write_audit(db, user, "meeting.attendee_create", "business_meeting", meeting.id, {"role": item.role}, client_ip(request))
    db.commit()
    return _attendee_out(item)


@router.patch("/business-meetings/{meeting_id}/attendees/{attendee_id}", response_model=dict)
def patch_attendee(
    meeting_id: str,
    attendee_id: str,
    payload: AttendeePatch,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting = _meeting(db, meeting_id)
    require_meeting_modify(db, meeting, user)
    item = db.get(MeetingAttendee, attendee_id)
    if not item or item.meeting_id != meeting.id:
        raise ProblemException(404, "ATTENDEE_NOT_FOUND", "参会记录不存在", "请刷新后重试。")
    if meeting.archived_at or item.archived_at:
        raise ProblemException(409, "ATTENDEE_ARCHIVED", "出席记录已归档", "请先恢复会议和出席记录。")
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "参会记录已更新", "请刷新后重试。")
    for key, value in payload.model_dump(exclude_unset=True).items():
        setattr(item, key, value)
    item.version += 1
    write_audit(db, user, "meeting.attendee_update", "business_meeting", meeting.id, {"attendee_id": item.id}, client_ip(request))
    db.commit()
    return _attendee_out(item)


@router.get("/business-meetings/{meeting_id}/attendees/{attendee_id}/deletion-impact", response_model=dict)
def attendee_deletion_impact(
    meeting_id: str,
    attendee_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, bool]:
    meeting = _meeting(db, meeting_id)
    require_meeting_modify(db, meeting, user)
    item = db.get(MeetingAttendee, attendee_id)
    if not item or item.meeting_id != meeting.id:
        raise ProblemException(404, "ATTENDEE_NOT_FOUND", "参会记录不存在", "请刷新后重试。")
    return {
        "attendance_statistics": item.attendance_status == "present",
        "voting_record": item.voting_eligible,
        "recoverable": True,
        "physical_delete": False,
    }


@router.delete("/business-meetings/{meeting_id}/attendees/{attendee_id}", response_model=dict)
def archive_attendee(
    meeting_id: str,
    attendee_id: str,
    payload: LifecycleReason,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting = _meeting(db, meeting_id)
    require_meeting_modify(db, meeting, user)
    item = db.get(MeetingAttendee, attendee_id)
    if not item or item.meeting_id != meeting.id:
        raise ProblemException(404, "ATTENDEE_NOT_FOUND", "参会记录不存在", "请刷新后重试。")
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "出席记录已更新", "请刷新后重试。")
    if not item.archived_at:
        item.archived_at = utcnow()
        item.archived_by = user.id
        item.archive_reason = payload.reason.strip()
        item.version += 1
        write_audit(db, user, "meeting.attendee_archive", "business_meeting", meeting.id, {"attendee_id": item.id, "reason": item.archive_reason}, client_ip(request))
        db.commit()
    return _attendee_out(item)


@router.post("/business-meetings/{meeting_id}/attendees/{attendee_id}/restore", response_model=dict)
def restore_attendee(
    meeting_id: str,
    attendee_id: str,
    payload: LifecycleReason,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting = _meeting(db, meeting_id)
    require_meeting_modify(db, meeting, user)
    if meeting.archived_at:
        raise ProblemException(409, "MEETING_ARCHIVED", "会议仍在归档中", "请先恢复会议，再恢复出席记录。")
    item = db.get(MeetingAttendee, attendee_id)
    if not item or item.meeting_id != meeting.id:
        raise ProblemException(404, "ATTENDEE_NOT_FOUND", "参会记录不存在", "请刷新后重试。")
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "出席记录已更新", "请刷新后重试。")
    if item.archived_at:
        item.archived_at = None
        item.archived_by = None
        item.archive_reason = ""
        item.version += 1
        write_audit(db, user, "meeting.attendee_restore", "business_meeting", meeting.id, {"attendee_id": item.id, "reason": payload.reason.strip()}, client_ip(request))
        db.commit()
    return _attendee_out(item)


@router.get("/business-meetings/{meeting_id}/actions", response_model=List[Dict])
def list_actions(
    meeting_id: str,
    lifecycle: str = Query(default="active", pattern=r"^(active|archived|all)$"),
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> list[dict[str, Any]]:
    _meeting(db, meeting_id)
    query = select(MeetingAction).where(MeetingAction.meeting_id == meeting_id)
    if lifecycle == "active":
        query = query.where(MeetingAction.archived_at.is_(None))
    elif lifecycle == "archived":
        query = query.where(MeetingAction.archived_at.is_not(None))
    return [_action_out(item) for item in db.scalars(query.order_by(MeetingAction.due_at, MeetingAction.created_at)).all()]


@router.post("/business-meetings/{meeting_id}/actions", response_model=dict, status_code=201)
def add_action(
    meeting_id: str,
    payload: ActionInput,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting = _meeting(db, meeting_id)
    require_meeting_modify(db, meeting, user)
    if meeting.archived_at:
        raise ProblemException(409, "MEETING_ARCHIVED", "会议已归档", "请先恢复会议，再新增落实项。")
    responsible = db.get(User, payload.responsible_user_id) if payload.responsible_user_id else user
    if not responsible or not responsible.active:
        raise ProblemException(422, "ACTION_OWNER_INVALID", "落实负责人不可用", "请选择启用用户。")
    task: Task | None = None
    if payload.create_task:
        task = Task(
            title=payload.title.strip(),
            description=f"来源：{meeting.title}会议决议",
            task_type=TaskType.STANDARD,
            status=TaskStatus.IN_PROGRESS,
            sensitivity=Sensitivity.NORMAL,
            priority=Priority.NORMAL,
            source="会议决议",
            source_kind="meeting_action",
            category="党务落实",
            formal_due_at=payload.due_at,
            internal_due_at=payload.due_at,
            owner_id=responsible.id,
            created_by=user.id,
            updated_by=user.id,
        )
        db.add(task)
        db.flush()
    item = MeetingAction(
        meeting_id=meeting.id,
        title=payload.title.strip(),
        responsible_user_id=responsible.id,
        due_at=payload.due_at,
        task_id=task.id if task else None,
        note=payload.note,
        created_by=user.id,
    )
    db.add(item)
    write_audit(db, user, "meeting.action_create", "business_meeting", meeting.id, {"task_id": item.task_id}, client_ip(request))
    db.commit()
    return _action_out(item)


@router.patch("/business-meetings/{meeting_id}/actions/{action_id}", response_model=dict)
def patch_action(
    meeting_id: str,
    action_id: str,
    payload: ActionPatch,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting = _meeting(db, meeting_id)
    item = db.get(MeetingAction, action_id)
    if not item or item.meeting_id != meeting.id:
        raise ProblemException(404, "ACTION_NOT_FOUND", "决议落实项不存在", "请刷新后重试。")
    if meeting.archived_at or item.archived_at:
        raise ProblemException(409, "ACTION_ARCHIVED", "落实项已归档", "请先恢复会议和落实项。")
    if not (can_modify_meeting(db, meeting, user) or item.responsible_user_id == user.id):
        raise ProblemException(403, "ACTION_MODIFY_FORBIDDEN", "没有修改落实项的权限", "仅会议记录负责人或本落实项负责人可以修改。")
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "落实项已更新", "请刷新后重试。")
    changes = payload.model_dump(exclude_unset=True)
    if "responsible_user_id" in changes and changes["responsible_user_id"]:
        responsible = db.get(User, changes["responsible_user_id"])
        if not responsible or not responsible.active:
            raise ProblemException(
                422,
                "ACTION_OWNER_INVALID",
                "落实负责人不可用",
                "请选择当前实例内已启用的账号。",
            )
    for key, value in changes.items():
        setattr(item, key, value.strip() if isinstance(value, str) else value)
    task = db.get(Task, item.task_id) if item.task_id else None
    if task:
        if "title" in changes:
            task.title = item.title
        if "due_at" in changes:
            task.formal_due_at = item.due_at
            task.internal_due_at = item.due_at
        if "responsible_user_id" in changes and item.responsible_user_id:
            task.owner_id = item.responsible_user_id
        if item.status == "completed":
            task.status = TaskStatus.COMPLETED
        task.version += 1
        task.updated_by = user.id
    item.version += 1
    write_audit(db, user, "meeting.action_update", "business_meeting", meeting.id, {"action_id": item.id, "fields": sorted(changes)}, client_ip(request))
    db.commit()
    return _action_out(item)


@router.get("/business-meetings/{meeting_id}/actions/{action_id}/deletion-impact", response_model=dict)
def action_deletion_impact(
    meeting_id: str,
    action_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, bool]:
    meeting = _meeting(db, meeting_id)
    require_meeting_modify(db, meeting, user)
    item = db.get(MeetingAction, action_id)
    if not item or item.meeting_id != meeting.id:
        raise ProblemException(404, "ACTION_NOT_FOUND", "决议落实项不存在", "请刷新后重试。")
    return {
        "linked_task": bool(item.task_id),
        "reminders_affected": bool(item.task_id),
        "recoverable": True,
        "physical_delete": False,
    }


@router.delete("/business-meetings/{meeting_id}/actions/{action_id}", response_model=dict)
def archive_action(
    meeting_id: str,
    action_id: str,
    payload: LifecycleReason,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting = _meeting(db, meeting_id)
    require_meeting_modify(db, meeting, user)
    item = db.get(MeetingAction, action_id)
    if not item or item.meeting_id != meeting.id:
        raise ProblemException(404, "ACTION_NOT_FOUND", "决议落实项不存在", "请刷新后重试。")
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "落实项已更新", "请刷新后重试。")
    if not item.archived_at:
        item.archived_at = utcnow()
        item.archived_by = user.id
        item.archive_reason = payload.reason.strip()
        item.version += 1
        if item.task_id and (task := db.get(Task, item.task_id)):
            item.task_status_before_archive = task.status.value
            task.status = TaskStatus.ARCHIVED
            task.version += 1
            task.updated_by = user.id
        write_audit(db, user, "meeting.action_archive", "business_meeting", meeting.id, {"action_id": item.id, "reason": item.archive_reason, "task_id": item.task_id}, client_ip(request))
        db.commit()
    return _action_out(item)


@router.post("/business-meetings/{meeting_id}/actions/{action_id}/restore", response_model=dict)
def restore_action(
    meeting_id: str,
    action_id: str,
    payload: LifecycleReason,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting = _meeting(db, meeting_id)
    require_meeting_modify(db, meeting, user)
    if meeting.archived_at:
        raise ProblemException(409, "MEETING_ARCHIVED", "会议仍在归档中", "请先恢复会议，再恢复落实项。")
    item = db.get(MeetingAction, action_id)
    if not item or item.meeting_id != meeting.id:
        raise ProblemException(404, "ACTION_NOT_FOUND", "决议落实项不存在", "请刷新后重试。")
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "落实项已更新", "请刷新后重试。")
    if item.archived_at:
        item.archived_at = None
        item.archived_by = None
        item.archive_reason = ""
        item.version += 1
        if item.task_id and (task := db.get(Task, item.task_id)) and task.status == TaskStatus.ARCHIVED:
            try:
                task.status = TaskStatus(item.task_status_before_archive)
            except ValueError:
                task.status = TaskStatus.IN_PROGRESS
            task.version += 1
            task.updated_by = user.id
        item.task_status_before_archive = ""
        write_audit(db, user, "meeting.action_restore", "business_meeting", meeting.id, {"action_id": item.id, "reason": payload.reason.strip(), "task_id": item.task_id}, client_ip(request))
        db.commit()
    return _action_out(item)


def _safe_cell(value: object) -> str:
    text = "" if value is None else str(value)
    return "'" + text if text.startswith(("=", "+", "-", "@")) else text


def _ledger_export(rows: list[dict[str, Any]], *, title: str, output_format: str) -> StreamingResponse:
    if output_format == "xlsx":
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "台账"
        headers = ["类型", "组织", "标题", "计划时间", "状态", "台账状态", "出席人数", "材料数", "逾期落实项"]
        sheet.append(headers)
        for row in rows:
            sheet.append([_safe_cell(value) for value in (
                row["meeting_type_label"], row["organization"], row["title"], row["scheduled_at"], row["status"], row["ledger_state"], row["present_count"], row["document_count"], row["overdue_action_count"],
            )])
        stream = io.BytesIO()
        workbook.save(stream)
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        filename = f"{title}.xlsx"
    else:
        document = Document()
        document.add_heading(title, level=1)
        document.add_paragraph("提示：系统台账用于材料完整性提醒，不替代党组织对会议有效性和程序合规性的判断。")
        for row in rows:
            document.add_heading(str(row["title"]), level=2)
            document.add_paragraph(f"类型：{row['meeting_type_label']}　组织：{row['organization']}")
            document.add_paragraph(f"计划时间：{row['scheduled_at'] or '待补充'}　台账状态：{row['ledger_state']}")
            if row["missing_items"]:
                document.add_paragraph("待补充：" + "、".join(row["missing_items"]))
        stream = io.BytesIO()
        document.save(stream)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{title}.docx"
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}", "Cache-Control": "no-store"},
    )


@router.get("/party-life/ledger/export.{output_format}")
def export_party_life_ledger(
    output_format: Literal["docx", "xlsx"],
    request: Request,
    year: int | None = Query(default=None, ge=2000, le=2200),
    organization: str = "",
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> StreamingResponse:
    meetings = db.scalars(_meeting_query(set(PARTY_LIFE_TYPES), year, organization)).all()
    for meeting in meetings:
        require_meeting_export(db, meeting, user)
    rows = [_ledger_row(db, meeting) for meeting in meetings]
    write_audit(db, user, "party_life.ledger_export", "business_meeting", None, {"format": output_format, "count": len(rows)}, client_ip(request))
    db.commit()
    return _ledger_export(rows, title="三会一课年度台账", output_format=output_format)


@router.get("/study-center/ledger/export.{output_format}")
def export_study_center_ledger(
    output_format: Literal["docx", "xlsx"],
    request: Request,
    year: int | None = Query(default=None, ge=2000, le=2200),
    organization: str = "",
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> StreamingResponse:
    meetings = db.scalars(_meeting_query({STUDY_CENTER_TYPE}, year, organization)).all()
    for meeting in meetings:
        require_meeting_export(db, meeting, user)
    rows = [_ledger_row(db, meeting) for meeting in meetings]
    write_audit(db, user, "study_center.ledger_export", "business_meeting", None, {"format": output_format, "count": len(rows)}, client_ip(request))
    db.commit()
    return _ledger_export(rows, title="理论学习中心组年度台账", output_format=output_format)

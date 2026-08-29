"""会议筹备、结构化在线文档与年度统计。"""

from __future__ import annotations

import hashlib
import io
from datetime import datetime, timedelta, timezone
from decimal import Decimal, InvalidOperation
from typing import Any, List
from urllib.parse import quote

from docx import Document
from fastapi import APIRouter, Body, Depends, File, Header, Query, Request, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy import func, select
from sqlalchemy.orm import Session
from starlette.concurrency import run_in_threadpool

from ..audit import write_audit
from ..database import get_session
from ..enums import Priority, Sensitivity, TaskStatus, TaskType, UserRole
from ..meeting_imports import (
    MAX_MEETING_IMPORT_BYTES,
    extract_meeting_text,
    propose_meeting,
)
from ..models import (
    BusinessDocument,
    BusinessDocumentRevision,
    BusinessMeeting,
    MeetingAction,
    MeetingAttendee,
    MeetingImportDraft,
    MeetingTopic,
    StudyPlan,
    Task,
    TaskStep,
    User,
    WorkflowTemplate,
    utcnow,
)
from ..problems import ProblemException
from ..security import get_current_user, require_admin
from ..time_utils import beijing_iso
from .router_utils import client_ip, parse_if_match

router = APIRouter(tags=["business-workflows"])

MEETING_TYPES = {
    "party_committee": "党委会",
    "branch_members": "支委会",
    "party_member_meeting": "党员大会",
    "party_group": "党小组会",
    "party_class": "党课",
    "study_group": "党委（党组）理论学习中心组学习",
    "theme_party_day": "主题党日",
    "organization_life": "组织生活会",
}

COMMITTEE_STEPS = [
    {"title": "会议议程起草", "offset_days": -20, "responsible_role": "承办人", "required_document": "agenda"},
    {"title": "人员通知", "offset_days": -10, "responsible_role": "会务", "required_document": "notice"},
    {"title": "会议资料准备", "offset_days": -7, "responsible_role": "承办人", "required_document": "materials"},
    {"title": "资料打印", "offset_days": -1, "responsible_role": "会务", "required_document": ""},
    {"title": "会议记录", "offset_days": 0, "responsible_role": "记录人", "required_document": "minutes"},
    {"title": "会后纪要", "offset_days": 3, "responsible_role": "记录人", "required_document": "summary"},
]


class WorkflowTemplateInput(BaseModel):
    name: str = Field(min_length=2, max_length=160)
    business_type: str = Field(min_length=2, max_length=48)
    description: str = Field(default="", max_length=4000)
    steps: list[dict[str, Any]] = Field(default_factory=list, max_length=50)
    recurrence: dict[str, Any] = Field(default_factory=dict)


class MeetingInput(BaseModel):
    meeting_type: str
    organization: str = Field(min_length=1, max_length=160)
    title: str = Field(min_length=1, max_length=240)
    scheduled_at: datetime | None = None
    workflow_template_id: str | None = None
    owner_id: str | None = None
    assignees: dict[str, str] = Field(default_factory=dict)
    recurrence_key: str = Field(default="", max_length=120)


class MeetingPatch(BaseModel):
    organization: str | None = Field(default=None, min_length=1, max_length=160)
    title: str | None = Field(default=None, min_length=1, max_length=240)
    scheduled_at: datetime | None = None
    status: str | None = Field(default=None, pattern=r"^(planned|in_progress|completed|cancelled|archived)$")


class MeetingImportPatch(BaseModel):
    """用户确认后的会议候选；模型或规则候选不会直接写入正式台账。"""

    meeting_type: str = "party_committee"
    organization: str = Field(min_length=1, max_length=160)
    title: str = Field(min_length=1, max_length=240)
    scheduled_at: datetime | None = None
    venue: str = Field(default="", max_length=240)
    host_name: str = Field(default="", max_length=120)
    attendee_text: str = Field(default="", max_length=2_000)
    owner_id: str | None = None
    topics: list[str] = Field(default_factory=list, max_length=50)


class TopicInput(BaseModel):
    title: str = Field(min_length=1, max_length=240)
    review_result: str = Field(default="", max_length=20_000)
    amount: str = "0"
    reviewed: bool = False
    amount_confirmed: bool = False


class TopicPatch(BaseModel):
    title: str | None = Field(default=None, min_length=1, max_length=240)
    review_result: str | None = Field(default=None, max_length=20_000)
    amount: str | None = None
    reviewed: bool | None = None
    amount_confirmed: bool | None = None


class DocumentInput(BaseModel):
    meeting_id: str | None = None
    task_step_id: str | None = None
    document_type: str = Field(
        pattern=r"^(agenda|notice|minutes|summary|materials|attendance|archive|other)$"
    )
    title: str = Field(min_length=1, max_length=240)
    content: dict[str, Any] = Field(default_factory=dict)


class DocumentPatch(BaseModel):
    title: str | None = Field(default=None, min_length=1, max_length=240)
    content: dict[str, Any] | None = None
    change_note: str = Field(default="", max_length=500)


class LifecycleReason(BaseModel):
    reason: str = Field(min_length=2, max_length=1000)


def ensure_committee_template(db: Session, user: User) -> WorkflowTemplate:
    template = db.scalar(select(WorkflowTemplate).where(WorkflowTemplate.name == "党委会筹备（六步）"))
    if template:
        return template
    template = WorkflowTemplate(
        name="党委会筹备（六步）",
        system_key="other-meeting.party-committee.v1",
        business_type="party_committee",
        description="议程、通知、资料、打印、记录、纪要六个固定步骤，可按月复用。",
        steps=COMMITTEE_STEPS,
        recurrence={"kind": "monthly", "enabled": False},
        built_in=True,
        template_version="1.0",
        created_by=user.id,
    )
    db.add(template)
    db.flush()
    return template


def template_out(item: WorkflowTemplate) -> dict[str, Any]:
    return {
        "id": item.id,
        "name": item.name,
        "system_key": item.system_key,
        "business_type": item.business_type,
        "description": item.description,
        "steps": item.steps,
        "recurrence": item.recurrence,
        "active": item.active,
        "built_in": item.built_in,
        "template_version": item.template_version,
        "version": item.version,
    }


def meeting_out(db: Session, item: BusinessMeeting) -> dict[str, Any]:
    steps = db.scalars(select(TaskStep).where(TaskStep.task_id == item.task_id).order_by(TaskStep.sort_order)).all() if item.task_id else []
    topics = db.scalars(select(MeetingTopic).where(MeetingTopic.meeting_id == item.id).order_by(MeetingTopic.created_at)).all()
    done = sum(1 for step in steps if step.done)
    return {
        "id": item.id,
        "meeting_type": item.meeting_type,
        "meeting_type_label": MEETING_TYPES.get(item.meeting_type, item.meeting_type),
        "organization": item.organization,
        "title": item.title,
        "scheduled_at": item.scheduled_at,
        "completed_at": item.completed_at,
        "status": item.status,
        "archived_at": item.archived_at,
        "archive_reason": item.archive_reason,
        "workflow_template_id": item.workflow_template_id,
        "host_id": item.host_id,
        "recorder_id": item.recorder_id,
        "venue": item.venue,
        "study_plan_id": item.study_plan_id,
        "business_data": item.business_data,
        "task_id": item.task_id,
        "recurrence_key": item.recurrence_key,
        "version": item.version,
        "progress": {"done": done, "total": len(steps), "percent": round(done * 100 / len(steps)) if steps else 0},
        "steps": [
            {"id": step.id, "title": step.title, "assignee_id": step.assignee_id, "due_at": step.due_at, "done": step.done, "version": step.version}
            for step in steps
        ],
        "topics": [_topic_out(topic) for topic in topics if not topic.archived_at],
        "archived_topics": [_topic_out(topic) for topic in topics if topic.archived_at],
    }


def _topic_out(topic: MeetingTopic) -> dict[str, Any]:
    return {
        "id": topic.id,
        "title": topic.title,
        "review_result": topic.review_result,
        "amount": f"{Decimal(topic.amount_cents) / 100:.2f}",
        "reviewed": topic.reviewed,
        "amount_confirmed": topic.amount_confirmed,
        "archived_at": topic.archived_at,
        "archive_reason": topic.archive_reason,
        "version": topic.version,
    }


def _amount_cents(value: str) -> int:
    try:
        amount = Decimal(value)
        if not amount.is_finite() or amount != amount.quantize(Decimal("0.01")):
            raise InvalidOperation
        amount_cents = int(amount * 100)
    except (InvalidOperation, ValueError):
        raise ProblemException(422, "MEETING_AMOUNT_INVALID", "议题金额无效", "请输入最多两位小数的非负金额。") from None
    if amount_cents < 0 or amount_cents > 9_000_000_000_000_000:
        raise ProblemException(422, "MEETING_AMOUNT_INVALID", "议题金额无效", "金额必须为非负数且不能超过系统统计上限。")
    return amount_cents


def _can_modify_meeting(db: Session, meeting: BusinessMeeting, user: User) -> bool:
    """旧通用接口与专属模块共用同一权限边界，防止旁路修改。"""

    if user.role == UserRole.ADMIN or user.id in {
        meeting.created_by,
        meeting.host_id,
        meeting.recorder_id,
    }:
        return True
    plan = db.get(StudyPlan, meeting.study_plan_id) if meeting.study_plan_id else None
    return bool(plan and plan.secretary_id == user.id)


def _require_modify_meeting(db: Session, meeting: BusinessMeeting, user: User) -> None:
    if not _can_modify_meeting(db, meeting, user):
        raise ProblemException(
            403,
            "MEETING_MODIFY_FORBIDDEN",
            "没有修改此台账的权限",
            "仅管理员、创建人、主持人、记录人或指定学习秘书可以修改。",
        )


def _require_export_meeting(db: Session, meeting: BusinessMeeting, user: User) -> None:
    if user.role == UserRole.ADMIN or user.id == meeting.recorder_id:
        return
    plan = db.get(StudyPlan, meeting.study_plan_id) if meeting.study_plan_id else None
    if plan and plan.secretary_id == user.id:
        return
    raise ProblemException(
        403,
        "MEETING_EXPORT_FORBIDDEN",
        "没有导出此材料的权限",
        "仅管理员、记录负责人或指定学习秘书可以导出。",
    )


def _create_meeting_record(
    db: Session,
    *,
    template: WorkflowTemplate,
    meeting_type: str,
    organization: str,
    title: str,
    scheduled_at: datetime | None,
    owner: User,
    created_by: User,
    assignees: dict[str, str] | None = None,
    recurrence_key: str = "",
) -> BusinessMeeting:
    """原子创建会议、任务和步骤，供手工入口与周期调度共享。"""

    if recurrence_key and db.scalar(
        select(BusinessMeeting.id).where(
            BusinessMeeting.recurrence_key == recurrence_key
        )
    ):
        raise ProblemException(
            409,
            "MEETING_RECURRENCE_EXISTS",
            "本周期会议已生成",
            "请打开已有会议，不要重复生成。",
        )
    task = Task(
        title=title,
        description=f"{MEETING_TYPES[meeting_type]}筹备流程",
        task_type=TaskType.PROJECT,
        status=TaskStatus.IN_PROGRESS,
        sensitivity=Sensitivity.NORMAL,
        priority=Priority.NORMAL,
        source="会议工作流",
        source_kind="workflow",
        category="党建会议",
        formal_due_at=scheduled_at,
        internal_due_at=scheduled_at,
        owner_id=owner.id,
        created_by=created_by.id,
        updated_by=created_by.id,
    )
    db.add(task)
    db.flush()
    assignments = assignees or {}
    assignment_ids = {str(value) for value in assignments.values() if value}
    if assignment_ids:
        active_ids = set(
            db.scalars(
                select(User.id).where(User.id.in_(assignment_ids), User.active.is_(True))
            ).all()
        )
        if active_ids != assignment_ids:
            raise ProblemException(
                422,
                "MEETING_ASSIGNEE_INVALID",
                "步骤负责人不可用",
                "请为每个角色选择仍在启用状态的用户。",
            )
    for order, definition in enumerate(template.steps):
        role = str(definition.get("responsible_role", ""))
        due_at = (
            scheduled_at + timedelta(days=int(definition.get("offset_days", 0)))
            if scheduled_at
            else None
        )
        db.add(
            TaskStep(
                task_id=task.id,
                title=str(definition.get("title", f"步骤{order + 1}"))[:240],
                assignee_id=assignments.get(role) or owner.id,
                due_at=due_at,
                sort_order=order,
            )
        )
    meeting = BusinessMeeting(
        meeting_type=meeting_type,
        organization=organization.strip(),
        title=title.strip(),
        scheduled_at=scheduled_at,
        workflow_template_id=template.id,
        task_id=task.id,
        recurrence_key=recurrence_key,
        created_by=created_by.id,
    )
    db.add(meeting)
    db.flush()
    return meeting


def generate_due_recurring_meetings(
    db: Session,
    admin: User,
    now: datetime,
) -> list[BusinessMeeting]:
    """幂等生成本月及预生成月份会议；无完整配置的模板保持停用。"""

    local_now = now if now.tzinfo else now.replace(tzinfo=timezone.utc)
    templates = db.scalars(
        select(WorkflowTemplate).where(WorkflowTemplate.active.is_(True))
    ).all()
    generated: list[BusinessMeeting] = []
    for template in templates:
        recurrence = template.recurrence if isinstance(template.recurrence, dict) else {}
        if recurrence.get("kind") != "monthly" or recurrence.get("enabled") is not True:
            continue
        organization = str(recurrence.get("organization", "")).strip()
        if not organization:
            continue
        day = max(1, min(28, int(recurrence.get("day", 1) or 1)))
        hour = max(0, min(23, int(recurrence.get("hour", 9) or 9)))
        minute = max(0, min(59, int(recurrence.get("minute", 0) or 0)))
        months_ahead = max(0, min(12, int(recurrence.get("months_ahead", 1) or 1)))
        owner = db.get(User, str(recurrence.get("owner_id", ""))) or admin
        if not owner.active:
            owner = admin
        for offset in range(months_ahead + 1):
            absolute_month = local_now.year * 12 + local_now.month - 1 + offset
            year, month_index = divmod(absolute_month, 12)
            month = month_index + 1
            recurrence_key = f"{template.id}:{year:04d}-{month:02d}"
            if db.scalar(
                select(BusinessMeeting.id).where(
                    BusinessMeeting.recurrence_key == recurrence_key
                )
            ):
                continue
            scheduled = datetime(year, month, day, hour, minute, tzinfo=local_now.tzinfo)
            pattern = str(recurrence.get("title_pattern", "{year}年{month}月{type}"))
            title = pattern.format(
                year=year,
                month=month,
                type=MEETING_TYPES.get(template.business_type, template.business_type),
                organization=organization,
            )[:240]
            generated.append(
                _create_meeting_record(
                    db,
                    template=template,
                    meeting_type=template.business_type,
                    organization=organization,
                    title=title,
                    scheduled_at=scheduled,
                    owner=owner,
                    created_by=admin,
                    recurrence_key=recurrence_key,
                )
            )
    return generated


@router.get("/workflow-templates", response_model=List[dict])
def list_workflow_templates(
    include_archived: bool = False,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> list[dict[str, Any]]:
    ensure_committee_template(db, user)
    db.commit()
    query = select(WorkflowTemplate)
    if not include_archived:
        query = query.where(WorkflowTemplate.active.is_(True))
    return [template_out(item) for item in db.scalars(query.order_by(WorkflowTemplate.name)).all()]


@router.post("/workflow-templates", response_model=dict, status_code=201)
def create_workflow_template(
    payload: WorkflowTemplateInput,
    request: Request,
    admin: User = Depends(require_admin),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    if not payload.steps:
        raise ProblemException(422, "WORKFLOW_STEPS_REQUIRED", "流程步骤不能为空", "请至少配置一个有完成条件的步骤。")
    item = WorkflowTemplate(**payload.model_dump(), created_by=admin.id)
    db.add(item)
    db.flush()
    write_audit(db, admin, "workflow_template.create", "workflow_template", item.id, {"step_count": len(item.steps)}, client_ip(request))
    db.commit()
    return template_out(item)


@router.delete("/workflow-templates/{template_id}", response_model=dict)
def archive_workflow_template(
    template_id: str,
    request: Request,
    payload: LifecycleReason | None = Body(default=None),
    admin: User = Depends(require_admin),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    item = db.get(WorkflowTemplate, template_id)
    if not item:
        raise ProblemException(404, "WORKFLOW_TEMPLATE_NOT_FOUND", "流程模板不存在", "请刷新后重试。")
    if item.built_in:
        raise ProblemException(409, "BUILT_IN_TEMPLATE_ARCHIVE_FORBIDDEN", "内置模板不能归档", "内置流程用于系统业务，可复制后维护自定义版本。")
    item.active = False
    item.archived_at = utcnow()
    item.archived_by = admin.id
    item.archive_reason = payload.reason.strip() if payload else "管理员归档自定义流程模板"
    item.version += 1
    write_audit(db, admin, "workflow_template.archive", "workflow_template", item.id, {"reason": item.archive_reason}, client_ip(request))
    db.commit()
    return {"archived": True, "id": item.id}


@router.post("/workflow-templates/{template_id}/restore", response_model=dict)
def restore_workflow_template(
    template_id: str,
    payload: LifecycleReason,
    request: Request,
    admin: User = Depends(require_admin),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    item = db.get(WorkflowTemplate, template_id)
    if not item:
        raise ProblemException(404, "WORKFLOW_TEMPLATE_NOT_FOUND", "流程模板不存在", "请刷新后重试。")
    item.active = True
    item.archived_at = None
    item.archived_by = None
    item.archive_reason = ""
    item.version += 1
    write_audit(db, admin, "workflow_template.restore", "workflow_template", item.id, {"reason": payload.reason.strip()}, client_ip(request))
    db.commit()
    return template_out(item)


def _meeting_import_draft(db: Session, draft_id: str, user: User) -> MeetingImportDraft:
    draft = db.get(MeetingImportDraft, draft_id)
    if not draft or (draft.created_by != user.id and user.role != UserRole.ADMIN):
        raise ProblemException(
            404,
            "MEETING_IMPORT_DRAFT_NOT_FOUND",
            "会议导入草稿不存在",
            "草稿可能已过期、被取消或不属于当前用户。",
        )
    expires_at = draft.expires_at
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if draft.status == "draft" and expires_at <= utcnow():
        draft.status = "expired"
        draft.version += 1
        db.commit()
        raise ProblemException(
            410,
            "MEETING_IMPORT_DRAFT_EXPIRED",
            "会议导入草稿已过期",
            "原始文件未被保存，请重新选择文件生成候选。",
        )
    return draft


def _meeting_import_out(draft: MeetingImportDraft) -> dict[str, Any]:
    return {
        "id": draft.id,
        "source_sha256": draft.source_sha256,
        "source_kind": draft.source_kind,
        "status": draft.status,
        "meeting": draft.proposed_meeting,
        "topics": draft.proposed_topics,
        "warnings": draft.warnings,
        "meeting_id": draft.meeting_id,
        "expires_at": draft.expires_at,
        "confirmed_at": draft.confirmed_at,
        "version": draft.version,
    }


def _utc_business_datetime(value: datetime | None) -> datetime | None:
    """无时区输入一律按北京时间解释，数据库只保存 UTC 瞬时。"""

    if value is None:
        return None
    if value.tzinfo is None:
        value = value.replace(tzinfo=timezone(timedelta(hours=8)))
    return value.astimezone(timezone.utc)


@router.post("/business-meetings/imports/inspect", response_model=dict, status_code=201)
async def inspect_meeting_import(
    request: Request,
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    """检查真实容器并生成候选；原文件、文件名和正文均不落库。"""

    data = await file.read(MAX_MEETING_IMPORT_BYTES + 1)
    filename_hint = file.filename or ""
    await file.close()
    if len(data) > MAX_MEETING_IMPORT_BYTES:
        raise ProblemException(
            413,
            "MEETING_IMPORT_SIZE_LIMIT",
            "会议文件超过限制",
            "单个会议导入文件不得超过 50 MiB。",
        )
    text, source_kind = await run_in_threadpool(extract_meeting_text, data)
    candidate = propose_meeting(text, filename_hint)
    topics = list(candidate.pop("topics", []))
    warnings = list(candidate.pop("warnings", []))
    draft = MeetingImportDraft(
        created_by=user.id,
        source_sha256=hashlib.sha256(data).hexdigest(),
        source_kind=source_kind,
        proposed_meeting=candidate,
        proposed_topics=topics,
        warnings=warnings,
        expires_at=utcnow() + timedelta(minutes=30),
    )
    db.add(draft)
    db.flush()
    write_audit(
        db,
        user,
        "business_meeting.import_inspect",
        "meeting_import_draft",
        draft.id,
        {
            "source_kind": source_kind,
            "source_sha256_prefix": draft.source_sha256[:12],
            "topic_candidates": len(topics),
        },
        client_ip(request),
    )
    db.commit()
    return _meeting_import_out(draft)


@router.get("/business-meetings/imports/{draft_id}", response_model=dict)
def get_meeting_import(
    draft_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    return _meeting_import_out(_meeting_import_draft(db, draft_id, user))


@router.patch("/business-meetings/imports/{draft_id}", response_model=dict)
def confirm_meeting_import_candidates(
    draft_id: str,
    payload: MeetingImportPatch,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    draft = _meeting_import_draft(db, draft_id, user)
    if draft.status != "draft":
        raise ProblemException(409, "MEETING_IMPORT_STATE_INVALID", "会议草稿状态不可修改", "请重新生成导入草稿。")
    if draft.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "会议导入草稿已更新", "请刷新后重新确认。")
    if payload.meeting_type not in MEETING_TYPES:
        raise ProblemException(422, "MEETING_TYPE_INVALID", "会议类型无效", "请选择系统支持的党建会议类型。")
    scheduled_at = _utc_business_datetime(payload.scheduled_at)
    draft.proposed_meeting = {
        "meeting_type": payload.meeting_type,
        "organization": payload.organization.strip(),
        "title": payload.title.strip(),
        "scheduled_at": beijing_iso(scheduled_at) if scheduled_at else None,
        "venue": payload.venue.strip(),
        "host_name": payload.host_name.strip(),
        "attendee_text": payload.attendee_text.strip(),
        "owner_id": payload.owner_id,
        "requires_confirmation": False,
    }
    unique_topics: list[str] = []
    for raw_title in payload.topics:
        title = " ".join(raw_title.split()).strip(" ：:;；。")[:240]
        if title and title not in unique_topics:
            unique_topics.append(title)
    draft.proposed_topics = [
        {"title": title, "confidence": 1.0, "source": "用户人工确认", "confirmed": True}
        for title in unique_topics
    ]
    draft.version += 1
    write_audit(
        db,
        user,
        "business_meeting.import_confirm",
        "meeting_import_draft",
        draft.id,
        {"topic_count": len(unique_topics), "meeting_type": payload.meeting_type},
        client_ip(request),
    )
    db.commit()
    return _meeting_import_out(draft)


@router.post("/business-meetings/imports/{draft_id}/commit", response_model=dict, status_code=201)
def commit_meeting_import(
    draft_id: str,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    draft = _meeting_import_draft(db, draft_id, user)
    if draft.status != "draft":
        if draft.status == "committed" and draft.meeting_id:
            meeting = db.get(BusinessMeeting, draft.meeting_id)
            if meeting:
                return {"draft": _meeting_import_out(draft), "meeting": meeting_out(db, meeting)}
        raise ProblemException(409, "MEETING_IMPORT_STATE_INVALID", "会议草稿不能提交", "请重新生成并确认导入草稿。")
    if draft.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "会议导入草稿已更新", "请刷新后重新确认。")
    values = dict(draft.proposed_meeting or {})
    if values.get("requires_confirmation", True):
        raise ProblemException(
            409,
            "MEETING_IMPORT_CONFIRMATION_REQUIRED",
            "会议候选尚未人工确认",
            "请先核对会议、时间、地点和议题，再提交正式台账。",
        )
    meeting_type = str(values.get("meeting_type", ""))
    if meeting_type not in MEETING_TYPES:
        raise ProblemException(422, "MEETING_TYPE_INVALID", "会议类型无效", "请选择系统支持的党建会议类型。")
    owner = db.get(User, values.get("owner_id")) if values.get("owner_id") else user
    if not owner or not owner.active:
        raise ProblemException(422, "MEETING_OWNER_INVALID", "负责人不可用", "请选择启用用户。")
    scheduled_raw = values.get("scheduled_at")
    scheduled_at = datetime.fromisoformat(scheduled_raw) if scheduled_raw else None
    template = ensure_committee_template(db, user)
    meeting = _create_meeting_record(
        db,
        template=template,
        meeting_type=meeting_type,
        organization=str(values.get("organization", "")),
        title=str(values.get("title", "")),
        scheduled_at=scheduled_at,
        owner=owner,
        created_by=user,
    )
    meeting.venue = str(values.get("venue", ""))
    meeting.business_data = {
        "import_source_sha256": draft.source_sha256,
        "import_source_kind": draft.source_kind,
        "host_name_candidate": str(values.get("host_name", "")),
        "attendee_text_candidate": str(values.get("attendee_text", "")),
    }
    for candidate in draft.proposed_topics or []:
        if not candidate.get("confirmed") or not str(candidate.get("title", "")).strip():
            continue
        db.add(MeetingTopic(meeting_id=meeting.id, title=str(candidate["title"]).strip()[:240]))
    draft.status = "committed"
    draft.meeting_id = meeting.id
    draft.confirmed_at = utcnow()
    draft.version += 1
    write_audit(
        db,
        user,
        "business_meeting.import_commit",
        "business_meeting",
        meeting.id,
        {"draft_id": draft.id, "topic_count": len(draft.proposed_topics or [])},
        client_ip(request),
    )
    db.commit()
    return {"draft": _meeting_import_out(draft), "meeting": meeting_out(db, meeting)}


@router.delete("/business-meetings/imports/{draft_id}", response_model=dict)
def cancel_meeting_import(
    draft_id: str,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    draft = _meeting_import_draft(db, draft_id, user)
    if draft.status == "committed":
        raise ProblemException(409, "MEETING_IMPORT_ALREADY_COMMITTED", "会议草稿已经入账", "如需撤销，请归档对应会议并保留审计记录。")
    if draft.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "会议导入草稿已更新", "请刷新后重试。")
    draft.status = "cancelled"
    draft.proposed_meeting = {}
    draft.proposed_topics = []
    draft.warnings = []
    draft.version += 1
    write_audit(db, user, "business_meeting.import_cancel", "meeting_import_draft", draft.id, {}, client_ip(request))
    db.commit()
    return {"cancelled": True, "id": draft.id, "version": draft.version}


@router.get("/business-meetings", response_model=List[dict])
def list_meetings(
    year: int | None = Query(default=None, ge=2000, le=2200),
    organization: str = "",
    meeting_type: str = "",
    scope: str = Query(default="", pattern=r"^(|other)$"),
    lifecycle: str = Query(default="active", pattern=r"^(active|archived|all)$"),
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> list[dict[str, Any]]:
    query = select(BusinessMeeting)
    if lifecycle == "active":
        query = query.where(BusinessMeeting.archived_at.is_(None))
    elif lifecycle == "archived":
        query = query.where(BusinessMeeting.archived_at.is_not(None))
    if scope == "other":
        query = query.where(
            BusinessMeeting.meeting_type.not_in(
                {"branch_members", "party_member_meeting", "party_group", "party_class", "study_group"}
            )
        )
    if organization:
        query = query.where(BusinessMeeting.organization == organization)
    if meeting_type:
        query = query.where(BusinessMeeting.meeting_type == meeting_type)
    if year:
        query = query.where(BusinessMeeting.scheduled_at >= datetime(year, 1, 1, tzinfo=timezone.utc), BusinessMeeting.scheduled_at < datetime(year + 1, 1, 1, tzinfo=timezone.utc))
    return [meeting_out(db, item) for item in db.scalars(query.order_by(BusinessMeeting.scheduled_at.desc())).all()]


@router.get("/business-meetings/{meeting_id}/deletion-impact", response_model=dict)
def meeting_deletion_impact(
    meeting_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    item = db.get(BusinessMeeting, meeting_id)
    if not item:
        raise ProblemException(404, "MEETING_NOT_FOUND", "会议不存在", "请刷新后重试。")
    _require_modify_meeting(db, item, user)
    steps = db.scalar(select(func.count(TaskStep.id)).where(TaskStep.task_id == item.task_id)) or 0
    topics = db.scalar(select(func.count(MeetingTopic.id)).where(MeetingTopic.meeting_id == item.id)) or 0
    documents = db.scalar(select(func.count(BusinessDocument.id)).where(BusinessDocument.meeting_id == item.id)) or 0
    attendees = db.scalar(select(func.count(MeetingAttendee.id)).where(MeetingAttendee.meeting_id == item.id)) or 0
    actions = db.scalar(select(func.count(MeetingAction.id)).where(MeetingAction.meeting_id == item.id)) or 0
    return {"steps": steps, "topics": topics, "documents": documents, "attendees": attendees, "actions": actions, "recoverable": True, "physical_delete": False}


@router.delete("/business-meetings/{meeting_id}", response_model=dict)
def archive_meeting(
    meeting_id: str,
    payload: LifecycleReason,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    item = db.get(BusinessMeeting, meeting_id)
    if not item:
        raise ProblemException(404, "MEETING_NOT_FOUND", "会议不存在", "请刷新后重试。")
    _require_modify_meeting(db, item, user)
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "会议已更新", "请刷新后重试。")
    if item.archived_at:
        return meeting_out(db, item)
    item.status_before_archive = item.status
    item.status = "archived"
    item.archived_at = utcnow()
    item.archived_by = user.id
    item.archive_reason = payload.reason.strip()
    item.version += 1
    if item.task_id and (task := db.get(Task, item.task_id)):
        task.status = TaskStatus.ARCHIVED
        task.version += 1
    write_audit(db, user, "business_meeting.archive", "business_meeting", item.id, {"reason": item.archive_reason}, client_ip(request))
    db.commit()
    return meeting_out(db, item)


@router.post("/business-meetings/{meeting_id}/restore", response_model=dict)
def restore_meeting(
    meeting_id: str,
    payload: LifecycleReason,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    item = db.get(BusinessMeeting, meeting_id)
    if not item:
        raise ProblemException(404, "MEETING_NOT_FOUND", "会议不存在", "请刷新后重试。")
    _require_modify_meeting(db, item, user)
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "会议已更新", "请刷新后重试。")
    if not item.archived_at:
        return meeting_out(db, item)
    item.status = item.status_before_archive if item.status_before_archive != "archived" else "planned"
    item.archived_at = None
    item.archived_by = None
    item.archive_reason = ""
    item.version += 1
    if item.task_id and (task := db.get(Task, item.task_id)):
        task.status = TaskStatus.COMPLETED if item.status == "completed" else TaskStatus.IN_PROGRESS
        task.version += 1
    write_audit(db, user, "business_meeting.restore", "business_meeting", item.id, {"reason": payload.reason.strip()}, client_ip(request))
    db.commit()
    return meeting_out(db, item)


@router.post("/business-meetings", response_model=dict, status_code=201)
def create_meeting(
    payload: MeetingInput,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    if payload.meeting_type not in MEETING_TYPES:
        raise ProblemException(422, "MEETING_TYPE_INVALID", "会议类型无效", "请选择系统支持的党建会议类型。")
    owner = db.get(User, payload.owner_id) if payload.owner_id else user
    if not owner or not owner.active:
        raise ProblemException(422, "MEETING_OWNER_INVALID", "负责人不可用", "请选择启用用户。")
    template = db.get(WorkflowTemplate, payload.workflow_template_id) if payload.workflow_template_id else ensure_committee_template(db, user)
    if not template or not template.active:
        raise ProblemException(
            422,
            "WORKFLOW_TEMPLATE_INVALID",
            "流程模板不可用",
            "请选择仍在启用状态的流程模板。",
        )
    meeting = _create_meeting_record(
        db,
        template=template,
        meeting_type=payload.meeting_type,
        organization=payload.organization,
        title=payload.title,
        scheduled_at=payload.scheduled_at,
        owner=owner,
        created_by=user,
        assignees=payload.assignees,
        recurrence_key=payload.recurrence_key,
    )
    write_audit(db, user, "business_meeting.create", "business_meeting", meeting.id, {"meeting_type": meeting.meeting_type, "task_id": meeting.task_id}, client_ip(request))
    db.commit()
    return meeting_out(db, meeting)


@router.post("/business-meetings/recurring/generate", response_model=dict)
def generate_recurring_meetings_now(
    request: Request,
    admin: User = Depends(require_admin),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    generated = generate_due_recurring_meetings(db, admin, utcnow())
    write_audit(
        db,
        admin,
        "business_meeting.recurring_generate",
        "workflow_template",
        None,
        {"count": len(generated)},
        client_ip(request),
    )
    db.commit()
    return {"generated": len(generated), "meeting_ids": [item.id for item in generated]}


@router.patch("/business-meetings/{meeting_id}", response_model=dict)
def patch_meeting(
    meeting_id: str,
    payload: MeetingPatch,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    item = db.get(BusinessMeeting, meeting_id)
    if not item:
        raise ProblemException(404, "MEETING_NOT_FOUND", "会议不存在", "请刷新后重试。")
    _require_modify_meeting(db, item, user)
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "会议信息已更新", "请刷新后重试。")
    changes = payload.model_dump(exclude_unset=True)
    old_schedule = item.scheduled_at
    for key, value in changes.items():
        setattr(item, key, value)
    if "scheduled_at" in changes and item.task_id:
        task = db.get(Task, item.task_id)
        if task:
            task.formal_due_at = item.scheduled_at
            task.internal_due_at = item.scheduled_at
            task.version += 1
            steps = db.scalars(
                select(TaskStep).where(TaskStep.task_id == task.id).order_by(TaskStep.sort_order)
            ).all()
            template = db.get(WorkflowTemplate, item.workflow_template_id) if item.workflow_template_id else None
            definitions = template.steps if template and isinstance(template.steps, list) else []
            for index, step in enumerate(steps):
                if index < len(definitions) and isinstance(definitions[index], dict):
                    step.due_at = (
                        item.scheduled_at + timedelta(days=int(definitions[index].get("offset_days", 0)))
                        if item.scheduled_at
                        else None
                    )
                elif item.scheduled_at and old_schedule and step.due_at:
                    # 历史自定义流程没有模板定义时，按绝对时差平移；SQLite 读取的
                    # 时间可能丢失 tzinfo，统一视作 UTC 后再相减。
                    new_aware = item.scheduled_at if item.scheduled_at.tzinfo else item.scheduled_at.replace(tzinfo=timezone.utc)
                    old_aware = old_schedule if old_schedule.tzinfo else old_schedule.replace(tzinfo=timezone.utc)
                    step.due_at += new_aware.astimezone(timezone.utc) - old_aware.astimezone(timezone.utc)
                elif item.scheduled_at is None:
                    step.due_at = None
                step.version += 1
    if item.status == "completed" and item.completed_at is None:
        item.completed_at = utcnow()
    item.version += 1
    write_audit(db, user, "business_meeting.update", "business_meeting", item.id, {"fields": sorted(changes)}, client_ip(request))
    db.commit()
    return meeting_out(db, item)


@router.post("/business-meetings/{meeting_id}/topics", response_model=dict, status_code=201)
def create_topic(
    meeting_id: str,
    payload: TopicInput,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting = db.get(BusinessMeeting, meeting_id)
    if not meeting:
        raise ProblemException(404, "MEETING_NOT_FOUND", "会议不存在", "请刷新后重试。")
    _require_modify_meeting(db, meeting, user)
    if meeting.archived_at:
        raise ProblemException(409, "MEETING_ARCHIVED", "会议已归档", "请先恢复会议，再新增议题。")
    amount_cents = _amount_cents(payload.amount)
    item = MeetingTopic(meeting_id=meeting_id, title=payload.title.strip(), review_result=payload.review_result, amount_cents=amount_cents, reviewed=payload.reviewed, amount_confirmed=payload.amount_confirmed)
    db.add(item)
    db.flush()
    write_audit(db, user, "meeting_topic.create", "meeting_topic", item.id, {"meeting_id": meeting_id}, client_ip(request))
    db.commit()
    return {**_topic_out(item), "meeting_id": meeting_id}


def _meeting_topic(
    db: Session,
    meeting_id: str,
    topic_id: str,
    user: User,
) -> tuple[BusinessMeeting, MeetingTopic]:
    meeting = db.get(BusinessMeeting, meeting_id)
    topic = db.get(MeetingTopic, topic_id)
    if not meeting:
        raise ProblemException(404, "MEETING_NOT_FOUND", "会议不存在", "请刷新后重试。")
    if not topic or topic.meeting_id != meeting.id:
        raise ProblemException(404, "MEETING_TOPIC_NOT_FOUND", "会议议题不存在", "请刷新后重试。")
    _require_modify_meeting(db, meeting, user)
    return meeting, topic


@router.get("/business-meetings/{meeting_id}/topics/{topic_id}/deletion-impact", response_model=dict)
def meeting_topic_deletion_impact(
    meeting_id: str,
    topic_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    _meeting, topic = _meeting_topic(db, meeting_id, topic_id, user)
    return {
        "annual_statistics": bool(topic.reviewed or topic.amount_confirmed),
        "confirmed_amount": bool(topic.amount_confirmed and topic.amount_cents),
        "recoverable": True,
        "physical_delete": False,
    }


@router.patch("/business-meetings/{meeting_id}/topics/{topic_id}", response_model=dict)
def patch_meeting_topic(
    meeting_id: str,
    topic_id: str,
    payload: TopicPatch,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting, topic = _meeting_topic(db, meeting_id, topic_id, user)
    if meeting.archived_at or topic.archived_at:
        raise ProblemException(409, "MEETING_TOPIC_ARCHIVED", "议题已归档", "请先恢复会议和议题，再继续修改。")
    if topic.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "议题已更新", "请刷新后重试。")
    changes = payload.model_dump(exclude_unset=True)
    if "amount" in changes:
        topic.amount_cents = _amount_cents(str(changes.pop("amount")))
    for key, value in changes.items():
        setattr(topic, key, value.strip() if isinstance(value, str) else value)
    topic.version += 1
    write_audit(db, user, "meeting_topic.update", "meeting_topic", topic.id, {"fields": sorted(payload.model_fields_set)}, client_ip(request))
    db.commit()
    return _topic_out(topic)


@router.delete("/business-meetings/{meeting_id}/topics/{topic_id}", response_model=dict)
def archive_meeting_topic(
    meeting_id: str,
    topic_id: str,
    payload: LifecycleReason,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting, topic = _meeting_topic(db, meeting_id, topic_id, user)
    if topic.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "议题已更新", "请刷新后重试。")
    if not topic.archived_at:
        topic.archived_at = utcnow()
        topic.archived_by = user.id
        topic.archive_reason = payload.reason.strip()
        topic.version += 1
        write_audit(db, user, "meeting_topic.archive", "meeting_topic", topic.id, {"reason": topic.archive_reason}, client_ip(request))
        db.commit()
    return _topic_out(topic)


@router.post("/business-meetings/{meeting_id}/topics/{topic_id}/restore", response_model=dict)
def restore_meeting_topic(
    meeting_id: str,
    topic_id: str,
    payload: LifecycleReason,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting, topic = _meeting_topic(db, meeting_id, topic_id, user)
    if meeting.archived_at:
        raise ProblemException(409, "MEETING_ARCHIVED", "会议仍在归档中", "请先恢复会议，再恢复议题。")
    if topic.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "议题已更新", "请刷新后重试。")
    if topic.archived_at:
        topic.archived_at = None
        topic.archived_by = None
        topic.archive_reason = ""
        topic.version += 1
        write_audit(db, user, "meeting_topic.restore", "meeting_topic", topic.id, {"reason": payload.reason.strip()}, client_ip(request))
        db.commit()
    return _topic_out(topic)


@router.get("/business-meetings/statistics/annual", response_model=dict)
def annual_meeting_statistics(
    year: int = Query(ge=2000, le=2200),
    organization: str = "",
    meeting_type: str = "",
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    query = select(BusinessMeeting).where(BusinessMeeting.status == "completed", BusinessMeeting.completed_at >= datetime(year, 1, 1, tzinfo=timezone.utc), BusinessMeeting.completed_at < datetime(year + 1, 1, 1, tzinfo=timezone.utc))
    if organization:
        query = query.where(BusinessMeeting.organization == organization)
    if meeting_type:
        query = query.where(BusinessMeeting.meeting_type == meeting_type)
    meetings = db.scalars(query).all()
    ids = [item.id for item in meetings]
    topics = db.scalars(select(MeetingTopic).where(MeetingTopic.meeting_id.in_(ids), MeetingTopic.reviewed.is_(True), MeetingTopic.archived_at.is_(None))).all() if ids else []
    confirmed_cents = sum(topic.amount_cents for topic in topics if topic.amount_confirmed)
    return {"year": year, "organization": organization, "meeting_type": meeting_type, "completed_meetings": len(meetings), "reviewed_topics": len(topics), "confirmed_amount": f"{Decimal(confirmed_cents) / 100:.2f}"}


@router.get("/business-documents", response_model=List[dict])
def list_documents(
    meeting_id: str | None = None,
    lifecycle: str = Query(default="active", pattern=r"^(active|archived|all)$"),
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> list[dict[str, Any]]:
    query = select(BusinessDocument)
    if lifecycle == "active":
        query = query.where(BusinessDocument.archived_at.is_(None))
    elif lifecycle == "archived":
        query = query.where(BusinessDocument.archived_at.is_not(None))
    if meeting_id:
        query = query.where(BusinessDocument.meeting_id == meeting_id)
    return [{"id": item.id, "meeting_id": item.meeting_id, "task_step_id": item.task_step_id, "document_type": item.document_type, "title": item.title, "content": item.content, "version": item.version, "archived_at": item.archived_at, "archive_reason": item.archive_reason, "updated_at": item.updated_at} for item in db.scalars(query.order_by(BusinessDocument.updated_at.desc())).all()]


@router.get("/business-documents/{document_id}/deletion-impact", response_model=dict)
def document_deletion_impact(
    document_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    item = db.get(BusinessDocument, document_id)
    if not item:
        raise ProblemException(404, "BUSINESS_DOCUMENT_NOT_FOUND", "文档不存在", "请刷新后重试。")
    meeting = db.get(BusinessMeeting, item.meeting_id) if item.meeting_id else None
    if meeting:
        _require_modify_meeting(db, meeting, user)
    elif item.created_by != user.id and user.role != UserRole.ADMIN:
        raise ProblemException(403, "DOCUMENT_ARCHIVE_FORBIDDEN", "没有归档此文档的权限", "仅创建人或管理员可以归档。")
    revisions = db.scalar(select(func.count(BusinessDocumentRevision.id)).where(BusinessDocumentRevision.document_id == item.id)) or 0
    return {"revisions": revisions, "meeting_linked": bool(item.meeting_id), "task_step_linked": bool(item.task_step_id), "recoverable": True, "physical_delete": False}


@router.delete("/business-documents/{document_id}", response_model=dict)
def archive_document(
    document_id: str,
    payload: LifecycleReason,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    item = db.get(BusinessDocument, document_id)
    if not item:
        raise ProblemException(404, "BUSINESS_DOCUMENT_NOT_FOUND", "文档不存在", "请刷新后重试。")
    meeting = db.get(BusinessMeeting, item.meeting_id) if item.meeting_id else None
    if meeting:
        _require_modify_meeting(db, meeting, user)
    elif item.created_by != user.id and user.role != UserRole.ADMIN:
        raise ProblemException(403, "DOCUMENT_ARCHIVE_FORBIDDEN", "没有归档此文档的权限", "仅创建人或管理员可以归档。")
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "DOCUMENT_VERSION_CONFLICT", "文档已被修改", "请刷新后重试。")
    item.archived_at = utcnow()
    item.archived_by = user.id
    item.archive_reason = payload.reason.strip()
    item.version += 1
    write_audit(db, user, "business_document.archive", "business_document", item.id, {"reason": item.archive_reason}, client_ip(request))
    db.commit()
    return {"id": item.id, "archived": True, "version": item.version}


@router.post("/business-documents/{document_id}/restore", response_model=dict)
def restore_document(
    document_id: str,
    payload: LifecycleReason,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    item = db.get(BusinessDocument, document_id)
    if not item:
        raise ProblemException(404, "BUSINESS_DOCUMENT_NOT_FOUND", "文档不存在", "请刷新后重试。")
    meeting = db.get(BusinessMeeting, item.meeting_id) if item.meeting_id else None
    if meeting:
        _require_modify_meeting(db, meeting, user)
    elif item.created_by != user.id and user.role != UserRole.ADMIN:
        raise ProblemException(403, "DOCUMENT_RESTORE_FORBIDDEN", "没有恢复此文档的权限", "仅创建人或管理员可以恢复。")
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "DOCUMENT_VERSION_CONFLICT", "文档已被修改", "请刷新后重试。")
    item.archived_at = None
    item.archived_by = None
    item.archive_reason = ""
    item.version += 1
    write_audit(db, user, "business_document.restore", "business_document", item.id, {"reason": payload.reason.strip()}, client_ip(request))
    db.commit()
    return {"id": item.id, "archived": False, "version": item.version}


@router.post("/business-documents", response_model=dict, status_code=201)
def create_document(
    payload: DocumentInput,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    meeting = db.get(BusinessMeeting, payload.meeting_id) if payload.meeting_id else None
    step = db.get(TaskStep, payload.task_step_id) if payload.task_step_id else None
    if payload.meeting_id and not meeting:
        raise ProblemException(422, "DOCUMENT_MEETING_INVALID", "关联会议不存在", "请刷新会议列表后重试。")
    if meeting:
        _require_modify_meeting(db, meeting, user)
    if payload.task_step_id and not step:
        raise ProblemException(422, "DOCUMENT_STEP_INVALID", "关联步骤不存在", "请刷新筹备步骤后重试。")
    if meeting and step and meeting.task_id != step.task_id:
        raise ProblemException(422, "DOCUMENT_LINK_MISMATCH", "会议与步骤不匹配", "请选择该会议筹备流程内的步骤。")
    item = BusinessDocument(**payload.model_dump(), created_by=user.id)
    db.add(item)
    db.flush()
    db.add(BusinessDocumentRevision(document_id=item.id, revision_no=1, content=item.content, change_note="创建文档", created_by=user.id))
    write_audit(db, user, "business_document.create", "business_document", item.id, {"document_type": item.document_type}, client_ip(request))
    db.commit()
    return {"id": item.id, "version": item.version, "updated_at": item.updated_at}


@router.patch("/business-documents/{document_id}", response_model=dict)
def update_document(
    document_id: str,
    payload: DocumentPatch,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, Any]:
    item = db.get(BusinessDocument, document_id)
    if not item or item.archived_at:
        raise ProblemException(404, "BUSINESS_DOCUMENT_NOT_FOUND", "文档不存在", "请刷新后重试。")
    meeting = db.get(BusinessMeeting, item.meeting_id) if item.meeting_id else None
    if meeting:
        _require_modify_meeting(db, meeting, user)
    elif item.created_by != user.id and user.role != UserRole.ADMIN:
        raise ProblemException(
            403,
            "DOCUMENT_MODIFY_FORBIDDEN",
            "没有修改此文档的权限",
            "仅创建人或管理员可以修改未关联会议的独立文档。",
        )
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "DOCUMENT_VERSION_CONFLICT", "文档已被其他人修改", "请刷新并合并修改后重试。", extra={"current_version": item.version})
    if payload.title is not None:
        item.title = payload.title.strip()
    if payload.content is not None:
        item.content = payload.content
    item.version += 1
    db.add(BusinessDocumentRevision(document_id=item.id, revision_no=item.version, content=item.content, change_note=payload.change_note, created_by=user.id))
    write_audit(db, user, "business_document.update", "business_document", item.id, {"version": item.version}, client_ip(request))
    db.commit()
    return {"id": item.id, "title": item.title, "content": item.content, "version": item.version, "updated_at": item.updated_at}


@router.get("/business-documents/{document_id}/revisions", response_model=List[dict])
def document_revisions(
    document_id: str,
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> list[dict[str, Any]]:
    item = db.get(BusinessDocument, document_id)
    if not item:
        raise ProblemException(404, "BUSINESS_DOCUMENT_NOT_FOUND", "文档不存在", "请刷新后重试。")
    return [{"id": row.id, "revision_no": row.revision_no, "content": row.content, "change_note": row.change_note, "created_by": row.created_by, "created_at": row.created_at} for row in db.scalars(select(BusinessDocumentRevision).where(BusinessDocumentRevision.document_id == document_id).order_by(BusinessDocumentRevision.revision_no.desc())).all()]


def _append_structured_content(document: Document, content: dict[str, Any]) -> None:
    blocks = content.get("blocks", []) if isinstance(content, dict) else []
    for block in blocks if isinstance(blocks, list) else []:
        if not isinstance(block, dict):
            continue
        kind = block.get("type", "paragraph")
        text = str(block.get("text", ""))
        if kind == "heading":
            document.add_heading(text, level=max(1, min(3, int(block.get("level", 1)))))
        elif kind == "list":
            document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(text)


@router.get("/business-documents/{document_id}/export.docx")
def export_business_document(
    document_id: str,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> StreamingResponse:
    item = db.get(BusinessDocument, document_id)
    if not item or item.archived_at:
        raise ProblemException(404, "BUSINESS_DOCUMENT_NOT_FOUND", "文档不存在", "请刷新后重试。")
    meeting = db.get(BusinessMeeting, item.meeting_id) if item.meeting_id else None
    if meeting:
        _require_export_meeting(db, meeting, user)
    doc = Document()
    doc.add_heading(item.title, 0)
    _append_structured_content(doc, item.content)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = "".join(char if char not in '\\/:*?\"<>|' else "_" for char in item.title)[:80] or "PartyOps文档"
    write_audit(db, user, "business_document.export", "business_document", item.id, {"version": item.version}, client_ip(request))
    db.commit()
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}.docx", "Cache-Control": "no-store"})

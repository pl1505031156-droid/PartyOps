"""党员发展规则计算、Word 导出和单位补充材料管理。"""

from __future__ import annotations

import hashlib
import io
import typing
from datetime import date, datetime, time, timedelta, timezone
from pathlib import Path
from urllib.parse import quote

from docx import Document
from fastapi import APIRouter, Depends, Header, Request, Response, status
from fastapi.responses import FileResponse, StreamingResponse
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from sqlalchemy import delete, func, select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session
from starlette.background import BackgroundTask

from ..audit import write_audit
from ..config import get_settings
from ..database import get_session
from ..models import (
    Notification,
    PartyDevelopmentCase,
    PartyDevelopmentMaterial,
    PartyDevelopmentMilestone,
    PartyDevelopmentPlanProfile,
    PartyDevelopmentProfile,
    PartyDevelopmentProgressEvent,
    User,
    WorkCalendarEntry,
    utcnow,
)
from ..party_development import (
    NATIONAL_MATERIALS,
    PHASE_LABELS,
    calculate_party_development,
    calculate_reference_plan,
    ensure_reference_plan_profile,
    ensure_reference_profile,
    export_result_docx,
    profile_to_dict,
    rule_metadata,
    safe_person_filename,
    supplemental_materials,
)
from ..problems import ProblemException
from ..schemas import (
    PartyDevelopmentActualDates,
    PartyDevelopmentCalculateRequest,
    PartyDevelopmentCaseCreate,
    PartyDevelopmentCaseLifecycleAction,
    PartyDevelopmentCasePatch,
    PartyDevelopmentFromCalculationCreate,
    PartyDevelopmentMaterialInput,
    PartyDevelopmentMilestonePatch,
    PartyDevelopmentProfileCreate,
    PartyDevelopmentProfileOut,
    PartyDevelopmentProfilePatch,
    PartyDevelopmentProgressEventCorrect,
    PartyDevelopmentProgressEventCreate,
    PartyDevelopmentProgressEventVoid,
    PartyDevelopmentReferencePlanPatch,
    PartyDevelopmentResultOut,
)
from ..security import get_current_user, require_admin
from ..spreadsheet_security import safe_spreadsheet_row
from .router_utils import client_ip, parse_if_match

router = APIRouter(tags=["party-development"])
settings = get_settings()


def _as_datetime(value: date | None) -> datetime | None:
    return datetime.combine(value, time.min, tzinfo=timezone.utc) if value else None


def _as_date(value: datetime | None) -> date | None:
    return value.date() if value else None


def _aware(value: datetime | None) -> datetime | None:
    if value is None:
        return None
    return value if value.tzinfo else value.replace(tzinfo=timezone.utc)


EVENT_TO_ACTUAL_FIELD = {
    "conversation": "conversation_date",
    "activist_date": "activist_date",
    "activist_publicity_start": "publicity_start_date",
    "development_object_date": "development_object_date",
    "training_completed": "training_completed_date",
    "political_review_completed": "political_review_completed_date",
    "pre_review_approved": "pre_review_approved_date",
    "branch_acceptance": "branch_acceptance_date",
    "committee_approval_actual": "committee_approval_date",
    "oath": "oath_date",
    "transition_application": "transition_application_date",
    "transition_branch_meeting": "transition_branch_meeting_date",
    "transition_approval": "transition_approval_date",
}


def _current_progress_events(
    db: Session, case_id: str
) -> dict[str, PartyDevelopmentProgressEvent]:
    rows = db.scalars(
        select(PartyDevelopmentProgressEvent)
        .where(
            PartyDevelopmentProgressEvent.case_id == case_id,
            PartyDevelopmentProgressEvent.status == "confirmed",
        )
        .order_by(PartyDevelopmentProgressEvent.created_at.desc())
    ).all()
    result: dict[str, PartyDevelopmentProgressEvent] = {}
    for row in rows:
        result.setdefault(row.milestone_type, row)
    return result


def _case_payload(
    item: PartyDevelopmentCase,
    events: dict[str, PartyDevelopmentProgressEvent] | None = None,
) -> PartyDevelopmentCalculateRequest:
    actual_values: dict[str, date | int | float | None] = {
        "activist_date": _as_date(item.activist_at),
        "development_object_date": _as_date(item.development_object_at),
        "branch_acceptance_date": _as_date(item.probationary_at),
        "transition_approval_date": _as_date(item.converted_at),
    }
    for milestone_type, event in (events or {}).items():
        target = EVENT_TO_ACTUAL_FIELD.get(milestone_type)
        if target:
            actual_values[target] = _as_date(event.actual_at)
    return PartyDevelopmentCalculateRequest(
        name=item.name,
        application_date=item.application_at.date(),
        actual_dates=PartyDevelopmentActualDates(**actual_values),
    )


def _case_out(db: Session, item: PartyDevelopmentCase) -> dict[str, typing.Any]:
    progress = list(_current_progress_events(db, item.id).values())
    milestones = db.scalars(
        select(PartyDevelopmentMilestone)
        .where(PartyDevelopmentMilestone.case_id == item.id)
        .order_by(PartyDevelopmentMilestone.planned_at, PartyDevelopmentMilestone.milestone_type)
    ).all()
    return {
        "id": item.id,
        "party_committee": item.party_committee,
        "party_branch": item.party_branch,
        "name": item.name,
        "gender": item.gender,
        "ethnicity": item.ethnicity,
        "birth_date": item.birth_date,
        "education": item.education,
        "application_at": item.application_at,
        "activist_at": item.activist_at,
        "training_contacts": item.training_contacts,
        "introducers": item.introducers,
        "development_object_at": item.development_object_at,
        "probationary_at": item.probationary_at,
        "converted_at": item.converted_at,
        "stage": item.stage,
        "status": item.status,
        "rule_version": item.rule_version,
        "planning_profile_id": item.planning_profile_id,
        "planning_profile_snapshot": item.planning_profile_snapshot,
        "extra_fields": item.extra_fields or {},
        "import_batch_id": item.import_batch_id,
        "version": item.version,
        "progress_events": [
            {
                "id": row.id,
                "milestone_type": row.milestone_type,
                "actual_at": row.actual_at,
                "evidence_note": row.evidence_note,
                "source_entity_type": row.source_entity_type,
                "source_entity_id": row.source_entity_id,
                "status": row.status,
                "supersedes_event_id": row.supersedes_event_id,
                "version": row.version,
                "created_by": row.created_by,
                "created_at": row.created_at,
            }
            for row in progress
        ],
        "milestones": [
            {
                "id": row.id,
                "milestone_type": row.milestone_type,
                "actual_at": row.actual_at,
                "legal_earliest_at": row.legal_earliest_at,
                "legal_deadline_at": row.legal_deadline_at,
                "planned_at": row.planned_at,
                "adjusted_at": row.adjusted_at,
                "rule_version": row.rule_version,
                "legal_basis": row.legal_basis,
                "planning_basis": row.planning_basis,
                "plan_kind": row.plan_kind,
                "reminder_days": row.reminder_days,
                "version": row.version,
            }
            for row in milestones
        ],
    }


def _calculated_reference_plan(
    db: Session, item: PartyDevelopmentCase
) -> dict[str, typing.Any]:
    snapshot = (
        item.planning_profile_snapshot
        if isinstance(item.planning_profile_snapshot, dict)
        else {}
    )
    assumptions = snapshot.get("assumptions")
    if not isinstance(assumptions, dict):
        profile = (
            db.get(PartyDevelopmentPlanProfile, item.planning_profile_id)
            if item.planning_profile_id
            else None
        )
        assumptions = profile.assumptions if profile else {}
    plan = calculate_reference_plan(
        application_date=item.application_at.date(),
        calendar_entries=db.scalars(select(WorkCalendarEntry)).all(),
        activist_date=_as_date(item.activist_at),
        development_object_date=_as_date(item.development_object_at),
        branch_acceptance_date=_as_date(item.probationary_at),
        assumptions=assumptions,
    )
    existing = {
        row.milestone_type: row
        for row in db.scalars(
            select(PartyDevelopmentMilestone).where(
                PartyDevelopmentMilestone.case_id == item.id,
                PartyDevelopmentMilestone.plan_kind == "reference",
            )
        ).all()
    }
    for node in plan["nodes"]:
        row = existing.get(str(node["key"]))
        node["persisted_reference_date"] = _as_date(row.planned_at) if row else None
        node["adjusted_date"] = _as_date(row.adjusted_at) if row else None
        node["effective_date"] = (
            _as_date(row.adjusted_at or row.planned_at)
            if row
            else node["reference_date"]
        )
        node["version"] = row.version if row else 0
    plan["profile_snapshot"] = snapshot
    return plan


def _apply_reference_plan(
    db: Session,
    item: PartyDevelopmentCase,
    *,
    adjustments: dict[str, date | None] | None = None,
    bump_case_version: bool = True,
) -> dict[str, typing.Any]:
    plan = _calculated_reference_plan(db, item)
    adjustment_values = adjustments or {}
    allowed = {str(node["key"]) for node in plan["nodes"]}
    unknown = sorted(set(adjustment_values) - allowed)
    if unknown:
        raise ProblemException(
            422,
            "REFERENCE_PLAN_NODE_INVALID",
            "参考计划节点无效",
            "请刷新参考计划后重新调整。",
            extra={"unknown_keys": unknown},
        )
    previous = {
        row.milestone_type: row
        for row in db.scalars(
            select(PartyDevelopmentMilestone).where(
                PartyDevelopmentMilestone.case_id == item.id,
                PartyDevelopmentMilestone.plan_kind == "reference",
            )
        ).all()
    }
    for node in plan["nodes"]:
        key = str(node["key"])
        row = previous.get(key)
        if row is None:
            row = PartyDevelopmentMilestone(
                case_id=item.id,
                milestone_type=key,
                plan_kind="reference",
            )
            db.add(row)
        else:
            row.version += 1
        row.actual_at = None
        row.legal_earliest_at = None
        row.legal_deadline_at = None
        row.planned_at = _as_datetime(node["reference_date"])
        row.adjusted_at = (
            _as_datetime(adjustment_values[key])
            if key in adjustment_values
            else row.adjusted_at
        )
        row.rule_version = item.rule_version
        row.legal_basis = ""
        row.planning_basis = str(node["planning_basis"])
    if bump_case_version:
        item.version += 1
    db.flush()
    return _calculated_reference_plan(db, item)


def _profile(db: Session, profile_id: str) -> PartyDevelopmentProfile:
    profile = db.get(PartyDevelopmentProfile, profile_id)
    if not profile:
        raise ProblemException(404, "PARTY_DEVELOPMENT_PROFILE_NOT_FOUND", "补充材料模板不存在", "请刷新后重试。")
    return profile


def _check_version(profile: PartyDevelopmentProfile, if_match: str | None) -> None:
    if profile.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "补充材料模板已更新", "请刷新后重试。")


def _add_items(
    db: Session,
    profile: PartyDevelopmentProfile,
    items: list[PartyDevelopmentMaterialInput],
    admin: User,
) -> None:
    identities: set[tuple[str, str]] = set()
    for item in items:
        name = item.name.strip()
        identity = (item.phase, name.casefold())
        if identity in identities:
            raise ProblemException(422, "MATERIAL_DUPLICATE", "补充材料重复", f"{name} 在同一阶段重复。")
        identities.add(identity)
        db.add(PartyDevelopmentMaterial(
            profile_id=profile.id,
            phase=item.phase,
            name=name,
            responsible_party=item.responsible_party.strip(),
            guidance=item.guidance.strip(),
            required=item.required,
            enabled=item.enabled,
            sort_order=item.sort_order,
            created_by=admin.id,
        ))


@router.get("/party-development/rules/current", response_model=dict)
def current_rules(_user: User = Depends(get_current_user)) -> dict:
    return rule_metadata()


@router.get("/party-development/materials", response_model=dict)
def material_checklist(
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    """返回国家规则材料和已启用单位补充项，供全体登录用户查阅。"""

    supplemental = supplemental_materials(db, [])
    phases: list[dict[str, typing.Any]] = []
    for phase, label in PHASE_LABELS.items():
        national = [
            {
                "name": name,
                "source": "国家规则",
                "responsible_party": "按组织程序确认",
                "guidance": "具体表单、签署和归档要求由本单位组织部门确认",
                "required": True,
                "national": True,
            }
            for name in NATIONAL_MATERIALS.get(phase, [])
        ]
        phases.append(
            {
                "phase": phase,
                "label": label,
                "items": national + supplemental.get(phase, []),
            }
        )
    return {
        "rule": rule_metadata(),
        "phases": phases,
        "disclaimer": "材料清单用于办理提醒，不替代组织部门对个案材料和程序的审核。",
    }


@router.post("/party-development/calculate", response_model=PartyDevelopmentResultOut)
def calculate(
    payload: PartyDevelopmentCalculateRequest,
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> PartyDevelopmentResultOut:
    entries = db.scalars(select(WorkCalendarEntry)).all()
    return calculate_party_development(
        payload,
        entries,
        supplemental_materials(db, payload.profile_ids),
    )


@router.get("/party-development/cases", response_model=typing.List[dict])
def list_cases(
    party_committee: str = "",
    party_branch: str = "",
    case_status: str = "active",
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> list[dict[str, typing.Any]]:
    query = select(PartyDevelopmentCase)
    if party_committee:
        query = query.where(PartyDevelopmentCase.party_committee == party_committee)
    if party_branch:
        query = query.where(PartyDevelopmentCase.party_branch == party_branch)
    if case_status:
        query = query.where(PartyDevelopmentCase.status == case_status)
    return [_case_out(db, item) for item in db.scalars(query.order_by(PartyDevelopmentCase.name)).all()]


@router.post("/party-development/cases", response_model=dict, status_code=201)
def create_case(
    payload: PartyDevelopmentCaseCreate,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    profile = ensure_reference_plan_profile(db, user)
    snapshot = {
        "system_key": profile.system_key,
        "name": profile.name,
        "version": profile.version,
        "assumptions": dict(profile.assumptions),
        "captured_at": utcnow().isoformat(),
    }
    item = PartyDevelopmentCase(
        party_committee=payload.party_committee.strip(),
        party_branch=payload.party_branch.strip(),
        name=payload.name.strip(),
        gender=payload.gender.strip(),
        ethnicity=payload.ethnicity.strip(),
        birth_date=_as_datetime(payload.birth_date),
        education=payload.education.strip(),
        application_at=_as_datetime(payload.application_date),
        activist_at=_as_datetime(payload.activist_date),
        training_contacts=[value.strip() for value in payload.training_contacts if value.strip()],
        introducers=[value.strip() for value in payload.introducers if value.strip()],
        development_object_at=_as_datetime(payload.development_object_date),
        probationary_at=_as_datetime(payload.probationary_date),
        converted_at=_as_datetime(payload.converted_date),
        rule_version=str(rule_metadata()["version"]),
        planning_profile_id=profile.id,
        planning_profile_snapshot=snapshot,
        created_by=user.id,
    )
    db.add(item)
    db.flush()
    # 新建档案的初始计划属于同一笔创建事务，不应把初始版本从 1 抬高到 2。
    _apply_reference_plan(db, item, bump_case_version=False)
    write_audit(db, user, "party_development.case_create", "party_development_case", item.id, {"party_branch": item.party_branch}, client_ip(request))
    db.commit()
    return _case_out(db, item)


def _case_or_404(db: Session, case_id: str) -> PartyDevelopmentCase:
    item = db.get(PartyDevelopmentCase, case_id)
    if not item:
        raise ProblemException(
            404,
            "PARTY_DEVELOPMENT_CASE_NOT_FOUND",
            "发展档案不存在",
            "请刷新后重试。",
        )
    return item


def _ensure_case_active(item: PartyDevelopmentCase) -> None:
    if item.status != "active":
        raise ProblemException(
            409,
            "PARTY_DEVELOPMENT_CASE_INACTIVE",
            "发展档案当前不可编辑",
            "请先恢复已归档档案，再维护真实进度。",
        )


def _sync_high_level_facts(
    db: Session,
    item: PartyDevelopmentCase,
    *,
    clear_missing: set[str] | None = None,
) -> None:
    events = _current_progress_events(db, item.id)
    mapping = {
        "activist_date": "activist_at",
        "development_object_date": "development_object_at",
        "branch_acceptance": "probationary_at",
        "transition_approval": "converted_at",
    }
    for event_type, column in mapping.items():
        event = events.get(event_type)
        if event or event_type in (clear_missing or set()):
            setattr(item, column, event.actual_at if event else None)
    if item.converted_at:
        item.stage = "completed"
    elif item.probationary_at:
        item.stage = "probationary"
    elif item.development_object_at:
        item.stage = "development_object"
    elif item.activist_at:
        item.stage = "activist"
    else:
        item.stage = "application"


def _append_progress_fact(
    db: Session,
    item: PartyDevelopmentCase,
    *,
    milestone_type: str,
    actual_at: datetime,
    evidence_note: str,
    source_entity_type: str,
    source_entity_id: str | None,
    user: User,
) -> PartyDevelopmentProgressEvent:
    allowed = {"application", *EVENT_TO_ACTUAL_FIELD}
    if milestone_type not in allowed:
        raise ProblemException(
            422,
            "PARTY_DEVELOPMENT_PROGRESS_TYPE_INVALID",
            "进度节点无效",
            "请选择时间轴中允许录入真实进度的节点。",
        )
    previous = db.scalar(
        select(PartyDevelopmentProgressEvent)
        .where(
            PartyDevelopmentProgressEvent.case_id == item.id,
            PartyDevelopmentProgressEvent.milestone_type == milestone_type,
            PartyDevelopmentProgressEvent.status == "confirmed",
        )
        .order_by(PartyDevelopmentProgressEvent.created_at.desc())
    )
    # SQLite 会丢失时区信息；直接比较 aware/naive datetime 会把同一事实
    # 错判为新记录，形成重复时间轴。统一转为 UTC 后再做幂等判断。
    if previous and _aware(previous.actual_at) == _aware(actual_at):
        raise ProblemException(
            409,
            "PARTY_DEVELOPMENT_PROGRESS_DUPLICATE",
            "该真实进度已经记录",
            "如需补充说明，请使用纠正功能保留修订链。",
        )
    if previous:
        previous.status = "superseded"
        previous.version += 1
    event = PartyDevelopmentProgressEvent(
        case_id=item.id,
        milestone_type=milestone_type,
        actual_at=actual_at,
        evidence_note=evidence_note.strip(),
        source_entity_type=source_entity_type.strip(),
        source_entity_id=source_entity_id,
        supersedes_event_id=previous.id if previous else None,
        created_by=user.id,
    )
    db.add(event)
    db.flush()
    _sync_high_level_facts(db, item)
    _apply_reference_plan(db, item, bump_case_version=False)
    item.version += 1
    return event


@router.post(
    "/party-development/cases/from-calculation",
    response_model=dict,
    status_code=201,
)
def create_case_from_calculation(
    payload: PartyDevelopmentFromCalculationCreate,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    """把快速测算转为正式档案，且只保存用户已填写的真实节点。"""

    base = PartyDevelopmentCaseCreate.model_validate(
        payload.model_dump(exclude={"actual_dates"})
    )
    profile = ensure_reference_plan_profile(db, user)
    item = PartyDevelopmentCase(
        party_committee=base.party_committee.strip(),
        party_branch=base.party_branch.strip(),
        name=base.name.strip(),
        gender=base.gender.strip(),
        ethnicity=base.ethnicity.strip(),
        birth_date=_as_datetime(base.birth_date),
        education=base.education.strip(),
        application_at=_as_datetime(base.application_date),
        training_contacts=[value.strip() for value in base.training_contacts if value.strip()],
        introducers=[value.strip() for value in base.introducers if value.strip()],
        rule_version=str(rule_metadata()["version"]),
        planning_profile_id=profile.id,
        planning_profile_snapshot={
            "system_key": profile.system_key,
            "name": profile.name,
            "version": profile.version,
            "assumptions": dict(profile.assumptions),
            "captured_at": utcnow().isoformat(),
        },
        created_by=user.id,
    )
    db.add(item)
    db.flush()
    reverse = {value: key for key, value in EVENT_TO_ACTUAL_FIELD.items()}
    for field_name, actual_date in payload.actual_dates.model_dump().items():
        if actual_date in (None, 0, 0.0) or field_name in {"training_days", "training_hours"}:
            continue
        event_type = reverse.get(field_name)
        if event_type:
            _append_progress_fact(
                db,
                item,
                milestone_type=event_type,
                actual_at=_as_datetime(actual_date),
                evidence_note="由快速测算建档时确认",
                source_entity_type="quick_calculation",
                source_entity_id=None,
                user=user,
            )
    _apply_reference_plan(db, item, bump_case_version=False)
    write_audit(db, user, "party_development.case_create_from_calculation", "party_development_case", item.id, {"actual_fact_count": len(_current_progress_events(db, item.id))}, client_ip(request))
    db.commit()
    return _case_out(db, item)


@router.get("/party-development/cases/{case_id}/timeline", response_model=dict)
def get_case_timeline(
    case_id: str,
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    item = _case_or_404(db, case_id)
    current = _case_out(db, item)
    now = utcnow()
    facts = {row["milestone_type"]: row for row in current["progress_events"]}
    timeline: list[dict[str, typing.Any]] = []
    for milestone in current["milestones"]:
        target = milestone["adjusted_at"] or milestone["legal_deadline_at"] or milestone["legal_earliest_at"] or milestone["planned_at"]
        actual = milestone["actual_at"]
        fact = facts.get(milestone["milestone_type"])
        if fact:
            actual = fact["actual_at"]
        aware_target = _aware(target)
        if actual:
            visual_state = "completed"
        elif aware_target and aware_target < now:
            visual_state = "overdue"
        elif aware_target and aware_target <= now + timedelta(days=60):
            visual_state = "upcoming"
        else:
            visual_state = "planned"
        timeline.append({**milestone, "actual_at": actual, "progress_event": fact, "visual_state": visual_state, "is_reference": bool(milestone["adjusted_at"] or milestone["plan_kind"] == "reference")})
    present = {row["milestone_type"] for row in timeline}
    for event_type, fact in facts.items():
        if event_type not in present:
            timeline.append({"id": f"fact:{fact['id']}", "milestone_type": event_type, "actual_at": fact["actual_at"], "legal_earliest_at": None, "legal_deadline_at": None, "planned_at": None, "adjusted_at": None, "legal_basis": "", "planning_basis": "", "plan_kind": "fact", "reminder_days": [], "version": fact["version"], "progress_event": fact, "visual_state": "completed", "is_reference": False})
    def sort_key(row: dict[str, typing.Any]) -> float:
        value = row["actual_at"] or row["adjusted_at"] or row["legal_deadline_at"] or row["planned_at"]
        if isinstance(value, str):
            try:
                value = datetime.fromisoformat(value.replace("Z", "+00:00"))
            except ValueError:
                return float("inf")
        return _aware(value).timestamp() if isinstance(value, datetime) else float("inf")

    timeline.sort(key=sort_key)
    return {"case": current, "timeline": timeline, "legend": {"completed": "已实际完成", "overdue": "已逾期未完成", "upcoming": "即将到期", "planned": "法规或参考计划", "reference": "人工调整或不确定预测"}}


@router.post(
    "/party-development/cases/{case_id}/progress-events",
    response_model=dict,
    status_code=201,
)
def create_progress_event(
    case_id: str,
    payload: PartyDevelopmentProgressEventCreate,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    item = _case_or_404(db, case_id)
    _ensure_case_active(item)
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "发展档案已更新", "请刷新时间轴后重试。")
    event = _append_progress_fact(
        db,
        item,
        milestone_type=payload.milestone_type,
        actual_at=_as_datetime(payload.actual_date),
        evidence_note=payload.evidence_note,
        source_entity_type=payload.source_entity_type,
        source_entity_id=payload.source_entity_id,
        user=user,
    )
    write_audit(db, user, "party_development.progress_created", "party_development_progress_event", event.id, {"case_id": item.id, "milestone_type": event.milestone_type}, client_ip(request))
    db.commit()
    return get_case_timeline(item.id, user, db)


@router.post(
    "/party-development/progress-events/{event_id}/correct",
    response_model=dict,
)
def correct_progress_event(
    event_id: str,
    payload: PartyDevelopmentProgressEventCorrect,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    previous = db.get(PartyDevelopmentProgressEvent, event_id)
    if not previous:
        raise ProblemException(404, "PARTY_DEVELOPMENT_PROGRESS_NOT_FOUND", "真实进度不存在", "请刷新时间轴。")
    if previous.version != parse_if_match(if_match) or previous.status != "confirmed":
        raise ProblemException(409, "VERSION_CONFLICT", "真实进度已更新", "请刷新时间轴后重试。")
    item = _case_or_404(db, previous.case_id)
    _ensure_case_active(item)
    event = _append_progress_fact(db, item, milestone_type=previous.milestone_type, actual_at=_as_datetime(payload.actual_date), evidence_note=payload.evidence_note, source_entity_type="correction", source_entity_id=previous.id, user=user)
    write_audit(db, user, "party_development.progress_corrected", "party_development_progress_event", event.id, {"case_id": item.id, "supersedes": previous.id}, client_ip(request))
    db.commit()
    return get_case_timeline(item.id, user, db)


@router.post(
    "/party-development/progress-events/{event_id}/void",
    response_model=dict,
)
def void_progress_event(
    event_id: str,
    payload: PartyDevelopmentProgressEventVoid,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    event = db.get(PartyDevelopmentProgressEvent, event_id)
    if not event:
        raise ProblemException(404, "PARTY_DEVELOPMENT_PROGRESS_NOT_FOUND", "真实进度不存在", "请刷新时间轴。")
    if event.version != parse_if_match(if_match) or event.status != "confirmed":
        raise ProblemException(409, "VERSION_CONFLICT", "真实进度已更新", "请刷新时间轴后重试。")
    item = _case_or_404(db, event.case_id)
    _ensure_case_active(item)
    event.status = "voided"
    event.voided_at = utcnow()
    event.evidence_note = f"{event.evidence_note}\n作废原因：{payload.reason}".strip()
    event.version += 1
    if event.supersedes_event_id:
        previous = db.get(PartyDevelopmentProgressEvent, event.supersedes_event_id)
        if previous and previous.status == "superseded":
            previous.status = "confirmed"
            previous.version += 1
    _sync_high_level_facts(db, item, clear_missing={event.milestone_type})
    _apply_reference_plan(db, item, bump_case_version=False)
    item.version += 1
    write_audit(db, user, "party_development.progress_voided", "party_development_progress_event", event.id, {"case_id": item.id, "reason": payload.reason}, client_ip(request))
    db.commit()
    return get_case_timeline(item.id, user, db)


@router.get(
    "/party-development/cases/{case_id}/deletion-impact",
    response_model=dict,
)
def case_deletion_impact(
    case_id: str,
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    item = _case_or_404(db, case_id)
    milestones = int(db.scalar(select(func.count()).select_from(PartyDevelopmentMilestone).where(PartyDevelopmentMilestone.case_id == item.id)) or 0)
    progress_events = int(db.scalar(select(func.count()).select_from(PartyDevelopmentProgressEvent).where(PartyDevelopmentProgressEvent.case_id == item.id)) or 0)
    active_notifications = int(db.scalar(select(func.count()).select_from(Notification).where(Notification.entity_type == "party_development_case", Notification.entity_id == item.id, Notification.read_at.is_(None), Notification.revoked_at.is_(None))) or 0)
    return {"case_id": item.id, "status": item.status, "milestones": milestones, "progress_events": progress_events, "active_notifications": active_notifications, "action": "archive", "physical_delete": False, "recoverable": True, "message": "作废后从在办台账和活动提醒中移除，节点、材料和审计历史保留，可由有权限人员恢复。"}


@router.delete("/party-development/cases/{case_id}", response_model=dict)
def archive_case(
    case_id: str,
    payload: PartyDevelopmentCaseLifecycleAction,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    item = _case_or_404(db, case_id)
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "发展档案已更新", "请刷新影响范围后重试。")
    if item.status == "archived":
        return _case_out(db, item)
    item.status = "archived"
    item.version += 1
    now = utcnow()
    notifications = db.scalars(select(Notification).where(Notification.entity_type == "party_development_case", Notification.entity_id == item.id, Notification.read_at.is_(None), Notification.revoked_at.is_(None))).all()
    for notification in notifications:
        notification.revoked_at = now
    write_audit(db, user, "party_development.case_archived", "party_development_case", item.id, {"reason": payload.reason, "revoked_notifications": len(notifications)}, client_ip(request))
    db.commit()
    return _case_out(db, item)


@router.post("/party-development/cases/{case_id}/restore", response_model=dict)
def restore_case(
    case_id: str,
    payload: PartyDevelopmentCaseLifecycleAction,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    item = _case_or_404(db, case_id)
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "发展档案已更新", "请刷新后重试。")
    if item.status != "archived":
        raise ProblemException(409, "PARTY_DEVELOPMENT_CASE_NOT_ARCHIVED", "档案不需要恢复", "只有已归档档案可以恢复。")
    item.status = "active"
    item.version += 1
    write_audit(db, user, "party_development.case_restored", "party_development_case", item.id, {"reason": payload.reason}, client_ip(request))
    db.commit()
    return _case_out(db, item)


@router.get("/party-development/cases/{case_id}/reference-plan", response_model=dict)
def get_reference_plan(
    case_id: str,
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    item = db.get(PartyDevelopmentCase, case_id)
    if not item:
        raise ProblemException(404, "PARTY_DEVELOPMENT_CASE_NOT_FOUND", "发展档案不存在", "请刷新后重试。")
    return _calculated_reference_plan(db, item)


@router.post("/party-development/cases/{case_id}/reference-plan/recalculate-preview", response_model=dict)
def preview_reference_plan(
    case_id: str,
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    """只读预览实际日期重锚后的计划，不静默覆盖已保存日期。"""

    item = db.get(PartyDevelopmentCase, case_id)
    if not item:
        raise ProblemException(404, "PARTY_DEVELOPMENT_CASE_NOT_FOUND", "发展档案不存在", "请刷新后重试。")
    result = _calculated_reference_plan(db, item)
    result["requires_confirmation"] = any(
        node["persisted_reference_date"] not in {None, node["reference_date"]}
        for node in result["nodes"]
    )
    return result


@router.put("/party-development/cases/{case_id}/reference-plan", response_model=dict)
def confirm_reference_plan(
    case_id: str,
    payload: PartyDevelopmentReferencePlanPatch,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    item = db.get(PartyDevelopmentCase, case_id)
    if not item:
        raise ProblemException(404, "PARTY_DEVELOPMENT_CASE_NOT_FOUND", "发展档案不存在", "请刷新后重试。")
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "发展档案已更新", "请刷新预览后重新确认。")
    result = _apply_reference_plan(db, item, adjustments=payload.adjustments)
    write_audit(
        db,
        user,
        "party_development.reference_plan_confirm",
        "party_development_case",
        item.id,
        {"adjusted_nodes": sorted(payload.adjustments), "profile_version": item.planning_profile_snapshot.get("version")},
        client_ip(request),
    )
    db.commit()
    return result


@router.patch("/party-development/cases/{case_id}", response_model=dict)
def patch_case(
    case_id: str,
    payload: PartyDevelopmentCasePatch,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    item = db.get(PartyDevelopmentCase, case_id)
    if not item:
        raise ProblemException(404, "PARTY_DEVELOPMENT_CASE_NOT_FOUND", "发展档案不存在", "请刷新后重试。")
    if item.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "发展档案已更新", "请刷新后重试。")
    changes = payload.model_dump(exclude_unset=True)
    if "application_date" in changes and changes["application_date"] is None:
        raise ProblemException(
            422,
            "APPLICATION_DATE_REQUIRED",
            "入党申请书日期不能为空",
            "请填写实际提交日期；系统将据此重新预览后续参考计划。",
        )
    date_fields = {
        "birth_date": "birth_date",
        "application_date": "application_at",
        "activist_date": "activist_at",
        "development_object_date": "development_object_at",
        "probationary_date": "probationary_at",
        "converted_date": "converted_at",
    }
    for incoming, stored in date_fields.items():
        if incoming in changes:
            setattr(item, stored, _as_datetime(changes.pop(incoming)))
    for key, value in changes.items():
        setattr(item, key, value.strip() if isinstance(value, str) else value)
    item.version += 1
    write_audit(db, user, "party_development.case_update", "party_development_case", item.id, {"fields": sorted(payload.model_fields_set)}, client_ip(request))
    db.commit()
    return _case_out(db, item)


@router.post("/party-development/cases/{case_id}/generate-milestones", response_model=dict)
def generate_case_milestones(
    case_id: str,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    """生成法定边界和参考计划，绝不把计算结果写入 actual_at。"""

    item = db.get(PartyDevelopmentCase, case_id)
    if not item:
        raise ProblemException(404, "PARTY_DEVELOPMENT_CASE_NOT_FOUND", "发展档案不存在", "请刷新后重试。")
    result = calculate_party_development(
        _case_payload(item),
        db.scalars(select(WorkCalendarEntry)).all(),
        supplemental_materials(db, []),
    )
    previous = {
        row.milestone_type: row
        for row in db.scalars(select(PartyDevelopmentMilestone).where(PartyDevelopmentMilestone.case_id == item.id)).all()
    }
    generated: set[str] = set()
    for node in result.nodes:
        if not any(
            (
                node.actual_at,
                node.legal_earliest_at,
                node.legal_deadline_at,
                node.reference_at,
            )
        ):
            continue
        generated.add(node.key)
        row = previous.get(node.key)
        is_new = row is None
        if is_new:
            row = PartyDevelopmentMilestone(case_id=item.id, milestone_type=node.key, version=1)
            db.add(row)
        row.actual_at = _as_datetime(node.actual_at)
        row.legal_earliest_at = _as_datetime(node.legal_earliest_at)
        row.legal_deadline_at = _as_datetime(node.legal_deadline_at)
        row.planned_at = _as_datetime(node.reference_at)
        row.rule_version = result.rule_version
        row.legal_basis = f"{node.article}：{node.basis}"
        row.planning_basis = node.reference_basis
        row.plan_kind = "reference" if node.reference_at else "legal"
        if not is_new:
            row.version += 1
    item.rule_version = result.rule_version
    item.version += 1
    write_audit(db, user, "party_development.milestones_generate", "party_development_case", item.id, {"rule_version": result.rule_version, "count": len(generated)}, client_ip(request))
    db.commit()
    return _case_out(db, item)


@router.patch("/party-development/milestones/{milestone_id}", response_model=dict)
def patch_milestone(
    milestone_id: str,
    payload: PartyDevelopmentMilestonePatch,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    row = db.get(PartyDevelopmentMilestone, milestone_id)
    if not row:
        raise ProblemException(404, "PARTY_DEVELOPMENT_MILESTONE_NOT_FOUND", "节点不存在", "请刷新后重试。")
    if row.version != parse_if_match(if_match):
        raise ProblemException(409, "VERSION_CONFLICT", "节点已更新", "请刷新后重试。")
    if "actual_date" in payload.model_fields_set:
        row.actual_at = _as_datetime(payload.actual_date)
    if "adjusted_date" in payload.model_fields_set:
        row.adjusted_at = _as_datetime(payload.adjusted_date)
    if payload.reminder_days is not None:
        if any(value < 0 or value > 3650 for value in payload.reminder_days):
            raise ProblemException(422, "REMINDER_DAYS_INVALID", "提醒天数无效", "提醒天数必须在 0 到 3650 天之间。")
        row.reminder_days = sorted(set(payload.reminder_days), reverse=True)
    row.version += 1
    write_audit(db, user, "party_development.milestone_update", "party_development_milestone", row.id, {"fields": sorted(payload.model_fields_set)}, client_ip(request))
    db.commit()
    return _case_out(db, db.get(PartyDevelopmentCase, row.case_id))


@router.get("/party-development/statistics", response_model=dict)
def development_statistics(
    party_committee: str = "",
    party_branch: str = "",
    _user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> dict[str, typing.Any]:
    query = select(PartyDevelopmentCase).where(PartyDevelopmentCase.status == "active")
    if party_committee:
        query = query.where(PartyDevelopmentCase.party_committee == party_committee)
    if party_branch:
        query = query.where(PartyDevelopmentCase.party_branch == party_branch)
    cases = db.scalars(query).all()
    stage_counts: dict[str, int] = {}
    for item in cases:
        stage_counts[item.stage] = stage_counts.get(item.stage, 0) + 1
    case_ids = [item.id for item in cases]
    milestones = db.scalars(select(PartyDevelopmentMilestone).where(PartyDevelopmentMilestone.case_id.in_(case_ids))).all() if case_ids else []
    now = utcnow()
    targets = [
        (row, _aware(row.adjusted_at or row.planned_at or row.legal_deadline_at))
        for row in milestones
    ]
    upcoming = sum(1 for row, target in targets if not row.actual_at and target and now <= target <= now + timedelta(days=60))
    overdue = sum(1 for row, target in targets if not row.actual_at and target and target < now)
    return {"total": len(cases), "stage_counts": stage_counts, "upcoming_60_days": upcoming, "overdue": overdue}


def _export_cases_query(party_committee: str, party_branch: str):
    query = select(PartyDevelopmentCase).where(PartyDevelopmentCase.status == "active")
    if party_committee:
        query = query.where(PartyDevelopmentCase.party_committee == party_committee)
    if party_branch:
        query = query.where(PartyDevelopmentCase.party_branch == party_branch)
    return query.order_by(PartyDevelopmentCase.party_committee, PartyDevelopmentCase.party_branch, PartyDevelopmentCase.name)


def _display_date(value: datetime | None) -> str:
    return value.date().isoformat() if value else ""


def _development_export_rows(cases: list[PartyDevelopmentCase]) -> list[list[str]]:
    return [
        [
            item.party_committee,
            item.party_branch,
            item.name,
            item.gender,
            item.ethnicity,
            _display_date(item.birth_date),
            item.education,
            _display_date(item.application_at),
            _display_date(item.activist_at),
            "、".join(item.training_contacts or []),
            "、".join(item.introducers or []),
            _display_date(item.development_object_at),
            _display_date(item.probationary_at),
            _display_date(item.converted_at),
            item.stage,
        ]
        for item in cases
    ]


DEVELOPMENT_EXPORT_HEADERS = [
    "所属党委",
    "所属党支部",
    "姓名",
    "性别",
    "民族",
    "出生年月日",
    "文化程度",
    "提交入党申请时间",
    "列为积极分子时间",
    "培养联系人",
    "入党介绍人",
    "列为发展对象时间",
    "列为预备党员时间",
    "预备党员转正时间",
    "当前阶段",
]


@router.get("/party-development/cases/export.docx")
def export_development_cases_docx(
    request: Request,
    party_committee: str = "",
    party_branch: str = "",
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> StreamingResponse:
    """导出当前筛选范围内的发展党员档案汇总表。"""

    cases = list(db.scalars(_export_cases_query(party_committee, party_branch)).all())
    document = Document()
    document.add_heading("党员发展情况统计表", level=1)
    document.add_paragraph(f"档案人数：{len(cases)} 人；规则版本：{rule_metadata()['version']}")
    table = document.add_table(rows=1, cols=len(DEVELOPMENT_EXPORT_HEADERS))
    table.style = "Table Grid"
    for index, heading in enumerate(DEVELOPMENT_EXPORT_HEADERS):
        table.rows[0].cells[index].text = heading
    for row in _development_export_rows(cases):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    output = io.BytesIO()
    document.save(output)
    export_hash = hashlib.sha256(output.getvalue()).hexdigest()
    output.seek(0)
    write_audit(db, user, "party_development.cases_exported", "party_development_case", None, {"format": "docx", "count": len(cases), "party_committee": party_committee, "party_branch": party_branch, "sha256": export_hash}, client_ip(request))
    db.commit()
    filename = quote("党员发展情况统计表.docx")
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}", "Cache-Control": "no-store"},
    )


@router.get("/party-development/cases/export.xlsx")
def export_development_cases_xlsx(
    request: Request,
    party_committee: str = "",
    party_branch: str = "",
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> StreamingResponse:
    """导出可继续筛选、统计的发展党员档案工作簿。"""

    cases = list(db.scalars(_export_cases_query(party_committee, party_branch)).all())
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "党员发展情况"
    sheet.append(DEVELOPMENT_EXPORT_HEADERS)
    for row in _development_export_rows(cases):
        sheet.append(safe_spreadsheet_row(row))
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for column in sheet.columns:
        letter = column[0].column_letter
        sheet.column_dimensions[letter].width = min(28, max(12, max(len(str(cell.value or "")) for cell in column) + 2))
    output = io.BytesIO()
    workbook.save(output)
    export_hash = hashlib.sha256(output.getvalue()).hexdigest()
    output.seek(0)
    write_audit(db, user, "party_development.cases_exported", "party_development_case", None, {"format": "xlsx", "count": len(cases), "party_committee": party_committee, "party_branch": party_branch, "sha256": export_hash}, client_ip(request))
    db.commit()
    filename = quote("党员发展情况统计表.xlsx")
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}", "Cache-Control": "no-store"},
    )


def _delete_temporary_export(path: Path) -> None:
    try:
        path.unlink(missing_ok=True)
    except OSError:
        # 失败时不影响已经完成的下载；诊断清理会处理 exports 中的旧临时文件。
        pass


@router.post("/party-development/export.docx")
def export_docx(
    payload: PartyDevelopmentCalculateRequest,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_session),
) -> FileResponse:
    entries = db.scalars(select(WorkCalendarEntry)).all()
    result = calculate_party_development(
        payload,
        entries,
        supplemental_materials(db, payload.profile_ids),
    )
    path = export_result_docx(result, settings.exports_dir)
    write_audit(
        db,
        user,
        "party_development.exported",
        "party_development_calculation",
        None,
        {"rule_version": result.rule_version, "node_count": len(result.nodes), "warning_count": len(result.warnings)},
        client_ip(request),
    )
    db.commit()
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{safe_person_filename(payload.name)}-党员发展时间节点.docx",
        background=BackgroundTask(_delete_temporary_export, path),
    )


@router.get("/admin/party-development/profiles", response_model=typing.List[PartyDevelopmentProfileOut])
def list_profiles(
    admin: User = Depends(require_admin),
    db: Session = Depends(get_session),
) -> list[dict]:
    ensure_reference_profile(db, admin)
    db.commit()
    profiles = db.scalars(
        select(PartyDevelopmentProfile).order_by(PartyDevelopmentProfile.created_at, PartyDevelopmentProfile.name)
    ).all()
    return [profile_to_dict(db, profile) for profile in profiles]


@router.post(
    "/admin/party-development/profiles",
    response_model=PartyDevelopmentProfileOut,
    status_code=status.HTTP_201_CREATED,
)
def create_profile(
    payload: PartyDevelopmentProfileCreate,
    request: Request,
    admin: User = Depends(require_admin),
    db: Session = Depends(get_session),
) -> dict:
    profile = PartyDevelopmentProfile(
        name=payload.name.strip(),
        description=payload.description.strip(),
        source_label=payload.source_label.strip(),
        active=payload.active,
        created_by=admin.id,
    )
    db.add(profile)
    try:
        db.flush()
        _add_items(db, profile, payload.items, admin)
        write_audit(db, admin, "party_development.profile_created", "party_development_profile", profile.id, {
            "active": profile.active, "item_count": len(payload.items),
        }, client_ip(request))
        db.commit()
    except IntegrityError as exc:
        db.rollback()
        raise ProblemException(409, "PROFILE_NAME_EXISTS", "模板名称已存在", "请换一个模板名称。") from exc
    db.refresh(profile)
    return profile_to_dict(db, profile)


@router.patch("/admin/party-development/profiles/{profile_id}", response_model=PartyDevelopmentProfileOut)
def patch_profile(
    profile_id: str,
    payload: PartyDevelopmentProfilePatch,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    admin: User = Depends(require_admin),
    db: Session = Depends(get_session),
) -> dict:
    profile = _profile(db, profile_id)
    _check_version(profile, if_match)
    changes = payload.model_dump(exclude_unset=True)
    for field in ("name", "description", "source_label"):
        if field in changes:
            changes[field] = changes[field].strip()
    for key, value in changes.items():
        setattr(profile, key, value)
    profile.version += 1
    write_audit(db, admin, "party_development.profile_updated", "party_development_profile", profile.id, {
        "fields": sorted(changes), "active": profile.active,
    }, client_ip(request))
    try:
        db.commit()
    except IntegrityError as exc:
        db.rollback()
        raise ProblemException(409, "PROFILE_NAME_EXISTS", "模板名称已存在", "请换一个模板名称。") from exc
    db.refresh(profile)
    return profile_to_dict(db, profile)


@router.delete(
    "/admin/party-development/profiles/{profile_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    response_class=Response,
)
def delete_profile(
    profile_id: str,
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    admin: User = Depends(require_admin),
    db: Session = Depends(get_session),
) -> Response:
    profile = _profile(db, profile_id)
    _check_version(profile, if_match)
    detail = {"active": profile.active}
    db.delete(profile)
    write_audit(db, admin, "party_development.profile_deleted", "party_development_profile", profile_id, detail, client_ip(request))
    db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@router.put("/admin/party-development/profiles/{profile_id}/items", response_model=PartyDevelopmentProfileOut)
def replace_profile_items(
    profile_id: str,
    payload: list[PartyDevelopmentMaterialInput],
    request: Request,
    if_match: str | None = Header(default=None, alias="If-Match"),
    admin: User = Depends(require_admin),
    db: Session = Depends(get_session),
) -> dict:
    if len(payload) > 200:
        raise ProblemException(422, "TOO_MANY_MATERIALS", "补充材料过多", "单个模板最多保存 200 条材料。")
    profile = _profile(db, profile_id)
    _check_version(profile, if_match)
    db.execute(delete(PartyDevelopmentMaterial).where(PartyDevelopmentMaterial.profile_id == profile.id))
    _add_items(db, profile, payload, admin)
    profile.version += 1
    write_audit(db, admin, "party_development.materials_replaced", "party_development_profile", profile.id, {
        "item_count": len(payload),
    }, client_ip(request))
    db.commit()
    db.refresh(profile)
    return profile_to_dict(db, profile)

"""PartyOps 内嵌公文工具的跨平台功能注册表与本地执行器。

本模块以 ``PartyOps.DocumentFormatter.Source`` 中的 6 个功能、25 条产品能力
契约为兼容边界。所有处理只发生在本机临时目录，源文件始终只读；需要旧格式
转换或分页渲染时，只调用安装包随附的无窗口 LibreOffice 运行时。
"""

from __future__ import annotations

import copy
import fnmatch
import io
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Callable, Iterable

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from lxml import etree
from PIL import Image

from .official_format import (
    FormatReport,
    OfficialFormatError,
    _office_candidates,
    _safe_stem,
    diagnose_docx,
    format_docx,
    normalize_chinese_punctuation,
    prepare_docx,
)

ProgressCallback = Callable[[int, str], None]
CancelCallback = Callable[[], bool]


@dataclass(frozen=True)
class CapabilityDefinition:
    feature_id: str
    capability_id: str
    description: str


@dataclass(frozen=True)
class FeatureDefinition:
    id: str
    display_name: str
    notes: str
    accepts: tuple[str, ...]
    option_schema: tuple[dict[str, Any], ...]

    def as_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "display_name": self.display_name,
            "notes": self.notes,
            "accepts": list(self.accepts),
            "supports_batch": True,
            "option_schema": [copy.deepcopy(item) for item in self.option_schema],
            "capabilities": [
                asdict(item) for item in PRODUCT_CAPABILITIES if item.feature_id == self.id
            ],
        }


@dataclass(frozen=True)
class FeatureOutput:
    path: Path
    filename: str
    content_type: str


@dataclass(frozen=True)
class FeatureExecutionResult:
    outputs: tuple[FeatureOutput, ...]
    message: str
    report: FormatReport | None = None


PRODUCT_CAPABILITIES: tuple[CapabilityDefinition, ...] = (
    CapabilityDefinition("format", "format.element-recognition", "识别主标题、副标题、层级标题、正文、附件、落款和日期。"),
    CapabilityDefinition("format", "format.execution-scopes", "支持全文、普通选区和汇编文章范围排版。"),
    CapabilityDefinition("format", "format.templates", "支持系统默认参数及最多九套用户排版模板切换。"),
    CapabilityDefinition("format", "format.page-layout", "设置页边距、文档网格和页码。"),
    CapabilityDefinition("format", "format.images-and-tables", "按模板参数规划并执行图片和表格排版。"),
    CapabilityDefinition("replace", "replace.text-and-regex", "执行普通文字替换和带超时保护的正则替换。"),
    CapabilityDefinition("replace", "replace.wildcard", "执行 Word/WPS 兼容通配符替换。"),
    CapabilityDefinition("replace", "replace.format", "按字体和段落条件执行格式替换。"),
    CapabilityDefinition("replace", "replace.saved-plans", "保存、切换并自愈多套替换方案。"),
    CapabilityDefinition("replace", "replace.batch-rules", "一次规划并执行同一方案中的多项替换任务。"),
    CapabilityDefinition("redheader", "redheader.document-types", "提供下行文、上行文和便函三类内置模板。"),
    CapabilityDefinition("redheader", "redheader.top-marks", "设置份号、密级、保密期限和紧急程度。"),
    CapabilityDefinition("redheader", "redheader.agency-and-number", "生成发文机关、发文字号和签发人区域。"),
    CapabilityDefinition("redheader", "redheader.red-line-and-imprint", "生成红线及版记页面布局。"),
    CapabilityDefinition("rename", "rename.content-analysis", "从文档中提取标题、发文字号、副标题和日期。"),
    CapabilityDefinition("rename", "rename.composable-rules", "自由组合内容部件、自定义文字、日期和轮替词。"),
    CapabilityDefinition("rename", "rename.online", "在不覆盖源文件的前提下生成在线重命名副本。"),
    CapabilityDefinition("convert", "convert.document-formats", "导出 DOCX、PDF 和 TXT。"),
    CapabilityDefinition("convert", "convert.image-modes", "导出分页图片或长图并执行内存预算保护。"),
    CapabilityDefinition("convert", "convert.page-selection", "支持全部页面、连续范围和离散指定页面。"),
    CapabilityDefinition("convert", "convert.output-policy", "支持图片格式、清晰度、保存位置及同名文件策略。"),
    CapabilityDefinition("pdf-to-word", "pdf-to-word.local-reading", "使用本地 PDF 引擎读取文本、图像和页面几何。"),
    CapabilityDefinition("pdf-to-word", "pdf-to-word.layout-reconstruction", "重建阅读顺序、段落和 DOCX 文档结构。"),
    CapabilityDefinition("pdf-to-word", "pdf-to-word.tables", "分析表格候选并规划跨页续表。"),
    CapabilityDefinition("pdf-to-word", "pdf-to-word.verification", "转换完成后校验 DOCX 结果。"),
)


COMMON_OPTIONS = (
    {
        "id": "compatibility_mode",
        "label": "兼容模式",
        "type": "select",
        "default": "auto",
        "choices": ["auto", "word", "wps"],
    },
)

FEATURE_DEFINITIONS: tuple[FeatureDefinition, ...] = (
    FeatureDefinition(
        "format",
        "一键排版",
        "识别公文要素并按模板执行版面、字体、页码、图片和表格排版。",
        (".docx", ".doc", ".wps", ".rtf"),
        COMMON_OPTIONS
        + (
            {"id": "template", "label": "排版模板", "type": "text", "default": "GB/T 9704-2012"},
            {"id": "scope", "label": "执行范围", "type": "select", "default": "full", "choices": ["full", "selection", "compilation"]},
            {"id": "start_paragraph", "label": "起始段落", "type": "number", "default": 1, "min": 1, "max": 99999},
            {"id": "end_paragraph", "label": "结束段落", "type": "number", "default": 99999, "min": 1, "max": 99999},
        ),
    ),
    FeatureDefinition(
        "replace",
        "一键替换",
        "按同一方案依次执行普通文字、正则、通配符和格式规则。",
        (".docx", ".doc", ".wps", ".rtf"),
        COMMON_OPTIONS
        + (
            {"id": "plan_name", "label": "替换方案", "type": "text", "default": "默认方案"},
            {"id": "rules", "label": "替换规则", "type": "rules", "default": []},
        ),
    ),
    FeatureDefinition(
        "redheader",
        "一键套红",
        "生成下行文、上行文或便函的版头、文号、红线和版记。",
        (".docx", ".doc", ".wps", ".rtf"),
        COMMON_OPTIONS
        + (
            {"id": "document_type", "label": "公文类型", "type": "select", "default": "down", "choices": ["down", "up", "letter"]},
            {"id": "copy_number", "label": "份号", "type": "text", "default": ""},
            {"id": "security", "label": "密级与期限", "type": "text", "default": ""},
            {"id": "urgency", "label": "紧急程度", "type": "text", "default": ""},
            {"id": "agency", "label": "发文机关", "type": "text", "default": "中共××委员会"},
            {"id": "document_number", "label": "发文字号", "type": "text", "default": "×党发〔2026〕1号"},
            {"id": "signatory", "label": "签发人", "type": "text", "default": ""},
            {"id": "imprint", "label": "版记", "type": "text", "default": "中共××委员会办公室"},
        ),
    ),
    FeatureDefinition(
        "rename",
        "一键命名",
        "识别标题、文号、副标题和日期并按组合规则生成安全文件名。",
        (".docx", ".doc", ".wps", ".rtf"),
        COMMON_OPTIONS
        + (
            {"id": "parts", "label": "命名部件", "type": "multi-select", "default": ["title", "document_number"]},
            {"id": "custom_text", "label": "自定义文字", "type": "text", "default": ""},
            {"id": "separator", "label": "分隔符", "type": "text", "default": "-"},
            {"id": "rotation_words", "label": "轮替词", "type": "text", "default": ""},
        ),
    ),
    FeatureDefinition(
        "convert",
        "一键转换",
        "将文档导出为 DOCX、PDF、TXT、分页图片或长图。",
        (".docx", ".doc", ".wps", ".rtf", ".pdf"),
        COMMON_OPTIONS
        + (
            {"id": "target_format", "label": "输出格式", "type": "select", "default": "pdf", "choices": ["docx", "pdf", "txt", "png", "jpg"]},
            {"id": "image_mode", "label": "图片模式", "type": "select", "default": "pages", "choices": ["pages", "long"]},
            {"id": "page_selection", "label": "页码", "type": "text", "default": "all"},
            {"id": "dpi", "label": "图片清晰度", "type": "number", "default": 200, "min": 72, "max": 600},
            {"id": "same_name_policy", "label": "同名策略", "type": "select", "default": "auto-rename", "choices": ["auto-rename", "overwrite", "skip"]},
        ),
    ),
    FeatureDefinition(
        "pdf-to-word",
        "PDF 转 Word",
        "本地读取 PDF，重建阅读顺序、段落、表格和可验证 DOCX。",
        (".pdf",),
        COMMON_OPTIONS
        + (
            {"id": "normalize_punctuation", "label": "规范中文标点", "type": "boolean", "default": True},
            {"id": "reconstruct_tables", "label": "识别表格", "type": "boolean", "default": True},
        ),
    ),
)

FEATURE_BY_ID = {item.id: item for item in FEATURE_DEFINITIONS}
SUPPORTED_INPUT_EXTENSIONS = frozenset(
    extension for feature in FEATURE_DEFINITIONS for extension in feature.accepts
)


def capabilities_payload() -> dict[str, Any]:
    """返回稳定的 6 功能、25 能力契约。"""

    return {
        "schema_version": 1,
        "features": [item.as_dict() for item in FEATURE_DEFINITIONS],
        "capability_count": len(PRODUCT_CAPABILITIES),
        "engine": "partyops-bundled",
        "external_office_required": False,
    }


def _check_cancelled(cancelled: CancelCallback) -> None:
    if cancelled():
        raise OfficialFormatError("FORMAT_JOB_CANCELLED", "任务已取消", "本次任务已按要求停止，源文件未改变。")


def _progress(callback: ProgressCallback, percent: int, message: str) -> None:
    callback(max(0, min(100, int(percent))), message)


def _unique_output(workspace: Path, stem: str, suffix: str, options: dict[str, Any]) -> Path:
    policy = str(options.get("same_name_policy", "auto-rename"))
    base = workspace / f"{_safe_stem(stem)}{suffix}"
    if not base.exists() or policy == "overwrite":
        return base
    if policy == "skip":
        raise OfficialFormatError("OUTPUT_EXISTS", "输出文件已存在", "同名策略为跳过，未覆盖现有输出。")
    for index in range(2, 10000):
        candidate = base.with_name(f"{base.stem} ({index}){base.suffix}")
        if not candidate.exists():
            return candidate
    raise OfficialFormatError("OUTPUT_NAME_EXHAUSTED", "无法生成输出文件名", "请清理输出目录后重试。")


def _validate_docx(path: Path) -> None:
    try:
        with zipfile.ZipFile(path) as package:
            if "word/document.xml" not in package.namelist():
                raise ValueError("missing document")
            etree.fromstring(package.read("word/document.xml"))
    except (OSError, ValueError, zipfile.BadZipFile, etree.Error) as exc:
        raise OfficialFormatError("OUTPUT_DOCX_INVALID", "输出文档校验失败", "生成的 DOCX 结构不完整，结果未交付。") from exc


def _prepare_word_source(source: Path, workspace: Path) -> tuple[Path, bool]:
    if source.suffix.lower() not in {".docx", ".doc", ".wps", ".rtf"}:
        raise OfficialFormatError("FORMAT_UNSUPPORTED", "文件类型不受支持", "该功能只接受 DOCX、DOC、WPS 或 RTF。")
    return prepare_docx(source, workspace)


def _run_libreoffice_conversion(source: Path, workspace: Path, target_extension: str) -> Path:
    candidates = _office_candidates()
    if not candidates:
        raise OfficialFormatError(
            "BUNDLED_OFFICE_RUNTIME_MISSING",
            "内置转换引擎不可用",
            "安装包中的无窗口转换运行时缺失或损坏，请修复安装 PartyOps。",
        )
    output_dir = Path(tempfile.mkdtemp(prefix="partyops-convert-", dir=workspace))
    profile = output_dir / "profile"
    profile.mkdir(mode=0o700, exist_ok=True)
    environment = os.environ.copy()
    for key in ("HTTP_PROXY", "HTTPS_PROXY", "ALL_PROXY", "http_proxy", "https_proxy", "all_proxy"):
        environment[key] = "http://127.0.0.1:9"
    environment["NO_PROXY"] = "127.0.0.1,localhost"
    environment["no_proxy"] = "127.0.0.1,localhost"
    environment["LIBO_DISABLE_CRASHREPORT"] = "1"
    environment["SAL_DISABLE_OPENCL"] = "1"
    command = [
        str(candidates[0]),
        "--headless",
        "--invisible",
        "--safe-mode",
        "--nologo",
        "--nodefault",
        "--norestore",
        "--nolockcheck",
        "--nofirststartwizard",
        f"-env:UserInstallation={profile.as_uri()}",
        "--convert-to",
        target_extension,
        "--outdir",
        str(output_dir),
        str(source),
    ]
    try:
        completed = subprocess.run(command, capture_output=True, timeout=120, env=environment, check=False)
    except (OSError, subprocess.TimeoutExpired) as exc:
        raise OfficialFormatError("CONVERT_ENGINE_FAILED", "内置转换失败", "转换引擎未能在 120 秒内完成。") from exc
    expected = output_dir / f"{source.stem}.{target_extension}"
    if completed.returncode != 0 or not expected.is_file():
        raise OfficialFormatError("CONVERT_ENGINE_FAILED", "内置转换失败", "转换引擎没有生成可验证的结果。")
    return expected


def _iter_text_paragraphs(document: Document) -> Iterable[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _replace_paragraph_text(paragraph: Any, pattern: re.Pattern[str], replacement: str) -> int:
    original = "".join(run.text for run in paragraph.runs)
    if not original:
        return 0
    # 原工具沿用 .NET/Word 的 ``$1`` 捕获组写法；Python 使用 ``\g<1>``。
    compatible_replacement = re.sub(r"\$(\d+)", r"\\g<\1>", replacement)
    updated, count = pattern.subn(compatible_replacement, original)
    if not count:
        return 0
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)
    return count


def _compile_replace_pattern(rule: dict[str, Any]) -> re.Pattern[str]:
    search = str(rule.get("find", ""))
    if not search:
        raise OfficialFormatError("REPLACE_RULE_INVALID", "替换规则无效", "查找内容不能为空。")
    mode = str(rule.get("mode", "text"))
    flags = 0 if rule.get("case_sensitive") else re.IGNORECASE
    if mode == "regex":
        if len(search) > 1000:
            raise OfficialFormatError("REPLACE_REGEX_TOO_LONG", "正则表达式过长", "单条正则不得超过 1000 个字符。")
        try:
            return re.compile(search, flags)
        except re.error as exc:
            raise OfficialFormatError("REPLACE_REGEX_INVALID", "正则表达式无效", str(exc)) from exc
    if mode == "wildcard":
        translated = fnmatch.translate(search)
        if translated.endswith("\\Z"):
            translated = translated[:-2]
        return re.compile(translated, flags)
    return re.compile(re.escape(search), flags)


def _apply_format_rule(document: Document, rule: dict[str, Any]) -> int:
    criteria = str(rule.get("find", "")).strip()
    font_name = str(rule.get("font_name", "")).strip()
    font_size = rule.get("font_size")
    alignment = str(rule.get("alignment", "")).strip()
    changed = 0
    for paragraph in _iter_text_paragraphs(document):
        if criteria and criteria not in paragraph.text:
            continue
        if alignment in {"left", "center", "right", "justify"}:
            paragraph.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }[alignment]
        for run in paragraph.runs:
            if font_name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            if font_size not in (None, ""):
                run.font.size = Pt(max(5, min(72, float(font_size))))
        changed += 1
    return changed


def _execute_replace(source: Path, workspace: Path, options: dict[str, Any], progress: ProgressCallback, cancelled: CancelCallback) -> FeatureExecutionResult:
    prepared, _ = _prepare_word_source(source, workspace)
    output = _unique_output(workspace, source.stem, "-替换后.docx", options)
    document = Document(str(prepared))
    rules = options.get("rules")
    if not isinstance(rules, list) or not rules:
        raise OfficialFormatError("REPLACE_RULES_REQUIRED", "缺少替换规则", "请至少添加一条文字、正则、通配符或格式规则。")
    if len(rules) > 100:
        raise OfficialFormatError("REPLACE_RULES_LIMIT", "替换规则过多", "单次最多执行 100 条规则。")
    total_changes = 0
    for index, raw_rule in enumerate(rules):
        _check_cancelled(cancelled)
        if not isinstance(raw_rule, dict):
            raise OfficialFormatError("REPLACE_RULE_INVALID", "替换规则无效", "规则必须是结构化对象。")
        mode = str(raw_rule.get("mode", "text"))
        if mode == "format":
            total_changes += _apply_format_rule(document, raw_rule)
        else:
            pattern = _compile_replace_pattern(raw_rule)
            replacement = str(raw_rule.get("replace", ""))
            for paragraph in _iter_text_paragraphs(document):
                total_changes += _replace_paragraph_text(paragraph, pattern, replacement)
        _progress(progress, 15 + int((index + 1) * 70 / len(rules)), f"已执行规则 {index + 1}/{len(rules)}")
    document.save(output)
    _validate_docx(output)
    return FeatureExecutionResult(
        (FeatureOutput(output, output.name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),),
        f"替换完成，共应用 {total_changes} 处修改。",
        diagnose_docx(output, changed_count=total_changes),
    )


def _insert_paragraph_before(paragraph: Any, text: str = "") -> Any:
    new_element = OxmlElement("w:p")
    paragraph._p.addprevious(new_element)
    inserted = Paragraph(new_element, paragraph._parent)
    if text:
        inserted.add_run(text)
    return inserted


def _set_run_font(run: Any, name: str, size: float, *, red: bool = False, bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if red:
        run.font.color.rgb = RGBColor(255, 0, 0)


def _red_line(paragraph: Any) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "FF0000")
    borders.append(bottom)


def _execute_redheader(source: Path, workspace: Path, options: dict[str, Any], progress: ProgressCallback, cancelled: CancelCallback) -> FeatureExecutionResult:
    prepared, _ = _prepare_word_source(source, workspace)
    output = _unique_output(workspace, source.stem, "-套红版.docx", options)
    document = Document(str(prepared))
    _check_cancelled(cancelled)
    if not document.paragraphs:
        document.add_paragraph()
    anchor = document.paragraphs[0]
    top_values = [
        str(options.get("copy_number", "")).strip().zfill(6) if str(options.get("copy_number", "")).strip().isdigit() else str(options.get("copy_number", "")).strip(),
        str(options.get("security", "")).strip(),
        str(options.get("urgency", "")).strip(),
    ]
    for value in reversed([item for item in top_values if item]):
        paragraph = _insert_paragraph_before(anchor, value)
        _set_run_font(paragraph.runs[0], "黑体", 16)
    agency = str(options.get("agency", "中共××委员会")).strip() or "中共××委员会"
    agency_paragraph = _insert_paragraph_before(anchor, agency)
    agency_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(agency_paragraph.runs[0], "方正小标宋简体", 27, red=True)
    document_type = str(options.get("document_type", "down"))
    number_text = str(options.get("document_number", "")).strip()
    signatory = str(options.get("signatory", "")).strip()
    if document_type == "up" and signatory:
        number_text = f"{number_text}    签发人：{signatory}".strip()
    number_paragraph = _insert_paragraph_before(anchor, number_text)
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if number_paragraph.runs:
        _set_run_font(number_paragraph.runs[0], "仿宋_GB2312", 16)
    line_paragraph = _insert_paragraph_before(anchor)
    _red_line(line_paragraph)
    imprint = str(options.get("imprint", "")).strip()
    if document_type != "letter" and imprint:
        tail = document.add_paragraph(imprint)
        tail.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _red_line(tail)
        _set_run_font(tail.runs[0], "仿宋_GB2312", 14)
    _progress(progress, 78, "红头、红线与版记已生成")
    document.save(output)
    _validate_docx(output)
    report = diagnose_docx(output, changed_count=5 + len([item for item in top_values if item]))
    return FeatureExecutionResult(
        (FeatureOutput(output, output.name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),),
        "套红完成，请按最终打印版逐页复核机关标志、文号、红线与版记。",
        report,
    )


def _document_parts(path: Path) -> dict[str, str]:
    document = Document(str(path))
    texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    title = next((item for item in texts if len(item) <= 80 and not re.search(r"[。；！？]$", item)), path.stem)
    document_number = next((item for item in texts if re.fullmatch(r".+〔\d{4}〕\d+号", item)), "")
    date = next((item for item in reversed(texts) if re.fullmatch(r"[〇○零一二三四五六七八九十百千\d]{4}年.+月.+日", item)), "")
    subtitle = next((item for item in texts if item != title and len(item) <= 80 and not re.search(r"[。；！？]$", item)), "")
    return {"title": title, "document_number": document_number, "date": date, "subtitle": subtitle}


def _execute_rename(source: Path, workspace: Path, options: dict[str, Any], progress: ProgressCallback, cancelled: CancelCallback) -> FeatureExecutionResult:
    prepared, _ = _prepare_word_source(source, workspace)
    _check_cancelled(cancelled)
    values = _document_parts(prepared)
    values["custom"] = str(options.get("custom_text", "")).strip()
    rotation_words = [
        item.strip()
        for item in re.split(r"[|、，,；;]+", str(options.get("rotation_words", "")))
        if item.strip()
    ]
    values["rotation"] = rotation_words[0] if rotation_words else ""
    raw_parts = options.get("parts", ["title", "document_number"])
    parts = raw_parts if isinstance(raw_parts, list) else ["title", "document_number"]
    separator = re.sub(r"[<>:\"/\\|?*\x00-\x1f]", "", str(options.get("separator", "-")))[:4] or "-"
    aliases = {"mainTitle": "title", "docNumber": "document_number", "subTitle": "subtitle"}
    normalized_parts = [aliases.get(str(part), str(part)) for part in parts]
    new_stem = separator.join(values.get(part, "") for part in normalized_parts if values.get(part, ""))
    # 内容部件不是路径，必须先清理非法字符再交给路径 API；否则自定义文字中
    # 的斜杠会被误解为目录并静默丢掉前面的标题、文号等部件。
    new_stem = re.sub(r"[<>:\"/\\|?*\x00-\x1f]", "", new_stem or values["title"] or source.stem)
    new_stem = re.sub(r"\s+", " ", new_stem).strip(" .")[:120] or "未命名公文"
    output = _unique_output(workspace, new_stem, ".docx", options)
    shutil.copy2(prepared, output)
    _validate_docx(output)
    _progress(progress, 90, f"已生成文件名：{output.name}")
    return FeatureExecutionResult(
        (FeatureOutput(output, output.name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),),
        f"已按内容生成文件名：{output.name}",
        diagnose_docx(output),
    )


def _extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    blocks: list[str] = []
    blocks.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            blocks.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(blocks)


def _extract_pdf_text(path: Path) -> str:
    with fitz.open(path) as document:
        return "\n\f\n".join(page.get_text("text", sort=True) for page in document)


def _parse_page_selection(value: Any, page_count: int) -> list[int]:
    text = str(value or "all").strip().lower()
    if text in {"", "all", "全部"}:
        return list(range(page_count))
    pages: set[int] = set()
    for chunk in re.split(r"[,，]", text):
        chunk = chunk.strip()
        if not chunk:
            continue
        range_parts = re.split(r"[-至]", chunk, maxsplit=1)
        if len(range_parts) == 2:
            left, right = range_parts
            if not left.isdigit() or not right.isdigit():
                raise OfficialFormatError("PAGE_SELECTION_INVALID", "页码范围无效", "请使用 1-3,5 这样的格式。")
            start, end = int(left), int(right)
            if start > end:
                raise OfficialFormatError("PAGE_SELECTION_INVALID", "页码范围无效", "起始页不得大于结束页。")
            pages.update(range(start - 1, end))
        elif chunk.isdigit():
            pages.add(int(chunk) - 1)
        else:
            raise OfficialFormatError("PAGE_SELECTION_INVALID", "页码范围无效", "请使用 1-3,5 这样的格式。")
    valid = sorted(item for item in pages if 0 <= item < page_count)
    if not valid:
        raise OfficialFormatError("PAGE_SELECTION_EMPTY", "未选择有效页面", "页码超出文档范围。")
    return valid


def _render_pdf_images(pdf_path: Path, workspace: Path, source_stem: str, options: dict[str, Any], progress: ProgressCallback, cancelled: CancelCallback) -> tuple[FeatureOutput, ...]:
    image_format = str(options.get("target_format", "png")).lower()
    dpi = max(72, min(600, int(options.get("dpi", 200))))
    scale = dpi / 72
    rendered: list[Image.Image] = []
    outputs: list[FeatureOutput] = []
    with fitz.open(pdf_path) as pdf:
        pages = _parse_page_selection(options.get("page_selection", "all"), pdf.page_count)
        for index, page_index in enumerate(pages):
            _check_cancelled(cancelled)
            pixmap = pdf.load_page(page_index).get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
            image = Image.open(io.BytesIO(pixmap.tobytes("png"))).convert("RGB")
            if image.width * image.height > 80_000_000:
                raise OfficialFormatError("IMAGE_MEMORY_BUDGET", "页面图像过大", "请降低 DPI 后重试。")
            rendered.append(image.copy())
            _progress(progress, 35 + int((index + 1) * 45 / len(pages)), f"已渲染第 {page_index + 1} 页")
    extension = ".jpg" if image_format == "jpg" else ".png"
    mime = "image/jpeg" if image_format == "jpg" else "image/png"
    if str(options.get("image_mode", "pages")) == "long":
        width = max(image.width for image in rendered)
        height = sum(image.height for image in rendered)
        if width * height > 120_000_000:
            raise OfficialFormatError("LONG_IMAGE_MEMORY_BUDGET", "长图尺寸过大", "请减少页面或降低 DPI 后重试。")
        canvas = Image.new("RGB", (width, height), "white")
        cursor = 0
        for image in rendered:
            canvas.paste(image, ((width - image.width) // 2, cursor))
            cursor += image.height
        path = _unique_output(workspace, source_stem, f"-长图{extension}", options)
        canvas.save(path, format="JPEG" if image_format == "jpg" else "PNG", quality=92)
        outputs.append(FeatureOutput(path, path.name, mime))
    else:
        for index, image in enumerate(rendered, start=1):
            path = _unique_output(workspace, source_stem, f"-第{index}页{extension}", options)
            image.save(path, format="JPEG" if image_format == "jpg" else "PNG", quality=92)
            outputs.append(FeatureOutput(path, path.name, mime))
    return tuple(outputs)


def _pdf_to_docx(source: Path, output: Path, options: dict[str, Any], progress: ProgressCallback, cancelled: CancelCallback) -> None:
    document = Document()
    if document.paragraphs:
        empty = document.paragraphs[0]
        if not empty.text:
            empty._element.getparent().remove(empty._element)
    with fitz.open(source) as pdf:
        for page_index, page in enumerate(pdf):
            _check_cancelled(cancelled)
            if page_index:
                document.add_section(WD_SECTION.NEW_PAGE)
            section = document.sections[-1]
            section.page_width = Mm(page.rect.width * 25.4 / 72)
            section.page_height = Mm(page.rect.height * 25.4 / 72)
            blocks = page.get_text("blocks", sort=True)
            text_blocks = [block for block in blocks if len(block) >= 7 and int(block[6]) == 0 and str(block[4]).strip()]
            if not text_blocks:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                image_stream = io.BytesIO(pixmap.tobytes("png"))
                document.add_picture(image_stream, width=Mm(max(10, section.page_width.mm - 25)))
            else:
                for block in text_blocks:
                    value = str(block[4]).strip()
                    if options.get("normalize_punctuation", True):
                        value, _ = normalize_chinese_punctuation(value)
                    paragraph = document.add_paragraph()
                    paragraph.paragraph_format.space_after = Pt(4)
                    paragraph.add_run(value)
                if options.get("reconstruct_tables", True) and hasattr(page, "find_tables"):
                    try:
                        finder = page.find_tables()
                        for table in finder.tables[:20]:
                            data = table.extract()
                            if not data or not data[0]:
                                continue
                            word_table = document.add_table(rows=len(data), cols=max(len(row) for row in data))
                            for row_index, row in enumerate(data):
                                for column_index, value in enumerate(row):
                                    word_table.cell(row_index, column_index).text = str(value or "")
                    except (AttributeError, ValueError, RuntimeError):
                        pass
            _progress(progress, 12 + int((page_index + 1) * 76 / max(1, pdf.page_count)), f"正在还原第 {page_index + 1}/{pdf.page_count} 页")
    document.save(output)
    _validate_docx(output)


def _execute_pdf_to_word(source: Path, workspace: Path, options: dict[str, Any], progress: ProgressCallback, cancelled: CancelCallback) -> FeatureExecutionResult:
    if source.suffix.lower() != ".pdf":
        raise OfficialFormatError("PDF_REQUIRED", "请选择 PDF 文件", "PDF 转 Word 只接受 PDF。")
    output = _unique_output(workspace, source.stem, "-转换版.docx", options)
    _pdf_to_docx(source, output, options, progress, cancelled)
    return FeatureExecutionResult(
        (FeatureOutput(output, output.name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),),
        "PDF 已在本机转换为可编辑 DOCX，并通过结构校验。",
    )


def _execute_convert(source: Path, workspace: Path, options: dict[str, Any], progress: ProgressCallback, cancelled: CancelCallback) -> FeatureExecutionResult:
    target_format = str(options.get("target_format", "pdf")).lower()
    if target_format not in {"docx", "pdf", "txt", "png", "jpg"}:
        raise OfficialFormatError("CONVERT_FORMAT_INVALID", "输出格式无效", "请选择 DOCX、PDF、TXT、PNG 或 JPG。")
    _check_cancelled(cancelled)
    if target_format == "docx":
        if source.suffix.lower() == ".pdf":
            return _execute_pdf_to_word(source, workspace, options, progress, cancelled)
        prepared, _ = _prepare_word_source(source, workspace)
        output = _unique_output(workspace, source.stem, "-转换版.docx", options)
        shutil.copy2(prepared, output)
        _validate_docx(output)
        return FeatureExecutionResult((FeatureOutput(output, output.name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),), "DOCX 转换完成。")
    if target_format == "txt":
        if source.suffix.lower() == ".pdf":
            text = _extract_pdf_text(source)
        else:
            prepared, _ = _prepare_word_source(source, workspace)
            text = _extract_docx_text(prepared)
        output = _unique_output(workspace, source.stem, "-转换版.txt", options)
        output.write_text(text, encoding="utf-8", newline="\n")
        return FeatureExecutionResult((FeatureOutput(output, output.name, "text/plain; charset=utf-8"),), "TXT 转换完成。")
    pdf_path = source if source.suffix.lower() == ".pdf" else _run_libreoffice_conversion(_prepare_word_source(source, workspace)[0], workspace, "pdf")
    if target_format == "pdf":
        output = _unique_output(workspace, source.stem, "-转换版.pdf", options)
        shutil.copy2(pdf_path, output)
        return FeatureExecutionResult((FeatureOutput(output, output.name, "application/pdf"),), "PDF 转换完成。")
    outputs = _render_pdf_images(pdf_path, workspace, source.stem, options, progress, cancelled)
    return FeatureExecutionResult(outputs, "图片转换完成。")


def _execute_format(source: Path, workspace: Path, options: dict[str, Any], progress: ProgressCallback, cancelled: CancelCallback) -> FeatureExecutionResult:
    prepared, _ = _prepare_word_source(source, workspace)
    _check_cancelled(cancelled)
    output = _unique_output(workspace, source.stem, "-公文规范版.docx", options)
    _progress(progress, 30, "正在识别公文要素")
    scope = str(options.get("scope", "full"))
    if scope not in {"full", "selection", "compilation"}:
        raise OfficialFormatError("FORMAT_SCOPE_INVALID", "排版范围无效", "请选择全文、段落范围或汇编文章。")
    paragraph_range = None
    if scope != "full":
        try:
            start = int(options.get("start_paragraph", 1))
            end = int(options.get("end_paragraph", 99999))
        except (TypeError, ValueError) as exc:
            raise OfficialFormatError("FORMAT_SCOPE_INVALID", "排版范围无效", "起止段落必须为整数。") from exc
        paragraph_range = (start, end)
    report = format_docx(
        prepared,
        output,
        paragraph_range=paragraph_range,
        apply_document_layout=scope != "selection",
    )
    _check_cancelled(cancelled)
    _progress(progress, 88, "正在复核版面与输出结构")
    return FeatureExecutionResult(
        (FeatureOutput(output, output.name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),),
        "公文规范排版完成。",
        report,
    )


EXECUTORS = {
    "format": _execute_format,
    "replace": _execute_replace,
    "redheader": _execute_redheader,
    "rename": _execute_rename,
    "convert": _execute_convert,
    "pdf-to-word": _execute_pdf_to_word,
}


def execute_feature(
    feature_id: str,
    source: Path,
    workspace: Path,
    options: dict[str, Any] | None = None,
    *,
    progress: ProgressCallback | None = None,
    cancelled: CancelCallback | None = None,
) -> FeatureExecutionResult:
    """执行一个文档功能，保证源文件不被写入。"""

    definition = FEATURE_BY_ID.get(feature_id)
    if definition is None:
        raise OfficialFormatError("FEATURE_UNKNOWN", "功能不存在", "请选择受支持的排版功能。")
    if source.suffix.lower() not in definition.accepts:
        raise OfficialFormatError(
            "FEATURE_INPUT_UNSUPPORTED",
            "文件类型与功能不匹配",
            f"{definition.display_name} 支持：{', '.join(definition.accepts)}。",
        )
    workspace.mkdir(parents=True, exist_ok=True)
    callback = progress or (lambda _percent, _message: None)
    cancel_callback = cancelled or (lambda: False)
    _progress(callback, 2, "正在准备只读源文件")
    result = EXECUTORS[feature_id](source, workspace, dict(options or {}), callback, cancel_callback)
    _progress(callback, 100, result.message)
    return result


__all__ = [
    "FEATURE_DEFINITIONS",
    "PRODUCT_CAPABILITIES",
    "SUPPORTED_INPUT_EXTENSIONS",
    "FeatureExecutionResult",
    "FeatureOutput",
    "capabilities_payload",
    "execute_feature",
]

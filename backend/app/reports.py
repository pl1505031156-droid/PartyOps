"""周期报告边界、自动归集与导出。"""

from __future__ import annotations

from datetime import datetime, timedelta, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from sqlalchemy import select
from sqlalchemy.orm import Session

from .config import get_settings
from .enums import PeriodReportStatus, PeriodType, ReportSection, TaskStatus
from .models import PeriodReport, PeriodReportItem, Task, User
from .schemas import PeriodReportItemOut, PeriodReportOut
from .spreadsheet_security import safe_spreadsheet_row
from .task_service import visible_tasks

LOCAL_TIMEZONE = timezone(timedelta(hours=8))
SECTION_LABELS = {
    ReportSection.COMPLETED: "本期完成",
    ReportSection.NEXT_PLAN: "下期计划",
    ReportSection.CARRY_OVER: "延续事项",
    ReportSection.RISK: "重点问题与风险",
    ReportSection.COORDINATION: "需要协调事项",
}


def report_sections(report: PeriodReport) -> list[ReportSection]:
    configured = (report.snapshot or {}).get("design", {}).get("sections", [])
    result: list[ReportSection] = []
    for value in configured:
        try:
            section = ReportSection(value)
        except ValueError:
            continue
        if section not in result:
            result.append(section)
    return result or list(ReportSection)


def _aware(value: datetime | None) -> datetime | None:
    if value is None:
        return None
    return value if value.tzinfo else value.replace(tzinfo=timezone.utc)


def period_bounds(
    period_type: PeriodType, anchor: datetime | None = None
) -> tuple[str, str, datetime, datetime]:
    value = (anchor or datetime.now(LOCAL_TIMEZONE)).astimezone(LOCAL_TIMEZONE)
    if period_type == PeriodType.WEEK:
        start = (value - timedelta(days=value.weekday())).replace(
            hour=0, minute=0, second=0, microsecond=0
        )
        end = start + timedelta(days=7)
        iso = start.isocalendar()
        key = f"{iso.year}-W{iso.week:02d}"
        title = f"{iso.year}年第{iso.week:02d}周工作"
    elif period_type == PeriodType.MONTH:
        start = value.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        if start.month == 12:
            end = start.replace(year=start.year + 1, month=1)
        else:
            end = start.replace(month=start.month + 1)
        key = f"{start.year}-{start.month:02d}"
        title = f"{start.year}年{start.month}月工作"
    elif period_type == PeriodType.QUARTER:
        quarter = (value.month - 1) // 3 + 1
        month = (quarter - 1) * 3 + 1
        start = value.replace(
            month=month, day=1, hour=0, minute=0, second=0, microsecond=0
        )
        end = (
            start.replace(year=start.year + 1, month=1)
            if month == 10
            else start.replace(month=month + 3)
        )
        key = f"{start.year}-Q{quarter}"
        title = f"{start.year}年第{quarter}季度工作"
    else:
        start = value.replace(
            month=1, day=1, hour=0, minute=0, second=0, microsecond=0
        )
        end = start.replace(year=start.year + 1)
        key = str(start.year)
        title = f"{start.year}年度工作"
    return key, title, start.astimezone(timezone.utc), end.astimezone(timezone.utc)


def report_to_out(db: Session, report: PeriodReport) -> PeriodReportOut:
    items = [
        PeriodReportItemOut.model_validate(item)
        for item in db.scalars(
            select(PeriodReportItem)
            .where(PeriodReportItem.report_id == report.id)
            .order_by(PeriodReportItem.section, PeriodReportItem.sort_order, PeriodReportItem.created_at)
        ).all()
    ]
    return PeriodReportOut.model_validate(report).model_copy(update={"items": items})


def auto_fill_report(db: Session, report: PeriodReport, user: User) -> int:
    """同步草稿报告中的任务投影，返回发生变化的条目数。

    已发布、已锁定报告不会调用本函数。任务重新打开、改期或删除时，
    对应草稿条目会被移除或移动栏目，避免周期汇总长期滞后。
    """

    tasks = visible_tasks(db, user)
    existing_items = db.scalars(
        select(PeriodReportItem)
        .where(PeriodReportItem.report_id == report.id)
        .order_by(PeriodReportItem.created_at, PeriodReportItem.id)
    ).all()
    start = _aware(report.start_at)
    end = _aware(report.end_at)
    assert start is not None and end is not None
    next_end = end + (end - start)
    order = max((item.sort_order for item in existing_items), default=-1) + 1
    changed = 0
    expected: dict[str, tuple[ReportSection, Task, bool]] = {}
    for task in tasks:
        section: ReportSection | None = None
        carried_over = False
        completed = _aware(task.completed_at)
        planned = _aware(task.planned_start_at or task.planned_end_at)
        due = _aware(task.internal_due_at or task.formal_due_at)
        if completed and start <= completed < end:
            section = ReportSection.COMPLETED
        elif (
            report.period_type == PeriodType.WEEK
            and task.status not in {TaskStatus.COMPLETED, TaskStatus.ARCHIVED}
            and ((planned and end <= planned < next_end) or (due and end <= due < next_end))
        ):
            section = ReportSection.NEXT_PLAN
        elif (
            task.status not in {TaskStatus.COMPLETED, TaskStatus.ARCHIVED}
            and _aware(task.planned_end_at) is not None
            and _aware(task.planned_end_at) < end
        ):
            section = ReportSection.CARRY_OVER
            carried_over = True
        if section is None:
            continue
        expected[task.id] = (section, task, carried_over)

    # 历史版本曾经在重试或投影恢复时写入重复任务条目。以来源任务为
    # 唯一键清理重复项，保留最早记录；这样旧数据也能在下一次幂等同步
    # 时自动恢复为“一项任务在一个周期内只出现一次”。
    existing_task_items: dict[str, PeriodReportItem] = {}
    for item in existing_items:
        if item.source_type != "task" or not item.source_id:
            continue
        if item.source_id in existing_task_items:
            db.delete(item)
            changed += 1
            continue
        existing_task_items[item.source_id] = item
    for task_id, item in existing_task_items.items():
        projection = expected.pop(task_id, None)
        if projection is None:
            db.delete(item)
            changed += 1
            continue
        section, task, carried_over = projection
        if (
            item.section != section
            or item.title != task.title
            or item.content != task.experience_notes
            or item.carried_over != carried_over
        ):
            item.section = section
            item.title = task.title
            item.content = task.experience_notes
            item.carried_over = carried_over
            item.version += 1
            changed += 1

    for task_id, (section, task, carried_over) in expected.items():
        db.add(
            PeriodReportItem(
                report_id=report.id,
                section=section,
                source_type="task",
                source_id=task.id,
                title=task.title,
                content=task.experience_notes,
                sort_order=order,
                carried_over=carried_over,
                created_by=user.id,
            )
        )
        order += 1
        changed += 1
    return changed


def ensure_period_reports(
    db: Session,
    user: User,
    anchor: datetime | None = None,
) -> tuple[list[PeriodReport], int, int]:
    """确保锚点所在周、月、季度、年度均有报告，并同步仍处于草稿的内容。"""

    reports: list[PeriodReport] = []
    created_count = 0
    added_count = 0
    for period_type in (
        PeriodType.WEEK,
        PeriodType.MONTH,
        PeriodType.QUARTER,
        PeriodType.YEAR,
    ):
        key, title, start, end = period_bounds(period_type, anchor)
        report = db.scalar(
            select(PeriodReport).where(PeriodReport.period_key == key)
        )
        created = report is None
        if report is None:
            report = PeriodReport(
                period_type=period_type,
                period_key=key,
                title=title,
                start_at=start,
                end_at=end,
                summary="",
                snapshot={},
                created_by=user.id,
                updated_by=user.id,
            )
            db.add(report)
            db.flush()
            created_count += 1
        if report.status == PeriodReportStatus.DRAFT:
            added = auto_fill_report(db, report, user)
            added_count += added
            if added and not created:
                report.version += 1
                report.updated_by = user.id
        reports.append(report)
    return reports, created_count, added_count


def report_snapshot(db: Session, report: PeriodReport) -> dict[str, object]:
    items = db.scalars(
        select(PeriodReportItem)
        .where(PeriodReportItem.report_id == report.id)
        .order_by(PeriodReportItem.section, PeriodReportItem.sort_order)
    ).all()
    return {
        "design": (report.snapshot or {}).get("design", {}),
        "period_key": report.period_key,
        "title": report.title,
        "summary": report.summary,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "items": [
            {
                "id": item.id,
                "section": item.section.value,
                "source_type": item.source_type,
                "source_id": item.source_id,
                "title": item.title,
                "content": item.content,
                "carried_over": item.carried_over,
            }
            for item in items
        ],
    }


def _period_export_payload(
    db: Session,
    report: PeriodReport,
) -> tuple[str, str, list[dict[str, object]]]:
    """发布/锁定报告使用不可变快照，草稿才读取实时条目。"""

    snapshot = report.snapshot or {}
    frozen_items = snapshot.get("items")
    if report.status != PeriodReportStatus.DRAFT and isinstance(frozen_items, list):
        return (
            str(snapshot.get("title") or report.title),
            str(snapshot.get("summary") or ""),
            [dict(item) for item in frozen_items if isinstance(item, dict)],
        )
    items = db.scalars(
        select(PeriodReportItem)
        .where(PeriodReportItem.report_id == report.id)
        .order_by(PeriodReportItem.section, PeriodReportItem.sort_order)
    ).all()
    return (
        report.title,
        report.summary,
        [
            {
                "section": item.section.value,
                "source_type": item.source_type,
                "title": item.title,
                "content": item.content,
                "carried_over": item.carried_over,
            }
            for item in items
        ],
    )


def export_period_docx(db: Session, report: PeriodReport) -> Path:
    path = get_settings().exports_dir / f"党建智办-{report.period_key}-{report.id[:8]}.docx"
    document = Document()
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title, summary, items = _period_export_payload(db, report)
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    if summary:
        document.add_paragraph(summary)
    for section in report_sections(report):
        section_items = [item for item in items if item.get("section") == section.value]
        if not section_items:
            continue
        document.add_heading(SECTION_LABELS[section], level=1)
        for item in section_items:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(str(item.get("title") or "")).bold = True
            if item.get("content"):
                paragraph.add_run(f"：{item['content']}")
    document.save(path)
    return path


def export_period_xlsx(db: Session, report: PeriodReport) -> Path:
    path = get_settings().exports_dir / f"党建智办-{report.period_key}-{report.id[:8]}.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = report.period_key[:31]
    sheet.append(["栏目", "事项", "说明", "来源", "是否延续"])
    for cell in sheet[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = PatternFill("solid", fgColor="B42318")
    _title, _summary, items = _period_export_payload(db, report)
    section_order = {section: index for index, section in enumerate(report_sections(report))}
    items.sort(
        key=lambda item: section_order.get(
            ReportSection(str(item.get("section") or ReportSection.COMPLETED.value)),
            999,
        )
    )
    for item in items:
        section = ReportSection(str(item.get("section") or ReportSection.COMPLETED.value))
        sheet.append(
            safe_spreadsheet_row([
                SECTION_LABELS[section],
                item.get("title") or "",
                item.get("content") or "",
                item.get("source_type") or "manual",
                "是" if item.get("carried_over") else "否",
            ])
        )
    for column, width in {"A": 20, "B": 42, "C": 60, "D": 14, "E": 12}.items():
        sheet.column_dimensions[column].width = width
    sheet.freeze_panes = "A2"
    workbook.save(path)
    return path

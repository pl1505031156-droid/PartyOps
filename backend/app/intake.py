"""离线收件解析：文本、Word、PDF 与本地 OCR。"""

from __future__ import annotations

import asyncio
import io
import calendar
import re
import tarfile
import zipfile
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path

import fitz
import pytesseract
from defusedxml import ElementTree
from docx import Document
from fastapi import UploadFile
from openpyxl import load_workbook
from PIL import Image

from .schemas import IntakeCandidate
from .problems import ProblemException
from .compat import to_thread


DATE_PATTERNS = [
    re.compile(r"(?P<year>20\d{2})[年./-](?P<month>\d{1,2})[月./-](?P<day>\d{1,2})日?"),
    re.compile(r"(?P<month>\d{1,2})月(?P<day>\d{1,2})日"),
]
LOCAL_TIMEZONE = timezone(timedelta(hours=8))
MAX_ARCHIVE_MEMBERS = 5_000
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 2 * 1024 * 1024 * 1024
MAX_ARCHIVE_COMPRESSION_RATIO = 200
MAX_INTAKE_UPLOAD_BYTES = 50 * 1024 * 1024
MAX_OFFICE_MEMBERS = 10_000
MAX_OFFICE_UNCOMPRESSED_BYTES = 500 * 1024 * 1024
MAX_PDF_PAGES = 500
MAX_PDF_OCR_PAGES = 20
MAX_IMAGE_PIXELS = 50_000_000
# Pillow 默认会按内容而不是扩展名加载大量插件。收件文件属于不可信输入，
# 因此 PartyOps 只允许 OCR 实际用到的常见位图签名，避免把伪装成 PNG/JPEG
# 的 PSD、FITS、GD、字体或 PDF 交给无关解码器。
_IMAGE_SIGNATURES = (
    b"\x89PNG\r\n\x1a\n",
    b"\xff\xd8\xff",
    b"BM",
    b"II*\x00",
    b"MM\x00*",
)
_PARSE_CONCURRENCY = asyncio.Semaphore(2)
WEEKDAY_MAP = {
    "一": 0,
    "二": 1,
    "三": 2,
    "四": 3,
    "五": 4,
    "六": 5,
    "日": 6,
    "天": 6,
}


def _time_parts(text: str) -> tuple[int, int]:
    match = re.search(r"(?<!\d)(?P<hour>[01]?\d|2[0-3])[:：](?P<minute>[0-5]\d)", text)
    if match:
        return int(match.group("hour")), int(match.group("minute"))
    chinese = re.search(
        r"(?P<period>上午|中午|下午|晚上)?(?P<hour>\d{1,2})[点时](?P<minute>\d{1,2})?分?",
        text,
    )
    if chinese:
        hour = int(chinese.group("hour"))
        if chinese.group("period") in {"下午", "晚上"} and hour < 12:
            hour += 12
        return min(hour, 23), int(chinese.group("minute") or 0)
    return 18, 0


def _extract_date(text: str, anchor: datetime) -> datetime | None:
    local_anchor = anchor.astimezone(LOCAL_TIMEZONE)
    hour, minute = _time_parts(text)
    for pattern in DATE_PATTERNS:
        match = pattern.search(text)
        if not match:
            continue
        year = int(match.groupdict().get("year") or local_anchor.year)
        month = int(match.group("month"))
        day = int(match.group("day"))
        try:
            value = datetime(year, month, day, hour, minute, tzinfo=LOCAL_TIMEZONE)
            if "year" not in match.groupdict() and value.date() < local_anchor.date():
                value = value.replace(year=year + 1)
            return value
        except ValueError:
            continue
    if "明天" in text:
        target = local_anchor + timedelta(days=1)
        return target.replace(hour=hour, minute=minute, second=0, microsecond=0)
    week_match = re.search(r"(本周|下周)([一二三四五六日天])", text)
    if week_match:
        target_weekday = WEEKDAY_MAP[week_match.group(2)]
        days = target_weekday - local_anchor.weekday()
        if week_match.group(1) == "下周":
            days += 7
        elif days < 0:
            days += 7
        target = local_anchor + timedelta(days=days)
        return target.replace(hour=hour, minute=minute, second=0, microsecond=0)
    if "月底" in text:
        day = calendar.monthrange(local_anchor.year, local_anchor.month)[1]
        return local_anchor.replace(
            day=day, hour=hour, minute=minute, second=0, microsecond=0
        )
    return None


def parse_text(text: str, source_kind: str = "text") -> IntakeCandidate:
    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    lines = normalized.splitlines()
    title = lines[0][:240] if lines else "未命名事项"
    for prefix in ("通知：", "关于", "【"):
        if title.startswith(prefix) and len(title) > 2:
            break
    requirements = [
        line[:300]
        for line in lines[1:]
        if any(token in line for token in ("报送", "提交", "材料", "要求", "反馈"))
    ][:10]
    due = _extract_date(normalized, datetime.now(timezone.utc))
    warnings: list[str] = []
    if due is None:
        warnings.append("未识别到明确截止时间，请人工填写。")
    if not normalized:
        warnings.append("未提取到正文，请人工录入。")
    return IntakeCandidate(
        title=title,
        formal_due_at=due,
        requirements=requirements,
        extracted_text=normalized[:30_000],
        source_kind=source_kind,
        warnings=warnings,
    )


def _extract_docx(data: bytes) -> str:
    _validate_office_container(data)
    document = Document(io.BytesIO(data))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(paragraphs)


def _validate_office_container(data: bytes) -> None:
    """在交给 Office 解析库前限制成员数、声明容量和压缩比。"""

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            members = archive.infolist()
            if len(members) > MAX_OFFICE_MEMBERS:
                raise ProblemException(
                    422,
                    "INTAKE_OFFICE_TOO_COMPLEX",
                    "文件结构异常",
                    "Office 文件内部条目过多，已停止解析；仍可在事项中作为原始附件保存。",
                )
            total_size = sum(max(0, item.file_size) for item in members)
            total_compressed = sum(max(1, item.compress_size) for item in members)
            if (
                total_size > MAX_OFFICE_UNCOMPRESSED_BYTES
                or total_size / max(1, total_compressed) > MAX_ARCHIVE_COMPRESSION_RATIO
            ):
                raise ProblemException(
                    422,
                    "INTAKE_OFFICE_EXPANSION_LIMIT",
                    "文件结构异常",
                    "Office 文件解压后容量或压缩比超过安全限制，已停止解析。",
                )
            for item in members:
                if item.is_dir() or not item.filename.lower().endswith(".xml"):
                    continue
                # Office XML 不需要 DTD 或实体。解析前统一拒绝，避免不同底层
                # 库对实体扩展的处理差异形成内存放大或外部实体读取入口。
                with archive.open(item) as source:
                    prefix = source.read(min(item.file_size, 64 * 1024)).upper()
                if b"<!DOCTYPE" in prefix or b"<!ENTITY" in prefix:
                    raise ProblemException(
                        422,
                        "INTAKE_OFFICE_XML_UNSAFE",
                        "文件结构不安全",
                        "Office 文件包含不允许的 XML 实体声明，已停止解析。",
                    )
    except zipfile.BadZipFile as exc:
        raise ProblemException(
            422,
            "INTAKE_OFFICE_INVALID",
            "Office 文件损坏",
            "文件无法安全读取，请确认文件完整；也可以只粘贴通知文字。",
        ) from exc


def _extract_pdf(data: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    pages: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as document:
        if document.needs_pass:
            raise ProblemException(422, "INTAKE_PDF_ENCRYPTED", "PDF 已加密", "请解密后再识别，或仅作为事项材料上传。")
        if document.page_count > MAX_PDF_PAGES:
            raise ProblemException(
                422,
                "INTAKE_PDF_PAGE_LIMIT",
                "PDF 页数过多",
                f"快速识别最多处理 {MAX_PDF_PAGES} 页；大文件可作为事项材料上传。",
            )
        ocr_pages = 0
        for page in document:
            text = page.get_text("text").strip()
            if text:
                pages.append(text)
                continue
            if ocr_pages >= MAX_PDF_OCR_PAGES:
                if not any("停止继续 OCR" in warning for warning in warnings):
                    warnings.append(f"扫描页超过 {MAX_PDF_OCR_PAGES} 页，已停止继续 OCR，请人工核对。")
                continue
            width = max(1, int(page.rect.width * 1.8))
            height = max(1, int(page.rect.height * 1.8))
            if width * height > MAX_IMAGE_PIXELS:
                warnings.append("扫描页像素规模过大，已跳过 OCR。")
                continue
            ocr_pages += 1
            try:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
                image = Image.open(io.BytesIO(pixmap.tobytes("png")))
                pages.append(pytesseract.image_to_string(image, lang="chi_sim", timeout=10))
            except (pytesseract.TesseractNotFoundError, RuntimeError):
                if not any("Tesseract" in warning for warning in warnings):
                    warnings.append("扫描页需要本地 Tesseract 中文 OCR；当前环境未检测到引擎或识别超时。")
    return "\n".join(pages), warnings


def _extract_image(data: bytes) -> tuple[str, list[str]]:
    is_webp = len(data) >= 12 and data.startswith(b"RIFF") and data[8:12] == b"WEBP"
    if not is_webp and not any(data.startswith(signature) for signature in _IMAGE_SIGNATURES):
        return "", ["图片内容与支持的 PNG、JPEG、BMP、TIFF 或 WebP 格式不符，已拒绝解析。"]
    try:
        image = Image.open(io.BytesIO(data))
        if image.width * image.height > MAX_IMAGE_PIXELS:
            return "", ["图片像素规模超过快速识别上限，请压缩后重试。"]
        return pytesseract.image_to_string(image, lang="chi_sim", timeout=20), []
    except pytesseract.TesseractNotFoundError:
        return "", ["当前环境未检测到 Tesseract 中文 OCR，请人工确认图片内容。"]
    except (OSError, RuntimeError):
        return "", ["图片无法识别，请确认文件格式。"]


def _parse_upload_data(data: bytes, safe_filename: str, pasted_text: str) -> IntakeCandidate:
    """同步解析主体由受控工作线程调用，避免 OCR/PDF 阻塞 API 事件循环。"""

    name = safe_filename.lower()
    warnings: list[str] = []
    parser_label = "仅保存原始附件"
    try:
        if name.endswith(".docx"):
            text = _extract_docx(data)
            source_kind = "word"
            parser_label = "Word 文档本地提取"
        elif name.endswith(".pdf"):
            text, warnings = _extract_pdf(data)
            source_kind = "pdf"
            parser_label = "PDF 文本与中文 OCR"
        elif name.endswith((".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp")):
            text, warnings = _extract_image(data)
            source_kind = "image"
            parser_label = "图片中文 OCR"
        else:
            try:
                text = data.decode("utf-8-sig")
                parser_label = "文本文件本地提取"
            except UnicodeDecodeError:
                text = ""
                warnings.append("暂不支持该文件正文提取，仍可作为原始附件保存。")
            source_kind = "file"
    except ProblemException:
        raise
    except Exception:
        # 收件箱的解析器故障只影响候选正文，不能让用户丢失本次录入。
        text = ""
        source_kind = "file"
        warnings.append("文件正文识别未完成，仍可人工填写并把原文件归档。")
    combined = "\n".join(part for part in (pasted_text, text) if part.strip())
    result = parse_text(combined, source_kind)
    if result.title == "未命名事项" and name:
        result.title = Path(name).stem[:240]
    result.warnings.extend(warnings)
    result.source_filename = safe_filename
    result.parser_label = parser_label
    return result


async def parse_upload(upload: UploadFile, pasted_text: str = "") -> IntakeCandidate:
    data = await upload.read(MAX_INTAKE_UPLOAD_BYTES + 1)
    if len(data) > MAX_INTAKE_UPLOAD_BYTES:
        raise ProblemException(
            413,
            "INTAKE_FILE_TOO_LARGE",
            "收件文件过大",
            "快速识别单个文件不超过 50 MB；大文件可在事项创建后作为材料上传。",
        )
    safe_filename = Path(upload.filename or "").name
    async with _PARSE_CONCURRENCY:
        return await to_thread(
            _parse_upload_data,
            data,
            safe_filename,
            pasted_text,
        )


@dataclass(frozen=True)
class PathExtraction:
    text: str = ""
    warnings: tuple[str, ...] = ()
    detected_type: str = "application/octet-stream"
    content_status: str = "metadata_only"
    error_code: str = ""
    archive_member_count: int = 0


def _file_signature(path: Path) -> bytes:
    with path.open("rb") as handle:
        return handle.read(16)


def _detect_type(path: Path, suffix: str) -> str:
    signature = _file_signature(path)
    if signature.startswith(b"PK\x03\x04"):
        return "application/zip"
    if signature.startswith(b"%PDF"):
        return "application/pdf"
    if signature.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"):
        return "application/x-ole-storage"
    if signature.startswith(b"\x89PNG"):
        return "image/png"
    if signature.startswith(b"\xFF\xD8\xFF"):
        return "image/jpeg"
    if signature.startswith((b"II*\x00", b"MM\x00*")):
        return "image/tiff"
    if suffix in {".txt", ".md", ".csv", ".json", ".log"}:
        return "text/plain"
    return "application/octet-stream"


def _extract_spreadsheet(path: Path) -> tuple[str, list[str]]:
    _validate_office_container(path.read_bytes())
    workbook = load_workbook(io.BytesIO(path.read_bytes()), read_only=True, data_only=True)
    rows: list[str] = []
    warnings: list[str] = []
    try:
        for sheet in workbook.worksheets:
            rows.append(f"[{sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value not in (None, "")]
                if values:
                    rows.append(" | ".join(values))
                if sum(len(item) for item in rows) >= 100_000:
                    warnings.append("表格正文较长，索引仅保留前 10 万字。")
                    return "\n".join(rows), warnings
    finally:
        workbook.close()
    return "\n".join(rows), warnings


def _extract_presentation(path: Path) -> tuple[str, list[str]]:
    """不执行宏，仅从 OOXML 演示文稿的幻灯片 XML 中读取文字。"""

    values: list[str] = []
    _validate_office_container(path.read_bytes())
    with zipfile.ZipFile(path) as archive:
        slide_names = sorted(
            name
            for name in archive.namelist()
            if name.startswith("ppt/slides/slide") and name.endswith(".xml")
        )
        if len(slide_names) > 5_000:
            return "", ["演示文稿页数异常，仅保存文件元数据。"]
        for name in slide_names:
            if archive.getinfo(name).file_size > 10 * 1024 * 1024:
                continue
            root = ElementTree.fromstring(archive.read(name))
            texts = [
                node.text.strip()
                for node in root.iter()
                if node.tag.endswith("}t") and node.text and node.text.strip()
            ]
            if texts:
                values.append(" ".join(texts))
            if sum(len(item) for item in values) >= 100_000:
                return "\n".join(values)[:100_000], ["演示文稿正文较长，索引仅保留前 10 万字。"]
    return "\n".join(values), []


def _extract_zip_listing(path: Path) -> tuple[str, list[str], int]:
    names: list[str] = []
    total_size = 0
    total_compressed = 0
    with zipfile.ZipFile(path) as archive:
        members = archive.infolist()
        if len(members) > MAX_ARCHIVE_MEMBERS:
            return "", ["压缩包文件数量超过安全限制，仅保存压缩包元数据。"], len(members)
        for member in members:
            total_size += max(0, member.file_size)
            total_compressed += max(1, member.compress_size)
            if total_size > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
                return "", ["压缩包声明容量超过安全限制，仅保存压缩包元数据。"], len(members)
            if total_size / total_compressed > MAX_ARCHIVE_COMPRESSION_RATIO:
                return "", ["压缩包压缩比异常，仅保存压缩包元数据。"], len(members)
            safe_name = member.filename.replace("\\", "/").lstrip("/")
            if safe_name and ".." not in Path(safe_name).parts:
                names.append(safe_name[:1_000])
    return "\n".join(names)[:100_000], [], len(names)


def _extract_tar_listing(path: Path) -> tuple[str, list[str], int]:
    # 仅处理未压缩 TAR；压缩 TAR 仍会被全类型元数据索引，但不在扫描阶段解压。
    names: list[str] = []
    total_size = 0
    with tarfile.open(path, mode="r:") as archive:
        for index, member in enumerate(archive):
            if index >= MAX_ARCHIVE_MEMBERS:
                return "", ["归档包文件数量超过安全限制，仅保存归档包元数据。"], index + 1
            total_size += max(0, member.size)
            if total_size > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
                return "", ["归档包声明容量超过安全限制，仅保存归档包元数据。"], index + 1
            safe_name = member.name.replace("\\", "/").lstrip("/")
            if safe_name and ".." not in Path(safe_name).parts:
                names.append(safe_name[:1_000])
    return "\n".join(names)[:100_000], [], len(names)


def extract_path_content(
    path: Path,
    maximum_bytes: int = 20 * 1024 * 1024,
    original_name: str | None = None,
) -> PathExtraction:
    """在单文件边界内完成安全识别；任何解析器异常都不会传播到目录扫描。"""

    try:
        if not path.is_file():
            return PathExtraction(
                warnings=("文件当前不可用。",),
                content_status="error",
                error_code="FILE_UNAVAILABLE",
            )
        suffix = (
            Path(original_name).suffix.lower() if original_name else path.suffix.lower()
        )
        detected_type = _detect_type(path, suffix)
        if path.stat().st_size > maximum_bytes and suffix not in {".zip", ".tar"}:
            return PathExtraction(
                warnings=("文件超过正文索引大小限制，仅保存名称和元数据。",),
                detected_type=detected_type,
                content_status="metadata_only",
                error_code="CONTENT_SIZE_LIMIT",
            )

        warnings: list[str] = []
        archive_member_count = 0
        if suffix == ".docx" or (suffix == ".wps" and detected_type == "application/zip"):
            text = _extract_docx(path.read_bytes())
            detected_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif suffix == ".pdf" or detected_type == "application/pdf":
            text, warnings = _extract_pdf(path.read_bytes())
            detected_type = "application/pdf"
        elif suffix in {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}:
            text, warnings = _extract_image(path.read_bytes())
        elif suffix in {".xlsx", ".xlsm"} or (
            suffix == ".et" and detected_type == "application/zip"
        ):
            text, warnings = _extract_spreadsheet(path)
            detected_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif suffix == ".pptx" or (
            suffix == ".dps" and detected_type == "application/zip"
        ):
            text, warnings = _extract_presentation(path)
            detected_type = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
        elif suffix in {".txt", ".md", ".csv", ".json", ".log"}:
            raw = path.read_bytes()
            text = ""
            for encoding in ("utf-8-sig", "gb18030"):
                try:
                    text = raw.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            if not text and raw:
                warnings.append("文本编码无法识别，仅保存文件元数据。")
        elif suffix == ".zip":
            text, warnings, archive_member_count = _extract_zip_listing(path)
            detected_type = "application/zip"
        elif suffix == ".tar":
            text, warnings, archive_member_count = _extract_tar_listing(path)
            detected_type = "application/x-tar"
        else:
            return PathExtraction(
                warnings=("该文件已纳管，当前仅索引文件名称和属性。",),
                detected_type=detected_type,
                content_status="unsupported",
                error_code="CONTENT_UNSUPPORTED",
            )

        if text:
            return PathExtraction(
                text=text[:100_000],
                warnings=tuple(warnings),
                detected_type=detected_type,
                content_status="indexed",
                archive_member_count=archive_member_count,
            )
        if warnings and any("OCR" in warning or "Tesseract" in warning for warning in warnings):
            return PathExtraction(
                warnings=tuple(warnings),
                detected_type=detected_type,
                content_status="pending_ocr",
                error_code="OCR_PENDING",
            )
        if warnings:
            return PathExtraction(
                warnings=tuple(warnings),
                detected_type=detected_type,
                content_status="metadata_only",
                error_code="CONTENT_METADATA_ONLY",
                archive_member_count=archive_member_count,
            )
        return PathExtraction(
            detected_type=detected_type,
            content_status="metadata_only",
            archive_member_count=archive_member_count,
        )
    except Exception:
        # 单个第三方解析器可能抛出 TypeError、XML 解析错误等非标准异常。
        # 对目录扫描而言，这只是一个文件没有正文，不应使整个任务失败。
        return PathExtraction(
            warnings=("正文识别失败，文件名称和属性已经保留。",),
            detected_type="application/octet-stream",
            content_status="error",
            error_code="CONTENT_PARSE_FAILED",
        )


def extract_path_text(
    path: Path,
    maximum_bytes: int = 20 * 1024 * 1024,
    original_name: str | None = None,
) -> tuple[str, list[str]]:
    """只读提取已授权主机文件正文，供文件索引与 AI 最小片段检索复用。"""
    result = extract_path_content(path, maximum_bytes, original_name)
    return result.text, list(result.warnings)

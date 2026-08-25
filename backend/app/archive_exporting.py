"""重要档案年度目录和归档包导出。"""

from __future__ import annotations

import hashlib
import json
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from sqlalchemy import select
from sqlalchemy.orm import Session

from .archive_service import can_view_category, safe_archive_name
from .backups import SCHEMA_VERSION
from .config import get_settings
from .models import ArchiveAttachment, ArchiveCategory, ArchiveRecord, FileBlob, User
from .spreadsheet_security import safe_spreadsheet_row


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while chunk := handle.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def _safe_file_name(value: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "_", value).strip(" .")[:160] or "未命名文件"


def _rows(
    db: Session,
    user: User,
    archive_year: int,
    category_id: str | None,
    keyword: str,
    device_id: str | None = None,
) -> list[tuple[ArchiveRecord, ArchiveCategory]]:
    statement = (
        select(ArchiveRecord, ArchiveCategory)
        .join(ArchiveCategory, ArchiveCategory.id == ArchiveRecord.category_id)
        .where(
            ArchiveRecord.archive_year == archive_year,
            ArchiveRecord.status != "voided",
            ArchiveCategory.active.is_(True),
        )
        .order_by(ArchiveCategory.name, ArchiveRecord.sequence_no)
    )
    if category_id:
        statement = statement.where(ArchiveRecord.category_id == category_id)
    items: list[tuple[ArchiveRecord, ArchiveCategory]] = []
    lowered = keyword.strip().lower()
    for record, category in db.execute(statement).all():
        if not can_view_category(db, category, user, device_id):
            continue
        if lowered and lowered not in record.search_text.lower():
            continue
        items.append((record, category))
    return items


def export_archive_package(
    db: Session,
    user: User,
    archive_year: int,
    category_id: str | None = None,
    keyword: str = "",
    device_id: str | None = None,
) -> Path:
    settings = get_settings()
    rows = _rows(db, user, archive_year, category_id, keyword, device_id)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    output = settings.exports_dir / f"党建智办-{archive_year}年重要档案包-{timestamp}.zip"
    manifest_files: list[dict[str, object]] = []
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "年度档案目录"
    sheet.append(
        [
            "序号",
            "档案类别",
            "文号",
            "文件标题",
            "涉及人员",
            "来源单位",
            "出文时间",
            "附件数量",
            "状态",
            "备注",
        ]
    )
    doc = Document()
    doc.core_properties.title = f"{archive_year}年度重要档案目录"
    heading = doc.add_heading(f"{archive_year}年度重要档案目录", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["序号", "类别", "文号", "文件标题", "涉及人员", "来源单位", "出文时间", "扫描件", "状态"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # 同名扫描件可能来自不同档案，导出时直接保存“目录内路径 -> 真实文件”
    # 映射，避免按文件名反查时误取另一条记录的附件。
    manifest_sources: dict[str, Path] = {}
    for record, category in rows:
        attachments = db.execute(
            select(ArchiveAttachment, FileBlob)
            .join(FileBlob, FileBlob.sha256 == ArchiveAttachment.blob_sha256)
            .where(
                ArchiveAttachment.record_id == record.id,
                ArchiveAttachment.status != "voided",
            )
            .order_by(ArchiveAttachment.version_no)
        ).all()
        date_text = record.document_date.strftime("%Y-%m-%d") if record.document_date else ""
        sheet.append(
            safe_spreadsheet_row([
                record.sequence_no,
                category.name,
                record.document_no,
                record.title,
                "、".join(record.involved_persons or [])
                or record.person_name,
                record.source_unit or record.organization,
                date_text,
                len(attachments),
                record.status.value,
                record.summary,
            ])
        )
        cells = table.add_row().cells
        values = [
            str(record.sequence_no),
            category.name,
            record.document_no,
            record.title,
            "、".join(record.involved_persons or []) or record.person_name,
            record.source_unit or record.organization,
            date_text,
            str(len(attachments)),
            record.status.value,
        ]
        for cell, value in zip(cells, values):
            cell.text = value
        for attachment, blob in attachments:
            source = settings.attachments_dir / blob.relative_path
            if not source.exists():
                continue
            relative = (
                f"扫描件/{safe_archive_name(category.name)}/"
                f"{record.sequence_no:03d}-{safe_archive_name(record.title)}/"
                f"{_safe_file_name(attachment.display_name or blob.original_name)}"
            )
            digest = _sha256(source)
            manifest_files.append(
                {
                    "path": relative,
                    "sha256": digest,
                    "size_bytes": source.stat().st_size,
                    "record_id": record.id,
                    "category": category.name,
                }
            )
            manifest_sources[relative] = source
    for column in sheet.columns:
        sheet.column_dimensions[column[0].column_letter].width = min(
            max(len(str(cell.value or "")) for cell in column) + 2,
            42,
        )
    workbook_path = settings.exports_dir / f".archive-{timestamp}.xlsx"
    docx_path = settings.exports_dir / f".archive-{timestamp}.docx"
    manifest = {
        "format": "partyops-important-archive",
        "version": 1,
        "archive_year": archive_year,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "generated_by": user.id,
        "database_schema_version": SCHEMA_VERSION,
        "record_count": len(rows),
        "files": manifest_files,
    }
    try:
        workbook.save(workbook_path)
        doc.save(docx_path)
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED, allowZip64=True) as archive:
            archive.write(workbook_path, f"{archive_year}年度档案目录.xlsx")
            archive.write(docx_path, f"{archive_year}年度档案目录.docx")
            for item in manifest_files:
                path = manifest_sources.get(str(item["path"]))
                if path and path.exists():
                    archive.write(path, str(item["path"]))
            archive.writestr(
                "manifest.json",
                json.dumps(manifest, ensure_ascii=False, indent=2),
            )
            checksums = [
                f"{item['sha256']}  {item['path']}"
                for item in manifest_files
            ]
            archive.writestr("SHA256SUMS", "\n".join(checksums) + ("\n" if checksums else ""))
    except Exception:
        output.unlink(missing_ok=True)
        raise
    finally:
        workbook_path.unlink(missing_ok=True)
        docx_path.unlink(missing_ok=True)
    return output
